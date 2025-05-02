from docx import Document
from docx.shared import Mm

def extract_doc_content(doc_path):
    # 加载文档
    doc = Document(doc_path)

    # 初试化string
    content = ""

    # 提取并打印所有段落文本
    for para in doc.paragraphs:
        # print(para.text)
        content += para.text + "\n"
    # 提取并打印所有表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # print(para.text)
                    content += para.text + "\n"

    return content

def extract_section_info(doc_path):
    doc = Document(doc_path)
    paper = {

    }
    section = doc.sections[0]
    # 获取分节类型
    section_type = section.start_type
    print(f"Section: {section_type} ({int(section_type)})")
    # 获取纸张信息
    print(f"Page width: {section.page_width}")
    print(f"Page height: {section.page_height}")

    paper["size"] = analysis_paper_size(section.page_width, section.page_height)
    # 获取页边距信息
    paper["margins"] = {
        "top": section.top_margin.cm,
        "bottom": section.bottom_margin.cm,
        "left": section.left_margin.cm,
        "right": section.right_margin.cm
    }
    # 获取页眉页脚信息
    paper["header"] = {
        "top": section.header_distance.cm,
        "bottom": section.footer_distance.cm
    }
    # 获取方向
    paper["orientation"] = "Portrait" if section.orientation == 0 else "Landscape"
    # 添加分节类型信息
    paper["section_type"] = int(section_type)
    return paper
def analysis_paper_size(width, height) -> str:
    # 分析纸张大小
    # 允许误差范围为2mm
    if abs(width - Mm(210)) < Mm(2) and abs(height - Mm(297)) < Mm(2):
        return "A4"
    elif abs(width - Mm(297)) < Mm(2) and abs(height - Mm(420)) < Mm(2):
        return "A3"
    else:
        # 如果无法判断，返回默认值
        return "Unknown"
