import docx
import xml.etree.ElementTree as ET
import json, re, os, zipfile
from backend.preparation.para_type import ParsedParaType, ParagraphManager
from docx.shared import RGBColor
from docx.oxml.ns import qn

# 尝试导入 extract_media 模块，如果不存在，则创建一个空函数
try:
    from backend.preparation.extract_media import add_media_to_manager
except ImportError:
    # 如果模块不存在，创建一个空函数
    def add_media_to_manager(manager, doc_path):
        print("未找到 extract_media 模块，无法提取图片和表格")
        return manager
# 定义工具函数
def get_alignment_string(alignment):
    """获取对齐方式的字符串表示"""
    if alignment == 0:
        return "left"
    elif alignment == 1:
        return "center"
    elif alignment == 2:
        return "right"
    elif alignment == 3:
        return "justify"
    else:
        return "unknown"

def get_alignment_display(alignment_str):
    """获取对齐方式的中文表示"""
    alignment_map = {
        "left": "左对齐",
        "center": "居中",
        "right": "右对齐",
        "justify": "两端对齐",
        "both": "两端对齐",
        "start": "左对齐",
        "end": "右对齐"
    }
    return alignment_map.get(alignment_str.lower(), "未知")

def standardize_color(color_str):
    """
    标准化颜色格式，统一添加#前缀
    """
    if not color_str:
        return None

    # 处理黑色的各种表示形式
    if color_str.lower() in ['auto', '000000', '0', '#0']:
        return 'black'

    # 去掉#前缀后再判断
    clean_color = color_str.lstrip('#').lower()
    if clean_color in ['0', '000000']:
        return 'black'

    # 为其他颜色添加#前缀
    if not color_str.startswith('#'):
        return f'#{color_str}'

    return color_str

def is_font_dict_empty(font_dict):
    """检查字体字典是否为空"""
    if not font_dict:
        return True

    for key, value in font_dict.items():
        if value and value != set():
            return False

    return True

def merge_font_dictionaries(dict1, dict2):
    """
    合并两个字体字典，按优先级选择值而非简单合并集合
    """
    result = {}
    all_keys = set(dict1.keys()) | set(dict2.keys())

    for key in all_keys:
        val1 = dict1.get(key)
        val2 = dict2.get(key)

        # 优先使用非空集合的值，若两者均为集合则根据规则合并
        if isinstance(val1, set) and isinstance(val2, set):
            merged = val1.union(val2)
            # 对特定字段应用规则
            if key in ['bold', 'italic']:
                # 只保留明确设置的值，避免猜测
                if True in merged:
                    result[key] = {True}  # 只有明确设置为True时才保留
                elif False in merged:
                    result[key] = {False}  # 如果有False值，则保留False
                elif not merged:
                    result[key] = {False}  # 如果集合为空，则设置为False
                else:
                    result[key] = {False}  # 其他情况也设置为False
            elif key == 'color':
                # 标准化颜色格式并去重
                standardized = {standardize_color(color) for color in merged if color}
                result[key] = standardized
            elif key == 'size':
                # 如果字体大小有多个值，保留所有值以便后续检查
                result[key] = merged
            else:
                result[key] = merged
        else:
            # 优先使用非空值
            if val1 not in (None, set()):
                result[key] = val1
            else:
                result[key] = val2

    return result

def is_all_caps_string(text):
    """检查文本是否全部是大写字母"""
    # 如果文本为空，返回False
    if not text:
        return False

    # 移除空格和标点符号
    import re
    text = re.sub(r'[\s\.,;:!?\"\'\\[\\]\\(\\)\\{\\}]', '', text)

    # 如果处理后的文本为空，返回False
    if not text:
        return False

    # 检查是否全部是大写字母
    return text.isupper()


def extract_para_format_from_style(style)-> dict:
    # 获取段落样式
    if style.type == 1:
        para_format = style.paragraph_format
        style_name = style.name
        # 创建一个字典来存储当前样式的信息
        current_style_info = {
            'style_name': style_name
        }
        # 获取对齐方式
        alignment = para_format.alignment if para_format.alignment and para_format.alignment is not None else 0
        format_alignment = analysise_alignment(alignment)
        current_style_info['alignment'] = format_alignment
        # 获取首行缩进
        first_line_indent = para_format.first_line_indent if para_format.first_line_indent and para_format.first_line_indent is not None else 0
        current_style_info['first_line_indent'] = first_line_indent
        # 获取左缩进（单位：磅）
        left_indent = para_format.left_indent if para_format.left_indent and para_format.left_indent is not None else 0
        current_style_info['left_indent'] = left_indent
        # 获取右缩进（单位：磅）
        right_indent = para_format.right_indent if para_format.right_indent and para_format.right_indent is not None else 0
        current_style_info['right_indent'] = right_indent
        # 获取段前间距（单位：磅）
        before_spacing = para_format.space_before if para_format.space_before and para_format.space_before is not None else 0
        current_style_info['before_spacing'] = before_spacing
        # 获取段后间距（单位：磅）
        after_spacing = para_format.space_after if para_format.space_after and para_format.space_after is not None else 0
        current_style_info['after_spacing'] = after_spacing
        # 获取行间距
        line_spacing = para_format.line_spacing
        if isinstance(line_spacing, float):
            current_style_info['line_spacing'] = f"{line_spacing} 倍行距"
        elif isinstance(line_spacing, int):
            current_style_info['line_spacing'] = f"{line_spacing / 20} 磅"
        else:
            # 如果行间距没有明确设置，则默认为单倍行距（1.0）
            current_style_info['line_spacing'] = "1.0"
        return current_style_info

def extract_para_format_info_from_paragraph_fromat(para)-> dict:
    # 解析para数据
    current_style_info = {}
    # 初始化段落格式信息
    alignment = None
    first_line_indent = None
    left_indent = None
    right_indent = None
    before_spacing = None
    after_spacing = None
    line_spacing = None

    try:
        # 获取段落xml数据
        para_xml = para._p.xml
        # print(para_xml)
        # 命名空间
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }
        root = ET.fromstring(para_xml)
        pPr = root.find('.//w:pPr', namespaces)
        # 输出节点信息
        # print(ET.tostring(pPr, encoding='utf-8').decode('utf-8'))
        if pPr is not None:
            # 查找缩进节点
            ind = pPr.find('w:ind', namespaces)
            if ind is not None:
                # 使用完整的命名空间URI获取属性
                # print(ET.tostring(ind, encoding='utf-8').decode('utf-8'))
                w_ns = namespaces['w']
                first_line_indent = ind.get(f'{{{w_ns}}}firstLineChars')
                if first_line_indent is not None:
                    first_line_indent = 0 if first_line_indent == 0 else int(first_line_indent) / 100
                # 如果没有首行字符缩进，则获取首行缩进值
                if first_line_indent is None:
                    first_line_indent = ind.get(f'{{{w_ns}}}firstLine')
        else:
            first_line_indent = None

        # 获取段落对齐方式
        if para.paragraph_format.alignment is not None:
            alignment = analysise_alignment(para.paragraph_format.alignment)
        else:
            # 尝试从 XML 中提取对齐方式
            if pPr is not None:
                jc = pPr.find('w:jc', namespaces)
                if jc is not None:
                    w_ns = namespaces['w']
                    val = jc.get(f'{{{w_ns}}}val')
                    # 使用 get_alignment_display 函数将对齐方式转换为中文表示
                    alignment = get_alignment_display(val)

        # 获取左缩进
        if para.paragraph_format.left_indent is not None:
            left_indent = para.paragraph_format.left_indent

        # 获取右缩进
        if para.paragraph_format.right_indent is not None:
            right_indent = para.paragraph_format.right_indent

        # 获取段前间距
        if para.paragraph_format.space_before is not None:
            before_spacing = para.paragraph_format.space_before

        # 获取段后间距
        if para.paragraph_format.space_after is not None:
            after_spacing = para.paragraph_format.space_after

        # 获取行间距
        if para.paragraph_format.line_spacing is not None:
            line_spacing = para.paragraph_format.line_spacing
    except Exception as e:
        print(f"解析段落格式信息出错: {str(e)}")
        # 发生错误时，使用默认值

    current_style_info['alignment'] = alignment
    current_style_info['first_line_indent'] = first_line_indent
    current_style_info['left_indent'] = left_indent
    current_style_info['right_indent'] = right_indent
    current_style_info['before_spacing'] = before_spacing
    current_style_info['after_spacing'] = after_spacing

    # 如果行间距没有明确设置，则默认为单倍行距（1.0）
    if line_spacing is None:
        current_style_info['line_spacing'] = "1.0"
    else:
        current_style_info['line_spacing'] = line_spacing

    return current_style_info

def extract_default_font_size_from_styles(docx_path):
    """
    从styles.xml中提取默认样式的字体大小信息
    """
    default_sizes = {
        'paragraph': None,  # 段落默认字体大小
        'character': None,  # 字符默认字体大小
        'table': None,      # 表格默认字体大小
        'numbering': None   # 编号默认字体大小
    }

    try:
        # 打开docx文件（实际上是一个zip文件）
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 检查styles.xml是否存在
            styles_path = 'word/styles.xml'
            if styles_path in docx_zip.namelist():
                # 读取styles.xml内容
                with docx_zip.open(styles_path) as styles_file:
                    styles_content = styles_file.read()

                # 解析XML
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                root = ET.fromstring(styles_content)

                # 首先从docDefaults中提取默认字体大小
                # 这是文档的全局默认设置
                doc_defaults = root.find('.//w:docDefaults', namespaces)
                if doc_defaults is not None:
                    # 查找默认运行属性
                    rPr_default = doc_defaults.find('.//w:rPrDefault/w:rPr', namespaces)
                    if rPr_default is not None:
                        # 查找默认字体大小
                        sz_element = rPr_default.find('./w:sz', namespaces)
                        if sz_element is not None:
                            sz_val = sz_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if sz_val:
                                try:
                                    # 字号单位是半点，需要除以2转换为磅
                                    font_size = float(sz_val) / 2
                                    # 将默认字体大小应用于所有类型
                                    for key in default_sizes:
                                        default_sizes[key] = font_size
                                    print(f"从docDefaults提取到默认字体大小: {font_size}pt")
                                except (ValueError, TypeError):
                                    pass

                        # 查找默认复杂脚本字体大小
                        szCs_element = rPr_default.find('./w:szCs', namespaces)
                        if szCs_element is not None:
                            szCs_val = szCs_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if szCs_val:
                                try:
                                    # 字号单位是半点，需要除以2转换为磅
                                    font_size = float(szCs_val) / 2
                                    # 如果还没有设置默认字体大小，则使用复杂脚本字体大小
                                    for key in default_sizes:
                                        if default_sizes[key] is None:
                                            default_sizes[key] = font_size
                                    print(f"从docDefaults提取到默认复杂脚本字体大小: {font_size}pt")
                                except (ValueError, TypeError):
                                    pass

                # 如果从docDefaults中没有找到默认字体大小，则继续查找默认样式
                if all(size is None for size in default_sizes.values()):
                    # 查找所有默认样式
                    default_styles = root.findall('.//w:style[@w:default="1"]', namespaces)
                    for style in default_styles:
                        # 获取样式类型
                        style_type = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')

                        # 查找字体大小信息
                        sz_element = style.find('.//w:sz', namespaces)
                        if sz_element is not None and '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val' in sz_element.attrib:
                            sz_val = sz_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            try:
                                # 字号单位是半点，需要除以2转换为磅
                                font_size = float(sz_val) / 2
                                if style_type in default_sizes:
                                    default_sizes[style_type] = font_size
                                    print(f"从styles.xml提取到{style_type}类型的默认字体大小: {font_size}pt")
                            except (ValueError, TypeError):
                                pass

                        # 如果没有找到sz，尝试查找szCs（复杂脚本字体大小）
                        if style_type in default_sizes and default_sizes[style_type] is None:
                            szCs_element = style.find('.//w:szCs', namespaces)
                            if szCs_element is not None and '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val' in szCs_element.attrib:
                                szCs_val = szCs_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                try:
                                    font_size = float(szCs_val) / 2
                                    default_sizes[style_type] = font_size
                                    print(f"从styles.xml提取到{style_type}类型的默认复杂脚本字体大小: {font_size}pt")
                                except (ValueError, TypeError):
                                    pass

                    # 如果没有找到默认样式，尝试查找名为Normal的样式
                    if default_sizes['paragraph'] is None:
                        # 修复XPath查询语法
                        for style in root.findall('.//w:style', namespaces):
                            # 查找style下的name元素
                            name_elem = style.find('./w:name', namespaces)
                            if name_elem is not None:
                                # 获取name元素的val属性
                                val_attr = name_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val_attr == 'Normal':
                                    # 找到Normal样式
                                    normal_style = style
                                    # 查找字体大小信息
                                    sz_element = normal_style.find('.//w:sz', namespaces)
                                    if sz_element is not None and '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val' in sz_element.attrib:
                                        sz_val = sz_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                        try:
                                            font_size = float(sz_val) / 2
                                            default_sizes['paragraph'] = font_size
                                            print(f"从styles.xml的Normal样式中提取到默认字体大小: {font_size}pt")
                                        except (ValueError, TypeError):
                                            pass

                                    # 如果还是没有找到，尝试查找szCs
                                    if default_sizes['paragraph'] is None:
                                        szCs_element = normal_style.find('.//w:szCs', namespaces)
                                        if szCs_element is not None and '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val' in szCs_element.attrib:
                                            szCs_val = szCs_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                            try:
                                                font_size = float(szCs_val) / 2
                                                default_sizes['paragraph'] = font_size
                                                print(f"从styles.xml的Normal样式中提取到默认复杂脚本字体大小: {font_size}pt")
                                            except (ValueError, TypeError):
                                                pass
                                    break

                print(f"从styles.xml提取的默认字体大小: {default_sizes}")
    except Exception as e:
        print(f"从styles.xml提取字体大小信息时出错: {str(e)}")

    return default_sizes

def extract_para_format_info(doc_path, manager: ParagraphManager):
    """
    从Word文档中提取段落格式信息，并将段落及其元数据添加到段落管理器中

    参数:
        doc_path: Word文档路径
        manager: 段落管理器实例

    返回:
        更新后的段落管理器实例
    """
    # 加载文档
    doc = docx.Document(doc_path)

    # 获取所有样式信息
    styles_info = [extract_para_format_from_style(style) for style in doc.styles if style.type == 1]
    print(f"从文档中提取到 {len(styles_info)} 个段落样式")

    # 从theme1.xml中提取字体信息
    theme_fonts = extract_font_from_theme(doc_path)
    print(f"从theme1.xml提取的字体信息：中文字体={theme_fonts['zh_family']}, 英文字体={theme_fonts['en_family']}")

    # 从styles.xml中提取默认字体大小信息
    default_font_sizes = extract_default_font_size_from_styles(doc_path)
    print(f"从styles.xml提取的默认字体大小：{default_font_sizes}")

    # 提取图片和表格信息
    manager = add_media_to_manager(manager, doc_path)
    print("已从文档中提取图片和表格信息")

    # 预处理段落，将摘要等信息提取出来
    processed_paras = pre_process_paragraphs(doc)
    print(f"预处理后的段落数量：{len(processed_paras)}")

    # 处理进度计数
    total_paras = len([p for p in processed_paras if p.text.strip()])
    processed_count = 0

    # 遍历每个段落
    for para in processed_paras:
        # 如果段落文本为空，则跳过
        if not para.text.strip():
            continue

        processed_count += 1
        if processed_count % 10 == 0 or processed_count == total_paras:
            print(f"处理进度：{processed_count}/{total_paras} ({int(processed_count/total_paras*100)}%)")

        # 创建段落的元数据字典
        meta_data = {}
        para_text_preview = para.text[:30] + ('...' if len(para.text) > 30 else '')

        # 获取段落和样式的XML数据
        style_xml_data = para.style.element.xml if hasattr(para.style, 'element') else None
        xml_data = para._p.xml if hasattr(para, '_p') else None

        # 打印段落信息
        print(f"\n处理段落 {processed_count}/{total_paras}: '{para_text_preview}'")
        print(f"  样式名称: {para.style.name}")

        # 解析段落格式
        para_format_info = extract_para_format_info_from_paragraph_fromat(para)
        style_info = find_style_by_name(para.style.name, styles_info)

        # 合并段落格式信息（段落格式优先，样式为补充）
        current_para_format_info = {}
        for key in para_format_info.keys():
            if para_format_info[key] is not None:
                current_para_format_info[key] = para_format_info[key]
                print(f"  使用段落格式的{key}: {para_format_info[key]}")
            elif style_info and key in style_info:
                current_para_format_info[key] = style_info[key]
                print(f"  使用样式的{key}: {style_info[key]}")
            else:
                current_para_format_info[key] = None
                print(f"  未找到{key}信息")

        # 将段落格式信息添加到元数据
        meta_data["paragraph_format"] = current_para_format_info
        print(f"  段落格式信息: {current_para_format_info}")

        # 初始化字体信息字典
        fonts = {
            'zh_family': set(),
            'en_family': set(),
            'size': set(),
            'color': set(),
            'bold': set(),  # 默认为False
            'italic': set(),  # 默认为False
            'isAllcaps': None
        }

        # 优先从 runs 提取字体信息
        fonts_from_runs = extract_font_info_from_runs(para.runs)
        print(f"  从runs中提取的字体信息: {fonts_from_runs}")

        # 解析XML中的字体信息
        fonts_from_xml = extract_font_info_from_xml(xml_data)
        if is_font_dict_empty(fonts_from_xml) and style_xml_data:
            fonts_from_xml = extract_font_info_from_xml(style_xml_data)
            print(f"  从样式XML中提取的字体信息: {fonts_from_xml}")
        else:
            print(f"  从段落XML中提取的字体信息: {fonts_from_xml}")

        # 合并字体信息（run优先，XML为补充）
        # 当runs中有有效数据时，优先使用runs中的数据
        if not is_font_dict_empty(fonts_from_runs):
            # 特殊处理加粗和斜体属性
            # 如果段落中只有部分文本是加粗或斜体，我们应该保留这些信息
            # 但不应该将整个段落标记为加粗或斜体
            if 'bold' in fonts_from_runs and True in fonts_from_runs['bold']:
                # 如果段落中有加粗文本，但不是全部文本都是加粗的
                # 我们应该将这个段落标记为"部分加粗"
                print(f"  段落中包含加粗文本")
            else:
                # 如果段落中没有加粗文本，我们应该将加粗属性设置为空集合
                fonts_from_runs['bold'] = set()
                print(f"  段落中不包含加粗文本")

            if 'italic' in fonts_from_runs and True in fonts_from_runs['italic']:
                # 如果段落中有斜体文本，但不是全部文本都是斜体的
                # 我们应该将这个段落标记为"部分斜体"
                print(f"  段落中包含斜体文本")
            else:
                # 如果段落中没有斜体文本，我们应该将斜体属性设置为空集合
                fonts_from_runs['italic'] = set()
                print(f"  段落中不包含斜体文本")

            fonts = merge_font_dictionaries(fonts_from_runs, fonts_from_xml)
            print(f"  优先使用runs中的字体信息进行合并")
        else:
            # 当runs中无有效数据时，使用XML中的数据
            fonts = fonts_from_xml
            print(f"  使用XML中的字体信息")

        print(f"  合并后的字体信息: {fonts}")

        # 处理中文字体信息
        if not fonts['zh_family'] or 'Unknown' in fonts['zh_family']:
            if theme_fonts['zh_family']:
                # 移除Unknown
                if 'Unknown' in fonts['zh_family']:
                    fonts['zh_family'].remove('Unknown')
                    # print(f"  移除了'Unknown'中文字体标记")

                # 添加theme中的中文字体
                original_zh_fonts = set(fonts['zh_family'])
                fonts['zh_family'].update(theme_fonts['zh_family'])

                if fonts['zh_family'] != original_zh_fonts:
                    print(f"  从theme1.xml补充中文字体: {theme_fonts['zh_family']}")

        # 处理英文字体信息
        if not fonts['en_family'] or 'Unknown' in fonts['en_family']:
            if theme_fonts['en_family']:
                # 移除Unknown
                if 'Unknown' in fonts['en_family']:
                    fonts['en_family'].remove('Unknown')
                    print(f"  移除了'Unknown'英文字体标记")

                # 添加theme中的英文字体
                original_en_fonts = set(fonts['en_family'])
                fonts['en_family'].update(theme_fonts['en_family'])

                if fonts['en_family'] != original_en_fonts:
                    print(f"  从theme1.xml补充英文字体: {theme_fonts['en_family']}")

        # 处理字体大小信息
        if not fonts['size'] or 'Unknown' in fonts['size']:
            # 获取段落默认字体大小
            default_size = default_font_sizes.get('paragraph')
            if default_size is not None:
                # 移除Unknown
                if 'Unknown' in fonts['size']:
                    fonts['size'].remove('Unknown')
                    # print(f"  移除了'Unknown'字体大小标记")

                # 添加默认字体大小
                original_sizes = set(fonts['size'])
                fonts['size'].add(default_size)

                if fonts['size'] != original_sizes:
                    print(f"  从styles.xml补充默认字体大小: {default_size}pt")

        # 最终的字体信息汇总
        print(f"  最终字体信息: zh_family={fonts['zh_family']}, en_family={fonts['en_family']}, "
              f"size={fonts['size']}, bold={fonts['bold']}, italic={fonts['italic']}, color={fonts['color']}")

        # 将字体信息添加到元数据
        meta_data["fonts"] = fonts

        # 将段落信息添加到段落管理器
        # 尝试确定段落类型
        try:
            # 确保我们使用的是枚举实例而不是枚举类
            para_type = ParsedParaType.BODY  # 默认为正文类型
            # print(f"  设置默认段落类型: {para_type}, 类型: {type(para_type)}")
        except Exception as e:
            # print(f"  获取默认段落类型失败: {str(e)}")
            # 使用字符串值作为备选
            para_type = 'body'
            # print(f"  使用字符串值'body'作为备选段落类型")

        # 检查是否是分节符或分页符
        if hasattr(para, '_p') and para._p is not None:
            # 检查是否包含分节符或分页符标记
            xml_str = para._p.xml if hasattr(para._p, 'xml') else ""
            if 'sectPr' in xml_str or 'w:br w:type="page"' in xml_str:
                try:
                    # 确保我们使用的是枚举实例而不是枚举类
                    para_type = ParsedParaType.OTHERS
                    # print(f"  检测到分节符或分页符，使用OTHERS类型: {para_type}, 类型: {type(para_type)}")
                except Exception as e:
                    # print(f"  获取OTHERS段落类型失败: {str(e)}")
                    # 使用字符串值作为备选
                    para_type = 'others'
                    # print(f"  使用字符串值'others'作为备选段落类型")

        # 确保para_type是ParsedParaType枚举类型
        if not isinstance(para_type, ParsedParaType):
            print(f"  警告：段落类型不是ParsedParaType枚举，而是{type(para_type)}，将使用BODY类型")
            # 尝试获取BODY枚举值
            try:
                # 直接使用枚举值
                para_type = 'body'  # 使用字符串值，让ParaInfo.__post_init__处理转换
                print(f"  使用字符串值'body'作为段落类型")
            except Exception as e:
                print(f"  获取BODY类型失败: {str(e)}")
                # 最后的尝试：使用枚举值字符串
                para_type = 'body'

        try:
            # 添加段落到管理器
            manager.add_para(
                para_type=para_type,
                content=para.text,
                meta=meta_data
            )
            print(f"  已将段落添加到管理器: id=para{len(manager.paragraphs)-1}")
        except Exception as e:
            print(f"  添加段落到管理器时出错: {str(e)}")
            # 尝试使用字符串值作为备选
            try:
                print(f"  尝试使用字符串值'others'作为备选段落类型")
                manager.add_para(
                    para_type='others',  # 使用字符串值，让ParaInfo.__post_init__处理转换
                    content=para.text,
                    meta=meta_data
                )
                print(f"  使用字符串值'others'添加段落成功: id=para{len(manager.paragraphs)-1}")
            except Exception as e2:
                print(f"  使用字符串值'others'添加段落失败: {str(e2)}")

                # 最后尝试使用'body'字符串值
                try:
                    print(f"  尝试使用字符串值'body'作为最后的备选")
                    manager.add_para(
                        para_type='body',  # 使用字符串值，让ParaInfo.__post_init__处理转换
                        content=para.text,
                        meta=meta_data
                    )
                    print(f"  使用字符串值'body'添加段落成功: id=para{len(manager.paragraphs)-1}")
                except Exception as e3:
                    print(f"  使用字符串值'body'添加段落也失败: {str(e3)}, 错误类型: {type(e3)}")

    print(f"\n文档处理完成，共处理 {processed_count} 个段落")
    return manager

def extract_font_info_from_runs(runs: list) -> dict:
    """
    从文档的runs中提取字体信息
    优化版本：直接从XML获取字体信息，避免基于名称的正则误判
    修复加粗和斜体检测问题：只有当明确设置为True时才添加
    """
    fonts = {
        'zh_family': set(),
        'en_family': set(),
        'size': set(),
        'color': set(),
        'bold': {False},  # 默认为False
        'italic': {False},  # 默认为False
        'isAllcaps': None
    }

    # 如果没有runs，返回空字典
    if not runs:
        return fonts

    # 设置默认值
    has_color_info = False
    has_size_info = False

    # 跟踪每个run的加粗和斜体状态
    run_bold_states = []
    run_italic_states = []

    for run in runs:
        if not run.text.strip():
            continue

        if run.font:
            # 提取字体名称 - 优化：不再使用正则判断中英文字体，而是通过XML获取
            font_name = run.font.name
            if font_name:
                # 优先从 XML 判断中英文字体，避免基于名称的正则误判
                fonts['en_family'].add(font_name)  # 暂时添加，后续由 XML 覆盖

            # 尝试获取字体大小
            if hasattr(run.font, 'size') and run.font.size is not None:
                font_size = run.font.size
                if font_size:
                    # 转换为磅
                    if hasattr(font_size, 'pt'):
                        fonts['size'].add(font_size.pt)
                    else:
                        # 如果没有pt属性，尝试直接获取数值并除以20（docx中字号单位转换）
                        try:
                            size_value = float(font_size) / 20
                            fonts['size'].add(size_value)
                        except (ValueError, TypeError):
                            pass
                    has_size_info = True

            # 从 XML 获取更准确的字体信息
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # 获取中英文字体
                    if hasattr(rPr, 'rFonts'):
                        rFonts = rPr.rFonts
                        if rFonts is not None:
                            # 获取中文字体
                            east_asia_font = rFonts.get(qn('w:eastAsia'))
                            if east_asia_font is not None:
                                fonts['zh_family'].add(east_asia_font)
                            # 获取英文字体
                            ascii_font = rFonts.get(qn('w:ascii'))
                            if ascii_font is not None:
                                fonts['en_family'].add(ascii_font)

                    # 获取字体大小
                    if not has_size_info and hasattr(rPr, 'sz'):
                        if hasattr(rPr.sz, 'val'):
                            sz_val = rPr.sz.val
                        elif hasattr(rPr.sz, 'get'):
                            sz_val = rPr.sz.get(qn('w:val'))
                        else:
                            sz_val = None
                        if sz_val is not None:
                            try:
                                fonts['size'].add(float(sz_val) / 2)  # 半点转为磅
                                has_size_info = True
                            except (ValueError, TypeError, AttributeError):
                                pass

                    # 获取加粗信息 - 只有当明确设置为True时才添加
                    if hasattr(rPr, 'b'):
                        b_val = rPr.b.get(qn('w:val')) if hasattr(rPr.b, 'get') else None
                        # 只有当b_val为None或不等于'0'时，才认为是加粗
                        is_bold = b_val is None or b_val != '0'
                        if is_bold:
                            fonts['bold'].add(True)

                    # 获取斜体信息 - 只有当明确设置为True时才添加
                    if hasattr(rPr, 'i'):
                        i_val = rPr.i.get(qn('w:val')) if hasattr(rPr.i, 'get') else None
                        # 只有当i_val为None或不等于'0'时，才认为是斜体
                        is_italic = i_val is None or i_val != '0'
                        if is_italic:
                            fonts['italic'].add(True)

                    # 获取颜色信息
                    if hasattr(rPr, 'color'):
                        if hasattr(rPr.color, 'get'):
                            color_val = rPr.color.get(qn('w:val'))
                        elif hasattr(rPr.color, 'val'):
                            color_val = rPr.color.val
                        else:
                            color_val = None
                        if color_val is not None:
                            fonts['color'].add(standardize_color(color_val))
                            has_color_info = True

            # 使用 python-docx API 补充信息
            if run.font.color and run.font.color.rgb:
                rgb_str = str(run.font.color.rgb)
                fonts['color'].add('black' if rgb_str == '000000' else rgb_str)
                has_color_info = True

            # 提取加粗信息 - 使用python-docx API
            # 记录每个run的加粗状态
            run_bold_state = None
            if run.font.bold is True:
                run_bold_state = True
            elif hasattr(run, 'bold') and run.bold is True:  # 有些版本可能直接在run上有bold属性
                run_bold_state = True
            elif run.font.bold is False or (hasattr(run, 'bold') and run.bold is False):
                run_bold_state = False

            # 只有当明确设置了加粗时才记录
            if run_bold_state is not None:
                run_bold_states.append((run.text, run_bold_state))

            # 提取斜体信息 - 使用python-docx API
            # 记录每个run的斜体状态
            run_italic_state = None
            if run.font.italic is True:
                run_italic_state = True
            elif hasattr(run, 'italic') and run.italic is True:  # 有些版本可能直接在run上有italic属性
                run_italic_state = True
            elif run.font.italic is False or (hasattr(run, 'italic') and run.italic is False):
                run_italic_state = False

            # 只有当明确设置了斜体时才记录
            if run_italic_state is not None:
                run_italic_states.append((run.text, run_italic_state))

            # 检查全大写
            if fonts['isAllcaps'] is None:
                fonts['isAllcaps'] = is_all_caps_string(run.text)

    # 处理收集到的加粗和斜体状态
    # 只有当段落中有明确设置为加粗的文本时，才将加粗状态添加到字体信息中
    if run_bold_states:
        # 计算加粗文本的总长度
        bold_text_length = sum(len(text) for text, is_bold in run_bold_states if is_bold)
        # 计算所有文本的总长度
        total_text_length = sum(len(text) for text, _ in run_bold_states)

        # 只有当有加粗文本时，才添加加粗状态
        if bold_text_length > 0:
            fonts['bold'].add(True)
            print(f"  段落中有加粗文本，加粗文本长度: {bold_text_length}, 总文本长度: {total_text_length}")
        else:
            # 确保加粗状态为空集合，并添加False值
            fonts['bold'] = {False}
            print(f"  段落中没有加粗文本，设置加粗状态为False")
    else:
        # 如果没有收集到任何加粗状态，确保加粗状态为False
        fonts['bold'] = {False}
        print(f"  未收集到加粗状态信息，设置加粗状态为False")

    # 只有当段落中有明确设置为斜体的文本时，才将斜体状态添加到字体信息中
    if run_italic_states:
        # 计算斜体文本的总长度
        italic_text_length = sum(len(text) for text, is_italic in run_italic_states if is_italic)
        # 计算所有文本的总长度
        total_text_length = sum(len(text) for text, _ in run_italic_states)

        # 只有当有斜体文本时，才添加斜体状态
        if italic_text_length > 0:
            fonts['italic'].add(True)
            print(f"  段落中有斜体文本，斜体文本长度: {italic_text_length}, 总文本长度: {total_text_length}")
        else:
            # 确保斜体状态为空集合，并添加False值
            fonts['italic'] = {False}
            print(f"  段落中没有斜体文本，设置斜体状态为False")
    else:
        # 如果没有收集到任何斜体状态，确保斜体状态为False
        fonts['italic'] = {False}
        print(f"  未收集到斜体状态信息，设置斜体状态为False")

    # 颜色默认值
    if not has_color_info:
        # 默认黑色
        fonts['color'].add('black')

    # 如果没有获取到字体大小信息，不设置为Unknown，而是保持为空集合
    # 这样在后续处理中可以从styles.xml中获取默认字体大小

    # 查找中文内容
    has_chinese_content = any(re.search(r'[\u4e00-\u9fa5]', run.text) for run in runs if run.text.strip())

    # 如果没有找到中文字体，但有中文内容，则尝试从段落样式中获取
    if not fonts['zh_family'] and has_chinese_content:
        # 检查段落样式中是否有中文字体设置
        for run in runs:
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # 尝试从段落样式中获取中文字体
                    if hasattr(run.font, '_element') and hasattr(run.font._element, 'rPr'):
                        style_rPr = run.font._element.rPr
                        if style_rPr is not None and hasattr(style_rPr, 'rFonts'):
                            style_rFonts = style_rPr.rFonts
                            if style_rFonts is not None:
                                east_asia_font = style_rFonts.get(qn('w:eastAsia'))
                                if east_asia_font is not None:
                                    fonts['zh_family'].add(east_asia_font)
                                    break

        # 只有在确实无法找到任何字体且有中文内容时才标记为Unknown
        # 否则保持为空集合，以便后续处理可以从theme1.xml中获取字体信息
        if not fonts['zh_family'] and has_chinese_content:
            fonts['zh_family'].add("Unknown")

    return fonts

def extract_font_from_theme(docx_path):
    """
    从theme1.xml中提取中文和英文字体信息
    """
    fonts = {
        'zh_family': set(),
        'en_family': set()
    }

    # 定义常见的中文字体列表
    common_cn_fonts = [
        '宋体',   # 宋体
        '黑体',   # 黑体
        '楷体',   # 楷体
        '仿宋',   # 仿宋
        '华文中宋', # 华文中宋
        '华文宋体', # 华文宋体
        '华文楷体', # 华文楷体
        '华文仿宋', # 华文仿宋
        '华文黑体', # 华文黑体
        '微软雅黑', # 微软雅黑
        '等线',    # 等线 (DengXian) - 同时支持中英文的字体
        '等线 Light', # 等线 Light
        '方正',    # 方正系列字体
        '思源',    # 思源系列字体
        '苹方',    # 苹方 (PingFang)
        '雅黑'     # 微软雅黑的简称
    ]

    # 定义双语字体列表（同时支持中英文的字体）
    bilingual_fonts = [
        '等线',     # 等线 (DengXian)
        '等线 Light', # 等线 Light
        '微软雅黑',   # 微软雅黑 (Microsoft YaHei)
        '思源黑体',   # 思源黑体 (Source Han Sans)
        '思源宋体',   # 思源宋体 (Source Han Serif)
        '苹方',      # 苹方 (PingFang)
        'Microsoft YaHei', # 微软雅黑英文名
        'DengXian',  # 等线英文名
        'SimSun',    # 宋体英文名
        'SimHei',    # 黑体英文名
        'KaiTi',     # 楷体英文名
        'FangSong'   # 仿宋英文名
    ]

    try:
        # 打开docx文件（实际上是一个zip文件）
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 检查theme1.xml是否存在
            theme_path = 'word/theme/theme1.xml'
            if theme_path in docx_zip.namelist():
                # 读取theme1.xml内容
                with docx_zip.open(theme_path) as theme_file:
                    theme_content = theme_file.read()
                    root = ET.fromstring(theme_content)

                # 解析XML
                namespaces = {
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                }

                # 提取中文字体 (Hans)
                for font in root.findall('.//a:font[@script="Hans"]', namespaces):
                    if 'typeface' in font.attrib:
                        typeface = font.get('typeface')
                        if typeface in common_cn_fonts or any(cn_font in typeface for cn_font in common_cn_fonts):
                            fonts['zh_family'].add(typeface)
                            print(f"从theme1.xml提取到Hans中文字体: {typeface}")

                # 如果没有找到Hans字体，尝试从东亚字体中获取
                if not fonts['zh_family']:
                    # 查找东亚字体(ea)
                    ea_fonts = root.findall('.//a:ea', namespaces)
                    for ea in ea_fonts:
                        if 'typeface' in ea.attrib:
                            typeface = ea.get('typeface')
                            if typeface in common_cn_fonts or any(cn_font in typeface for cn_font in common_cn_fonts):
                                fonts['zh_family'].add(typeface)
                                print(f"从theme1.xml提取到东亚字体: {typeface}")

                # 如果还是没有找到，尝试从fontScheme中获取
                if not fonts['zh_family']:
                    font_scheme = root.find('.//a:fontScheme', namespaces)
                    if font_scheme is not None:
                        # 查找中文默认字体
                        default_font = font_scheme.find('.//a:font[@script="Hans"]', namespaces)
                        if default_font is not None and 'typeface' in default_font.attrib:
                            typeface = default_font.get('typeface')
                            if typeface in common_cn_fonts or any(cn_font in typeface for cn_font in common_cn_fonts):
                                fonts['zh_family'].add(typeface)
                                print(f"从fontScheme提取到Hans中文字体: {typeface}")

                # 提取英文字体 (Latn 或 latin)
                for font in root.findall('.//a:font[@script="Latn"]', namespaces):
                    if 'typeface' in font.attrib:
                        typeface = font.get('typeface')
                        fonts['en_family'].add(typeface)
                        print(f"从theme1.xml提取到Latn英文字体: {typeface}")
                        # 如果是双语字体，也添加到中文字体集合中
                        if typeface in bilingual_fonts or any(bi_font in typeface for bi_font in bilingual_fonts):
                            fonts['zh_family'].add(typeface)
                            print(f"将双语字体添加到中文字体集合: {typeface}")

                # 如果没有找到Latn字体，尝试从latin节点获取
                if not fonts['en_family']:
                    for latin in root.findall('.//a:latin', namespaces):
                        if 'typeface' in latin.attrib:
                            typeface = latin.get('typeface')
                            fonts['en_family'].add(typeface)
                            print(f"从theme1.xml提取到latin字体: {typeface}")
                            # 如果是双语字体，也添加到中文字体集合中
                            if typeface in bilingual_fonts or any(bi_font in typeface for bi_font in bilingual_fonts):
                                fonts['zh_family'].add(typeface)
                                print(f"将双语字体添加到中文字体集合: {typeface}")

                # 如果还是没有找到，尝试从ascii节点获取
                if not fonts['en_family']:
                    for ascii_font in root.findall('.//a:font[@script="Ascii"]', namespaces):
                        if 'typeface' in ascii_font.attrib:
                            typeface = ascii_font.get('typeface')
                            fonts['en_family'].add(typeface)
                            print(f"从theme1.xml提取到Ascii英文字体: {typeface}")
                            # 如果是双语字体，也添加到中文字体集合中
                            if typeface in bilingual_fonts or any(bi_font in typeface for bi_font in bilingual_fonts):
                                fonts['zh_family'].add(typeface)
                                print(f"将双语字体添加到中文字体集合: {typeface}")

                # 最后一次尝试，查找majorFont和minorFont中的typeface
                if not fonts['en_family']:
                    for font_type in ['a:majorFont', 'a:minorFont']:
                        font_elem = root.find(f'.//a:fontScheme/{font_type}', namespaces)
                        if font_elem is not None:
                            latin = font_elem.find('./a:latin', namespaces)
                            if latin is not None and 'typeface' in latin.attrib:
                                typeface = latin.get('typeface')
                                fonts['en_family'].add(typeface)
                                print(f"从{font_type}提取到英文字体: {typeface}")
                                # 如果是双语字体，也添加到中文字体集合中
                                if typeface in bilingual_fonts or any(bi_font in typeface for bi_font in bilingual_fonts):
                                    fonts['zh_family'].add(typeface)
                                    print(f"将双语字体添加到中文字体集合: {typeface}")

                print(f"从theme1.xml提取的字体: 中文字体={fonts['zh_family']}, 英文字体={fonts['en_family']}")
    except Exception as e:
        print(f"从theme1.xml提取字体信息时出错: {str(e)}")

    return fonts

def extract_font_info_from_xml(xml_data):
    fonts = {
        'zh_family': set(),
        'en_family': set(),
        'size': set(),
        'color': set(),
        'bold': set(),
        'italic': set(),
        'isAllcaps': None
    }

    if not xml_data:
        return fonts

    try:
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        root = ET.fromstring(xml_data)

        has_color_info = False
        found_bold = False
        found_italic = False

        # 提取字体信息
        for rFonts in root.findall('.//w:rFonts', namespaces):
            east_asia = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
            if east_asia:
                fonts['zh_family'].add(east_asia)
            ascii = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
            if ascii:
                fonts['en_family'].add(ascii)

        # 提取字体大小
        for sz in root.findall('.//w:sz', namespaces) + root.findall('.//w:szCs', namespaces):
            val = sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if val:
                try:
                    fonts['size'].add(float(val) / 2)
                except (ValueError, TypeError):
                    pass

        # 提取颜色
        for color in root.findall('.//w:color', namespaces):
            val = color.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if val:
                fonts['color'].add(standardize_color(val))
                has_color_info = True

        # 提取加粗 - 只有当明确设置为True时才添加
        for bold in root.findall('.//w:b', namespaces):
            val = bold.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            # 只有当val为None或不等于'0'时，才认为是加粗
            is_bold = val is None or val != '0'
            if is_bold:
                fonts['bold'].add(True)
                found_bold = True

        # 提取斜体 - 只有当明确设置为True时才添加
        for italic in root.findall('.//w:i', namespaces):
            val = italic.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            # 只有当val为None或不等于'0'时，才认为是斜体
            is_italic = val is None or val != '0'
            if is_italic:
                fonts['italic'].add(True)
                found_italic = True

        # 提取全大写
        caps = root.find('.//w:caps', namespaces)
        if caps:
            val = caps.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            fonts['isAllcaps'] = (val is None or val != '0')

        # 颜色默认值仅在需要时添加
        if not has_color_info and root.find('.//w:t', namespaces) is not None:
            # 只有当段落有文本内容时才添加默认颜色
            fonts['color'].add('black')

        # 如果没有找到加粗设置，则设置为False
        if not found_bold:
            fonts['bold'] = {False}

        # 如果没有找到斜体设置，则设置为False
        if not found_italic:
            fonts['italic'] = {False}

    except Exception as e:
        print(f"从 XML 提取字体信息时出错: {str(e)}")
        # 添加默认颜色
        fonts['color'].add('black')
        # 确保加粗和斜体设置为False
        fonts['bold'] = {False}
        fonts['italic'] = {False}

    return fonts

def analysise_alignment(alignment: int) -> str:
    """分析对齐方式，返回中文表示

    参数:
    alignment: 对齐方式枚举值

    返回:
    str: 对齐方式的中文表示
    """
    # 如果对齐方式为 None，默认为左对齐
    if alignment is None:
        return "左对齐"

    # 先使用 get_alignment_string 函数获取对齐方式的字符串表示
    alignment_str = get_alignment_string(alignment)

    return alignment_str

def find_style_by_name(style_name, styles_info):
    """根据样式名称查找样式信息"""
    for style_info in styles_info:
        if style_info and 'style_name' in style_info and style_info['style_name'] == style_name:
            return style_info
    # 如果没有找到匹配的样式，返回默认值
    return {
        'style_name': 'Default',
        'alignment': '左对齐',
        'first_line_indent': 0,
        'left_indent': 0,
        'right_indent': 0,
        'before_spacing': 0,
        'after_spacing': 0,
        'line_spacing': '1.0 倍行距'
    }

def split_paragraph(paragraph, split_pos):
    """在指定位置分割段落"""
    if split_pos <= 0 or split_pos >= len(paragraph.text):
        return None

    # 获取分割点前后的文本
    first_part = paragraph.text[:split_pos]
    second_part = paragraph.text[split_pos:]

    # 更新原段落的文本
    paragraph.clear()
    paragraph.add_run(first_part)

    # 创建新段落并插入到原段落后
    parent = paragraph._parent
    new_para = parent.add_paragraph()
    new_para.style = paragraph.style

    # 将新段落移动到原段落后
    paragraph._p.addnext(new_para._p)

    # 复制段落格式
    para_format = paragraph.paragraph_format
    new_format = new_para.paragraph_format
    for attr in ['alignment', 'left_indent', 'right_indent', 'first_line_indent',
                 'line_spacing', 'space_before', 'space_after', 'line_spacing_rule']:
        setattr(new_format, attr, getattr(para_format, attr))

    # 添加第二部分文本
    new_para.add_run(second_part)

    return new_para

def split_paragraph(paragraph, split_pos):
    # 获取段落所有 run 并计算总文本
    runs = paragraph.runs
    full_text = ''.join([run.text for run in runs])

    # 如果文本为空或分割位置无效，则返回 None
    if not full_text or split_pos <= 0 or split_pos >= len(full_text):
        return None

    # 找到分割点所在的 run 和偏移量
    current_pos = 0
    split_run_idx = 0
    split_run_offset = 0
    for idx, run in enumerate(runs):
        run_len = len(run.text)
        if current_pos + run_len > split_pos:
            split_run_idx = idx
            split_run_offset = split_pos - current_pos
            break
        current_pos += run_len
    else:
        return None  # 分割位置在段落末尾无需处理

    # 创建新段落并插入到原段落后
    parent = paragraph._parent
    new_para = parent.add_paragraph()
    new_para.style = paragraph.style

    # 将新段落移动到原段落后
    paragraph._p.addnext(new_para._p)

    # 复制段落格式
    para_format = paragraph.paragraph_format
    new_format = new_para.paragraph_format
    for attr in ['alignment', 'left_indent', 'right_indent', 'first_line_indent',
                 'line_spacing', 'space_before', 'space_after', 'line_spacing_rule']:
        setattr(new_format, attr, getattr(para_format, attr))

    # 处理 run 的分割并复制字体样式
    def copy_run_style(source, target):
        target.font.bold = source.font.bold
        target.font.italic = source.font.italic
        target.font.underline = source.font.underline
        if source.font.color.rgb:
            target.font.color.rgb = source.font.color.rgb
        target.font.name = source.font.name
        if source.font.size:
            target.font.size = source.font.size

        # 修复FutureWarning
        if (hasattr(source.font, '_element') and
            hasattr(source.font._element, 'rPr') and
            source.font._element.rPr is not None and
            hasattr(source.font._element.rPr, 'rFonts') and
            source.font._element.rPr.rFonts is not None):

            target.font._element.rPr.rFonts.set(qn('w:eastAsia'),
                source.font._element.rPr.rFonts.get(qn('w:eastAsia')))

    # 分割 run
    split_run = runs[split_run_idx]
    text_before = split_run.text[:split_run_offset]
    text_after = split_run.text[split_run_offset:]

    # 修改原 run 的文本
    split_run.text = text_before

    # 将分割后的内容和后续 run 转移到新段落
    if text_after:
        new_run = new_para.add_run(text_after)
        copy_run_style(split_run, new_run)

    for run in runs[split_run_idx + 1:]:
        new_run = new_para.add_run(run.text)
        copy_run_style(run, new_run)
        paragraph._p.remove(run._r)  # 从原段落移除 run

    return new_para

# 预处理段落，将摘要等信息提取出来
def pre_process_paragraphs(doc):
    """
    预处理文档段落，将包含摘要/关键词的段落拆分为独立段落
    """
    while True:
        processed = False
        # 使用副本避免修改时影响迭代
        for para in list(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            # 匹配中英文关键词（摘要、Abstract、关键词、Keywords）
            pattern = re.compile(r'\b(摘要|Abstract|关键词|Keywords)\b\s*[:：]', re.IGNORECASE)
            match = pattern.search(text)
            if not match:
                continue

            # 根据匹配的关键词设置段落类型
            # keyword = match.group(1).lower()
            # if keyword == "摘要":
            #     para_type = ParsedParaType.ABSTRACT_ZH
            # elif keyword == "abstract":
            #     para_type = ParsedParaType.ABSTRACT_EN
            # elif keyword == "关键词":
            #     para_type = ParsedParaType.KEYWORDS_ZH
            # elif keyword == "keywords":
            #     para_type = ParsedParaType.KEYWORDS_EN

            # print(f"正在处理段落: {text}")
            # 计算分割位置（关键词后的位置）
            split_pos = match.end()

            # 分割段落
            new_para = split_paragraph(para, split_pos)
            if new_para:
                # print(f"分割后新段落: {new_para.text}")
                processed = True
        if not processed:
            break
    return doc.paragraphs


