from flask import Flask, request, jsonify, send_from_directory, send_file, Response
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from werkzeug.utils import secure_filename
from typing import List, Dict, Optional
import os, time, sys
import json
import sys
import os

# 添加项目根目录到系统路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from backend.agents.setting import LLMs
from backend.editors.format_editor import generate_formatted_doc
from backend.editors.document_marker import mark_document_errors
from backend.preparation.para_type import ParagraphManager
from backend.checkers.checker import check_format
from datetime import datetime
import backend.preparation.docx_parser as docx_parser
from backend.agents.advice_agent import AdviceAgent
from backend.agents.editor_agent import EditorAgent
from backend.agents.format_agent import FormatAgent
from backend.agents.communicate_agent import CommunicateAgent
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn  # 导入qn函数，用于XML命名空间
import tempfile
from backend.utils.utils import parse_llm_json_response

# 导入agents包中的功能
import backend.agents as agents

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*")

# sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 创建一个全局的Agent中心
agents_config = {
    "format_model": "qwen-plus",
    "editor_model": "deepseek-r1",
    "advice_model": "deepseek-r1",
    "communicate_model": "deepseek-r1"
}
agents = {
    "format": FormatAgent(agents_config["format_model"]),
    "editor": EditorAgent(agents_config["editor_model"]),
    "advice": AdviceAgent(agents_config["advice_model"]),
    "communicate": CommunicateAgent(agents_config["communicate_model"])
}

llm = LLMs()

# 配置处理的ParagraphManager和文件名
analysised_para_manager = []

# 配置上传文件夹
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
CACHES_FOLDER = os.path.join(os.path.dirname(__file__), 'caches')
ALLOWED_EXTENSIONS = {'docx'}  # 允许的文件类型
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CACHES_FOLDER'] = CACHES_FOLDER

# 确保上传目录和缓存目录存在
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(CACHES_FOLDER):
    os.makedirs(CACHES_FOLDER)

# 配置允许的文件类型
ALLOWED_EXTENSIONS = {'docx'}

# 初始化全局变量
page_title = "文档格式分析系统"
messages = []
execution_steps = [
    {"id": 1, "text": "上传文件", "status": "pending"},
    {"id": 2, "text": "分析文档格式", "status": "pending"},
    {"id": 3, "text": "检查格式规范", "status": "pending"},
    {"id": 4, "text": "生成分析报告", "status": "pending"}
]


# 检查文件类型是否允许
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 文件上传路由
@app.route('/api/upload-files', methods=['POST'])
def upload_file():
    try:
        # 检查是否有文件在请求中
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '没有上传文件'
            }), 400

        file = request.files['file']

        # 检查文件名是否为空或文件类型是否允许
        if file.filename == '':
            return jsonify({
                'success': False,
                'message': '没有选择文件'
            }), 400

        if file and allowed_file(file.filename):
            # 确保文件名安全并生成唯一文件名
            filename = secure_filename(file.filename)
            unique_filename = f"{os.path.splitext(filename)[0]}_{int(time.time())}{os.path.splitext(filename)[1]}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename) + ".docx"  # 文件名唯一

            # 保存文件
            file.save(file_path)

            # 返回成功响应
            return jsonify({
                'success': True,
                'message': '文件上传成功',
                'file': {
                    'filename': unique_filename,
                    'originalname': filename,
                    'path': file_path,
                }
            }), 200
        else:
            return jsonify({
                'success': False,
                'message': '只支持上传 .docx 文件'
            }), 400

    except Exception as e:
        print(f"文件上传错误: {str(e)}")
        return jsonify({
            'success': False,
            'message': f"文件上传失败: {str(e)}"
        }), 500

# 获取所有代理的模型配置
@app.route('/api/agent-models')
def get_agent_models():
    try:
        # 获取当前代理配置
        config = {
            "format_model": agents["format"].model,
            "editor_model": agents["editor"].model,
            "advice_model": agents["advice"].model,
            "communicate_model": agents["communicate"].model
        }
        return jsonify(config)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 设置指定代理的模型
@app.route('/api/set-agent-model', methods=['POST'])
def set_agent_model():
    data = request.get_json()
    if not data or 'agent_type' not in data or 'model_name' not in data:
        return jsonify({'error': '未提供代理类型或模型名称'}), 400

    agent_type = data['agent_type']
    model_name = data['model_name']

    if agent_type not in ["format", "editor", "advice", "communicate"]:
        return jsonify({'error': '无效的代理类型'}), 400

    try:
        # 检查模型是否存在于配置中
        if model_name not in llm.models_config:
            return jsonify({'error': f'模型 {model_name} 不存在于配置中'}), 400

        # 完全重新初始化代理，使用LLMs.set_model来更新所有配置
        if agent_type == "format":
            agents[agent_type] = FormatAgent(model_name)
        elif agent_type == "editor":
            agents[agent_type] = EditorAgent(model_name)
        elif agent_type == "advice":
            agents[agent_type] = AdviceAgent(model_name)
        elif agent_type == "communicate":
            agents[agent_type] = CommunicateAgent(model_name)

        return jsonify({'message': f'已将{agent_type}代理的模型完全更新为 {model_name}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# 获取所有可用模型
@app.route('/api/models')
def get_models():
    try:
        # 使用全局的LLMs实例
        models_config = llm.models_config
        return jsonify({"models": models_config})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 添加新模型
@app.route('/api/add-model', methods=['POST'])
def add_model():
    data = request.get_json()
    required_fields = ['name', 'base_url', 'api_key', 'model_name']
    if not data or not all(field in data for field in required_fields):
        return jsonify({'error': '缺少必要的模型信息'}), 400

    try:
        # 使用全局的LLMs实例
        llm.add_model(data['name'], data['base_url'], data['api_key'], data['model_name'])
        return jsonify({'message': f'模型 {data["name"]} 添加成功'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 删除模型
@app.route('/api/delete-model/<model_name>', methods=['DELETE'])
def delete_model(model_name):
    try:
        # 使用全局的LLMs实例
        llm.delete_model(model_name)
        return jsonify({'message': f'模型 {model_name} 删除成功'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 获取完整模型配置
@app.route('/keys.json')
def get_keys():
    try:
        # 使用全局的LLMs实例
        models_config = llm.models_config
        return jsonify(models_config)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 创建配置文件
@app.route('/api/create-config', methods=['POST'])
def create_config():
    data = request.get_json()
    if not data or 'config_path' not in data or 'config_data' not in data:
        return jsonify({"error": "缺少必要参数"}), 400

    try:
        config_path = data['config_path']
        config_data = data['config_data']

        # 确保目录存在
        os.makedirs(os.path.dirname(config_path), exist_ok=True)

        # 保存配置文件
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, ensure_ascii=False, indent=4)

        return jsonify({
            "success": True,
            "message": f"配置文件已保存到 {config_path}"
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"发生错误：{str(e)}"
        }), 500

# 获取示例配置文件
@app.route('/api/get-config-example')
def get_config_example():
    try:
        config_path = os.path.join(os.path.dirname(__file__), 'utils/config_example.json')
        if not os.path.exists(config_path):
            return jsonify({"error": "示例配置文件不存在"}), 404

        with open(config_path, 'r', encoding='utf-8') as f:
            config_data = json.load(f)

        return jsonify({
            "success": True,
            "config": config_data
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"发生错误：{str(e)}"
        }), 500

# 获取配置文件
@app.route('/api/get-config')
def get_config():
    try:
        config_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '../config.json'))  # 配置文件路径
        if not os.path.exists(config_path):
            return jsonify({"error": "配置文件不存在"}), 404

        with open(config_path, 'r', encoding='utf-8') as f:
            config_data = json.load(f)

        return jsonify(config_data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 保存配置文件
@app.route('/api/set-config', methods=['POST'])
def set_config():
    data = request.get_json()
    if not data:
        return jsonify({'error': '配置数据不能为空'}), 400

    try:
        config_path = 'config.json'  # 配置文件路径
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)

        return jsonify({'message': '配置保存成功'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 上传格式文件
@app.route('/api/upload-format', methods=['POST'])
def upload_format():
    if 'file' not in request.files:
        return jsonify({"success": False, "message": "未找到上传的文件"}), 400

    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({"success": False, "message": "未选择文件"}), 400

    try:
        filename = secure_filename(uploaded_file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

        if file_ext == 'json':
            # 直接保存JSON文件
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename) + '.json'
            uploaded_file.save(file_path)
            return jsonify({
                "success": True,
                "message": "格式Json上传成功",
                "file_path": file_path,
                "filename": filename
            })
        elif file_ext == 'docx':
            # 处理docx格式文件
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            uploaded_file.save(file_path)

            # 使用format_agent处理docx文件，提取格式信息
            format_json = llm.parse_format(file_path, '{}')

            # 保存提取的格式信息
            json_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename.rsplit('.', 1)[0]}_format.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                f.write(format_json)

            return jsonify({
                "success": True,
                "message": "格式文件上传并处理成功",
                "file_path": json_path,
                "filename": f"{filename.rsplit('.', 1)[0]}_format.json"
            })
        else:
            return jsonify({
                "success": False,
                "message": f"不支持的文件类型，仅支持 .docx 或 .json 文件"
            }), 400
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"格式文件上传失败: {str(e)}"
        }), 500

# 分析格式要求文档
@app.route('/api/analyse-format-doc', methods=['POST'])
def analyse_format_doc():
    if 'file' not in request.files:
        return jsonify({"success": False, "message": "未找到上传的文件"}), 400

    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({"success": False, "message": "未选择文件"}), 400

    try:
        if not allowed_file(uploaded_file.filename):
            return jsonify({"success": False, "message": "不支持的文件类型"}), 400
        filename = secure_filename(uploaded_file.filename)

        # 生成唯一文件名
        unique_filename = f"{os.path.splitext(filename)[0]}_{int(time.time())}.docx"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)

        # 保存文件
        uploaded_file.save(file_path)

        # 使用 FormatAgent 解析文档
        doc_content = docx_parser.extract_doc_content(file_path)
        format_agent = agents["format"]

        # 解析格式要求
        # 正确读取config.json文件
        config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config.json')
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config_json = json.load(f)
            # 将配置文件传递给parse_format方法
            format_json_str = format_agent.parse_format(doc_content, json.dumps(config_json))
            # 处理返回的原始字符串

        except FileNotFoundError:
            # 如果文件不存在，使用空对象
            print(f"Warning: Config file not found at {config_path}")
            format_json_str = format_agent.parse_format(doc_content, "{}")

        # 使用新的解析函数来处理大模型返回的JSON
        format_json = parse_llm_json_response(format_json_str)

        # 将解析结果保存到config_test.json文件中
        test_config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config_test.json')
        with open(test_config_path, 'w', encoding='utf-8') as f:
            json.dump(format_json, f, ensure_ascii=False, indent=4)

        # 生成唯一的缓存文件名
        cache_filename = f"format_{int(time.time())}.json"
        cache_path = os.path.join(app.config['CACHES_FOLDER'], cache_filename)

        # 保存解析结果到缓存文件夹
        with open(cache_path, 'w', encoding='utf-8') as f:
            json.dump(format_json, f, ensure_ascii=False, indent=4)

        return jsonify({
            "success": True,
            "message": "格式要求文档解析成功",
            "format": format_json,
            "file_path": file_path,
            "cache_path": cache_path
        })
    except Exception as e:
        print(f"解析格式要求文档错误: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"解析格式要求文档失败: {str(e)}"
        }), 500

# 使用默认格式
@app.route('/api/use-default-format', methods=['GET'])
def use_default_format():
    try:
        # 获取默认格式文件路径
        default_format_path = os.path.join(os.path.dirname(__file__), 'utils', 'config_example.json')

        # 确保文件存在
        if not os.path.exists(default_format_path):
            return jsonify({
                "success": False,
                "message": "默认格式文件不存在"
            }), 404

        # 读取默认配置文件内容
        with open(default_format_path, 'r', encoding='utf-8') as f:
            config_data = json.load(f)

        return jsonify({
            "success": True,
            "message": "已应用默认格式",
            "config_path": default_format_path,
            "config_data": config_data
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"应用默认格式失败: {str(e)}"
        }), 500

@app.route('/api/get-advice', methods=['POST'])
def process_file():
    try:
        data = request.get_json()
        if not data or 'doc_path' not in data:
            return jsonify({"error": "缺少必要参数"}), 400
        file_path = data.get('doc_path')
        # 处理文件信息
        doc_content = docx_parser.extract_doc_content(file_path)
        # 初始化advice_agent
        advice_agent = AdviceAgent('qwen-plus')
        # 传送给格式建议
        result = advice_agent.get_advice(doc_content)
        return jsonify({"success": True,
                        "result": result})
    except Exception as e:
        return jsonify({"success": False,
                        "message": str(e)}), 500
# 格式修改Agent
# @app.route('/api/use-format-agent', methods=['POST'])
# def use_format_agent():
#     try:
#         data = request.get_json()
#         if not data or 'messages' not in data:
#             return jsonify({"error": "缺少必要参数"}), 400


#         return jsonify({"success": True,
#                         "result": result})
#     except Exception as e:
#         return jsonify({"success": False,
#                         "message": str(e)}), 500
# 检查文档格式
@app.route('/api/check-format', methods=['POST'])
def start_check_format():
    try:
        data = request.get_json()
        # 修复逻辑错误，使用正确的条件判断
        if not data or ('doc_path' not in data or 'config_path' not in data):
            return jsonify({"success": False, "message": "缺少必要参数"}), 400

        # 获取文件路径
        file_path = data.get('doc_path')
        config_path = data.get('config_path')

        # 验证文件存在
        if not os.path.exists(file_path):
            return jsonify({"success": False, "message": "文档文件不存在"}), 404

        if not os.path.exists(config_path):
            return jsonify({"success": False, "message": "配置文件不存在"}), 404

        # 处理文件信息
        try:
            docx_errors, para_manager = check_format(file_path, config_path, agents["format"])

            # 检查返回值是否有效
            if para_manager is None:
                para_manager = ParagraphManager()

            # 存储当前的para_manager
            analysised_para_manager.append({"doc_path": file_path, "para_manager": para_manager})
        except Exception as e:
            print(f"检查格式时出错: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"检查格式时出错: {str(e)}"
            }), 500

        # 将文档抽取为JSON格式，便于调试
        try:
            doc_json = para_manager.to_dict()
            # 保存抽取的JSON到缓存文件
            json_cache_path = os.path.join(app.config['CACHES_FOLDER'], f"doc_json_{int(time.time())}.json")
            with open(json_cache_path, 'w', encoding='utf-8') as f:
                json.dump(doc_json, f, ensure_ascii=False, indent=4)
        except Exception as json_err:
            print(f"文档JSON抽取错误: {str(json_err)}")

        # 将para_manager转换为可序列化的字典
        para_manager_dict = para_manager.to_dict() if para_manager else {}

        # 以json格式返回错误信息和para_manager
        if docx_errors:
            return jsonify({
                "success": True,
                "message": "分析成功",
                "errors": docx_errors,  # 修改键名与前端一致
                "para_manager": para_manager_dict  # 返回para_manager字典
            })
        else:
            return jsonify({
                "success": True,
                "message": "未发现错误",
                "errors": [],  # 返回空数组而不是None
                "para_manager": para_manager_dict  # 返回para_manager字典
            })
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"格式检查失败: {str(e)}"
        }), 500

# 生成分析报告
@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    try:
        data = request.get_json()
        if not data or 'doc_path' not in data:
            return jsonify({"success": False, "message": "缺少必要参数"}), 400

        # 获取文件路径
        file_path = data.get('doc_path')
        errors = data.get('errors', [])
        original_filename = data.get('original_filename', '')

        # 验证文件存在
        if not os.path.exists(file_path):
            return jsonify({"success": False, "message": "文档文件不存在"}), 404

        # 为了保持与之前功能兼容，如果没有提供错误列表，则调用check_format
        if not errors and 'config_path' in data:
            config_path = data.get('config_path')
            if not config_path or not os.path.exists(config_path):
                return jsonify({"success": False, "message": "配置文件不存在"}), 404
            errors, _ = check_format(file_path, config_path, agents["format"])

            # 确保errors是列表类型
            if errors is None:
                errors = []
            elif isinstance(errors, str):
                # 如果返回的是字符串，将其作为单个错误消息添加到列表中
                errors = [{"message": errors, "location": "未知"}]
            elif not isinstance(errors, list):
                # 如果不是列表类型，尝试转换或创建新列表
                try:
                    errors = list(errors)
                except:
                    errors = [{"message": str(errors), "location": "未知"}]

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(file_path)

        # 创建临时文件用于存储标记错误的文档
        temp_dir = tempfile.gettempdir()
        doc_name = os.path.basename(file_path)
        marked_doc_path = os.path.join(temp_dir, f"marked_{doc_name}")

        # 打印错误列表，便于调试
        print(f"错误列表: {errors}")
        print(f"错误数量: {len(errors)}")

        # 使用新的mark_document_errors函数标记错误
        try:
            # 查找对应的para_manager
            para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == file_path), None)
            if not para_manager:
                # 如果没有找到对应的para_manager，创建一个新的
                para_manager = ParagraphManager()
                print(f"未找到对应的para_manager，创建了新的实例")

            print(f"开始标记文档错误: {file_path}")
            print(f"错误数量: {len(errors)}")

            # 确保输出目录存在
            os.makedirs(os.path.dirname(marked_doc_path), exist_ok=True)

            # 调用mark_document_errors函数标记错误
            marked_doc_path = mark_document_errors(file_path, errors, para_manager, marked_doc_path)
            print(f"文档错误标记完成，保存到: {marked_doc_path}")
        except Exception as e:
            print(f"标记文档错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({"success": False, "message": f"标记文档错误: {str(e)}"}), 500

        # 生成分析报告文本
        report = "文档格式分析报告\n\n"
        report += f"分析时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += f"文档: {os.path.basename(file_path)}\n\n"

        # 添加错误摘要
        if errors:
            report += f"发现 {len(errors)} 个格式问题:\n\n"
            for i, error in enumerate(errors):
                if isinstance(error, dict):
                    msg = error.get('message', '未知错误')
                    location = error.get('location', '')
                    report += f"{i+1}. {msg}"
                    if location:
                        report += f" (位置: {location})"
                else:
                    report += f"{i+1}. {str(error)}"
                report += "\n"
        else:
            report += "没有发现格式问题。\n"

        return jsonify({
            "success": True,
            "message": "报告生成成功",
            "report": report,
            "marked_doc_path": marked_doc_path
        })
    except Exception as e:
        print(f"报告生成失败: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"报告生成失败: {str(e)}"
        }), 500

# 生成报告
@app.route('/api/download-report', methods=['POST'])
def download_report():
    try:
        data = request.get_json()
        if not data or 'doc_path' not in data:
            return jsonify({"success": False, "message": "缺少文档路径参数"}), 400

        # 获取文件路径
        file_path = data.get('doc_path')
        errors = data.get('errors', [])
        original_filename = data.get('original_filename', '')

        # 验证文件存在
        if not os.path.exists(file_path):
            return jsonify({"success": False, "message": "文档文件不存在"}), 404

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(file_path)

        # 打印错误列表，便于调试
        print(f"报告错误列表: {errors}")
        print(f"报告错误数量: {len(errors)}")

        # 创建一个新的Word文档作为报告
        doc = Document()

        # 添加标题 - 使用更通用的方式
        title_para = doc.add_paragraph('文档格式分析报告')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)
        # 设置字体
        title_run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加基本信息
        doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n文档: {original_filename}')

        # 添加错误摘要
        if errors and len(errors) > 0:
            subtitle_para = doc.add_paragraph(f'发现 {len(errors)} 个格式问题:')
            subtitle_run = subtitle_para.runs[0]
            subtitle_run.bold = True
            subtitle_run.font.size = Pt(14)
            # 设置字体
            subtitle_run.font.name = 'Times New Roman'  # 先设置英文字体
            # 设置中文字体
            subtitle_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            for i, error in enumerate(errors):
                p = doc.add_paragraph()
                error_run = p.add_run(f"{i+1}. {error.get('message', '')}")
                error_run.bold = True
                # 设置字体
                error_run.font.name = 'Times New Roman'  # 先设置英文字体
                # 设置中文字体
                error_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                if 'location' in error and error['location']:
                    location_run = p.add_run(f" (位置: {error['location']})")
                    location_run.italic = True
                    # 设置字体
                    location_run.font.name = 'Times New Roman'  # 先设置英文字体
                    # 设置中文字体
                    location_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            no_error_para = doc.add_paragraph('没有发现格式问题。')
            no_error_run = no_error_para.runs[0]
            no_error_run.bold = True
            # 设置字体
            no_error_run.font.name = 'Times New Roman'  # 先设置英文字体
            # 设置中文字体
            no_error_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 保存报告到临时文件
        temp_dir = tempfile.gettempdir()
        report_filename = f"report_{os.path.basename(file_path)}"
        report_path = os.path.join(temp_dir, report_filename)
        doc.save(report_path)

        # 返回成功响应和报告路径信息，让前端使用GET请求下载
        return jsonify({
            "success": True,
            "message": "报告生成成功",
            "report_path": report_path,
            "original_filename": original_filename
        })
    except Exception as e:
        print(f"生成报告失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"生成报告失败: {str(e)}"
        }), 500

# 下载报告
@app.route('/api/get-report', methods=['GET'])
def get_report():
    try:
        doc_path = request.args.get('doc_path')
        original_filename = request.args.get('original_filename', '')

        if not doc_path:
            return jsonify({"success": False, "message": "缺少文档路径参数"}), 400

        # 检查文件是否存在
        if not os.path.exists(doc_path):
            return jsonify({"success": False, "message": "报告文件不存在，请先生成报告"}), 404

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(doc_path)
            download_name = f"report_{original_filename}"
        else:
            download_name = f"report_{original_filename}"

        # 直接以二进制模式读取文件并返回
        with open(doc_path, 'rb') as f:
            file_data = f.read()

        # 对文件名进行URL编码，避免非ASCII字符导致的编码问题
        from urllib.parse import quote
        encoded_filename = quote(download_name)

        response = Response(
            file_data,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{encoded_filename}"; filename*=UTF-8\'\'{encoded_filename}'
            }
        )
        return response
    except Exception as e:
        print(f"下载报告失败: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"下载报告失败: {str(e)}"
        }), 500

# 下载标记的文档
@app.route('/api/download-marked-document', methods=['POST'])
def download_marked_document():
    try:
        data = request.get_json()
        if not data or 'doc_path' not in data:
            return jsonify({"success": False, "message": "缺少文档路径参数"}), 400

        doc_path = data.get('doc_path')
        errors = data.get('errors', [])
        original_filename = data.get('original_filename', '')
        frontend_para_manager = data.get('para_manager', None)

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(doc_path)

        # 构造标记文档的路径
        temp_dir = tempfile.gettempdir()
        doc_name = os.path.basename(doc_path)
        marked_doc_path = os.path.join(temp_dir, f"marked_{doc_name}")

        # 如果前端传来了para_manager，则使用前端的para_manager
        if frontend_para_manager:
            print(f"使用前端传来的para_manager")
            try:
                # 创建一个新的ParagraphManager实例
                para_manager = ParagraphManager()

                # 将前端传来的数据加载到para_manager中
                if isinstance(frontend_para_manager, list):
                    for para_data in frontend_para_manager:
                        para_manager.add_paragraph_from_dict(para_data)

                # 更新或添加到analysised_para_manager中
                for i, item in enumerate(analysised_para_manager):
                    if item['doc_path'] == doc_path:
                        analysised_para_manager[i]['para_manager'] = para_manager
                        break
                else:
                    # 如果没有找到对应的条目，添加新的
                    analysised_para_manager.append({"doc_path": doc_path, "para_manager": para_manager})
            except Exception as e:
                print(f"加载前端传来的para_manager失败: {str(e)}")
                # 如果加载失败，则使用后端的para_manager
                para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == doc_path), None)
                if not para_manager:
                    # 如果没有找到对应的para_manager，创建一个新的
                    para_manager = ParagraphManager()
                    print(f"未找到对应的para_manager，创建了新的实例")
        else:
            # 如果前端没有传来para_manager，则使用后端的para_manager
            para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == doc_path), None)
            if not para_manager:
                # 如果没有找到对应的para_manager，创建一个新的
                para_manager = ParagraphManager()
                print(f"未找到对应的para_manager，创建了新的实例")

        # 打印错误列表，便于调试
        print(f"错误列表: {errors}")
        print(f"错误数量: {len(errors)}")

        # 使用前端传来的错误列表标记文档
        marked_doc_path = mark_document_errors(doc_path, errors, para_manager, marked_doc_path)

        # 返回成功响应和标记文档路径信息，让前端使用GET请求下载
        return jsonify({
            "success": True,
            "message": "文档标记成功",
            "marked_doc_path": marked_doc_path,
            "original_filename": original_filename
        })
    except Exception as e:
        print(f"标记文档失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"标记文档失败: {str(e)}"
        }), 500

# 下载已标记的文档
@app.route('/api/get-marked-document', methods=['GET'])
def get_marked_document():
    try:
        doc_path = request.args.get('doc_path')
        original_filename = request.args.get('original_filename', '')

        if not doc_path:
            return jsonify({"success": False, "message": "缺少文档路径参数"}), 400

        # 检查文件是否存在
        if not os.path.exists(doc_path):
            return jsonify({"success": False, "message": "标记文档不存在，请先生成报告"}), 404

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(doc_path)
            download_name = f"marked_{original_filename}"
        else:
            download_name = f"marked_{original_filename}"

        # 直接以二进制模式读取文件并返回
        with open(doc_path, 'rb') as f:
            file_data = f.read()

        # 对文件名进行URL编码，避免非ASCII字符导致的编码问题
        from urllib.parse import quote
        encoded_filename = quote(download_name)

        response = Response(
            file_data,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{encoded_filename}"; filename*=UTF-8\'\'{encoded_filename}'
            }
        )
        return response
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"下载标记文档失败: {str(e)}"
        }), 500

# 应用格式
@app.route('/api/apply-format', methods=['POST'])
def apply_format():
    try:
        # 检查内容类型，如果是JSON，则解析请求体
        if request.is_json:
            data = request.get_json()
        else:
            # 如果不是JSON，可能是表单数据或其他格式
            return Response(
                json.dumps({"success": False, "message": "请求格式不正确，需要JSON数据"}),
                status=400,
                mimetype='application/json'
            )

        if not data or 'doc_path' not in data or 'config_path' not in data:
            return Response(
                json.dumps({"success": False, "message": "缺少必要参数"}),
                status=400,
                mimetype='application/json'
            )

        doc_path = data.get('doc_path')
        config_path = data.get('config_path')
        errors = data.get('errors', [])
        original_filename = data.get('original_filename', '')
        frontend_para_manager = data.get('para_manager', None)

        # 验证文件存在
        if not os.path.exists(doc_path):
            return Response(
                json.dumps({"success": False, "message": "文档文件不存在"}),
                status=404,
                mimetype='application/json'
            )

        if not os.path.exists(config_path):
            return Response(
                json.dumps({"success": False, "message": "配置文件不存在"}),
                status=404,
                mimetype='application/json'
            )

        # 如果前端传来了para_manager，则使用前端的para_manager
        para_manager = None
        if frontend_para_manager:
            try:
                # 创建一个新的ParagraphManager实例
                para_manager = ParagraphManager()

                # 将前端传来的数据加载到para_manager中
                if isinstance(frontend_para_manager, list):
                    for para_data in frontend_para_manager:
                        para_manager.add_paragraph_from_dict(para_data)

                # 更新或添加到analysised_para_manager中
                for i, item in enumerate(analysised_para_manager):
                    if item['doc_path'] == doc_path:
                        analysised_para_manager[i]['para_manager'] = para_manager
                        break
                else:
                    # 如果没有找到对应的条目，添加新的
                    analysised_para_manager.append({"doc_path": doc_path, "para_manager": para_manager})

                print(f"应用格式时使用前端传来的para_manager")
            except Exception as e:
                print(f"加载前端传来的para_manager失败: {str(e)}")
                para_manager = None

        # 如果没有从前端获取到有效的para_manager
        if not para_manager:
            # 尝试从后端已存储的para_manager中获取
            para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == doc_path), None)
            if not para_manager:
                # 如果没有找到对应的para_manager，创建一个新的
                print(f"应用格式时未找到对应的para_manager，创建了新的实例")
                # 调用check_format获取para_manager
                try:
                    errors, para_manager = check_format(doc_path, config_path, agents["format"])
                    # 存储当前的para_manager
                    analysised_para_manager.append({"doc_path": doc_path, "para_manager": para_manager})
                except Exception as check_error:
                    print(f"检查格式创建para_manager失败: {str(check_error)}")
                    return Response(
                        json.dumps({"success": False, "message": f"检查格式创建para_manager失败: {str(check_error)}"}),
                        status=500,
                        mimetype='application/json'
                    )

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(doc_path)

        # 使用原始文件名生成输出文件名，但去除特殊字符防止路径问题
        safe_filename = ''.join(c for c in os.path.splitext(original_filename)[0] if c.isalnum() or c in '._- ')
        output_filename = f"{safe_filename}_formatted.docx"

        # 确保输出到临时目录，防止权限问题
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, output_filename)

        # 如果没有有效的para_manager，重新检查获取
        if not para_manager or not para_manager.paragraphs:
            print(f"重新检查格式获取错误和para_manager")
            try:
                errors, para_manager = check_format(doc_path, config_path, agents["format"])
                # 更新存储的para_manager
                for i, item in enumerate(analysised_para_manager):
                    if item['doc_path'] == doc_path:
                        analysised_para_manager[i]['para_manager'] = para_manager
                        break
                else:
                    # 如果没有找到对应的条目，添加新的
                    analysised_para_manager.append({"doc_path": doc_path, "para_manager": para_manager})
            except Exception as check_error:
                print(f"重新检查格式失败: {str(check_error)}")
                return Response(
                    json.dumps({"success": False, "message": f"重新检查格式失败: {str(check_error)}"}),
                    status=500,
                    mimetype='application/json'
                )

        # 确保para_manager有效
        if not para_manager or not para_manager.paragraphs:
            return Response(
                json.dumps({"success": False, "message": "无法获取文档段落信息，请重新检查格式"}),
                status=500,
                mimetype='application/json'
            )

        # 应用格式，并将错误信息传递给 generate_formatted_doc 函数
        try:
            # 传递原文档路径，使其在原文档基础上修改内容，保持所有元素的相对位置不变
            output_path = generate_formatted_doc(config_path, para_manager, output_path, errors, doc_path=doc_path)
        except Exception as e:
            print(f"生成格式化文档失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response(
                json.dumps({"success": False, "message": f"生成格式化文档失败: {str(e)}"}),
                status=500,
                mimetype='application/json'
            )

        # 检查文件是否存在且可读
        if not os.path.exists(output_path):
            return Response(
                json.dumps({"success": False, "message": "生成格式化文档后找不到文件"}),
                status=404,
                mimetype='application/json'
            )

        try:
            # 尝试获取文件大小，确保可读
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                return Response(
                    json.dumps({"success": False, "message": "生成的文档为空文件"}),
                    status=500,
                    mimetype='application/json'
                )

            # 检查文件大小是否合理 (不超过100MB)
            if file_size > 100 * 1024 * 1024:
                return Response(
                    json.dumps({"success": False, "message": "生成的文档过大，请重试"}),
                    status=500,
                    mimetype='application/json'
                )

            print(f"格式应用成功，文件大小: {file_size} 字节，准备返回文件: {output_path}")

            # 生成下载文件名 (仅字母数字和部分符号，避免编码问题)
            download_name = f"{safe_filename}_formatted.docx"

            # 使用更安全的方式返回文件
            from werkzeug.utils import secure_filename

            # 直接以二进制模式读取文件并返回
            with open(output_path, 'rb') as f:
                file_data = f.read()

            # 对文件名进行URL编码，避免非ASCII字符导致的编码问题
            from urllib.parse import quote
            encoded_filename = quote(download_name)

            # 设置响应头，支持断点续传和正确的MIME类型
            headers = {
                'Content-Disposition': f'attachment; filename="{encoded_filename}"; filename*=UTF-8\'\'{encoded_filename}',
                'Content-Length': str(file_size),
                'Accept-Ranges': 'bytes'
            }

            # 创建响应对象
            response = Response(
                file_data,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers=headers
            )

            return response

        except Exception as file_error:
            print(f"读取或准备文件数据失败: {str(file_error)}")
            import traceback
            traceback.print_exc()
            return Response(
                json.dumps({"success": False, "message": f"读取或准备文件数据失败: {str(file_error)}"}),
                status=500,
                mimetype='application/json'
            )

    except Exception as e:
        print(f"应用格式时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        # 返回JSON错误响应
        return Response(
            json.dumps({"success": False, "message": str(e)}),
            status=500,
            mimetype='application/json'
        )


# 下载格式化文档
@app.route('/api/get-formatted-document', methods=['GET'])
def get_formatted_document():
    try:
        doc_path = request.args.get('doc_path')
        original_filename = request.args.get('original_filename', '')

        if not doc_path:
            return jsonify({"success": False, "message": "缺少文档路径参数"}), 400

        # 检查文件是否存在
        if not os.path.exists(doc_path):
            return jsonify({"success": False, "message": "格式化文档不存在，请先应用格式"}), 404

        # 如果没有提供原始文件名，则使用文档路径中的文件名
        if not original_filename:
            original_filename = os.path.basename(doc_path)
            download_name = f"{os.path.splitext(original_filename)[0]}_formatted.docx"
        else:
            download_name = f"{os.path.splitext(original_filename)[0]}_formatted.docx"

        # 直接以二进制模式读取文件并返回
        with open(doc_path, 'rb') as f:
            file_data = f.read()

        # 对文件名进行URL编码，避免非ASCII字符导致的编码问题
        from urllib.parse import quote
        encoded_filename = quote(download_name)

        response = Response(
            file_data,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{encoded_filename}"; filename*=UTF-8\'\'{encoded_filename}'
            }
        )
        return response
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"下载格式化文档失败: {str(e)}"
        }), 500


# 和大模型交互
@app.route('/api/send-message', methods=['POST'])
def send_message():
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({"success": False, "message": "缺少必要参数"}), 400
        message = data.get('message')
        doc_path = data.get('doc_path')
        config_path = data.get('config_path')
        doc_content = docx_parser.extract_doc_content(doc_path)

        # 获取段落管理器
        para_manager = None
        frontend_para_manager = data.get('para_manager', None)

        # 如果前端传来了para_manager，则使用前端的para_manager
        if frontend_para_manager:
            try:
                # 创建一个新的ParagraphManager实例
                para_manager = ParagraphManager()

                # 将前端传来的数据加载到para_manager中
                if isinstance(frontend_para_manager, list):
                    for para_data in frontend_para_manager:
                        para_manager.add_paragraph_from_dict(para_data)

                # 更新或添加到analysised_para_manager中
                for i, item in enumerate(analysised_para_manager):
                    if item['doc_path'] == doc_path:
                        analysised_para_manager[i]['para_manager'] = para_manager
                        break
                else:
                    # 如果没有找到对应的条目，添加新的
                    analysised_para_manager.append({"doc_path": doc_path, "para_manager": para_manager})
            except Exception as e:
                print(f"加载前端传来的para_manager失败: {str(e)}")
                para_manager = None

        # 如果没有从前端获取到有效的para_manager
        if not para_manager:
            # 尝试从后端已存储的para_manager中获取
            para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == doc_path), None)

        # 使用CommunicateAgent处理消息，包含意图分析和分发，同时传入文档全文和段落管理器
        response = agents["communicate"].get_response(message, doc_content, para_manager, config_path)

        return jsonify({"success": True, "message": response})
    except Exception as e:
        print(f"发送消息时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500

# 段落分析API
@app.route('/api/analyze-paragraph', methods=['POST'])
def analyze_paragraph():
    try:
        data = request.get_json()
        if not data or 'doc_path' not in data or 'para_index' not in data:
            return jsonify({"success": False, "message": "缺少必要参数"}), 400

        doc_path = data.get('doc_path')
        para_index = data.get('para_index')
        context_range = data.get('context_range', 2)  # 默认前后各取2个段落作为上下文

        # 获取段落管理器
        para_manager = None
        frontend_para_manager = data.get('para_manager', None)

        # 如果前端传来了para_manager，则使用前端的para_manager
        if frontend_para_manager:
            try:
                # 创建一个新的ParagraphManager实例
                para_manager = ParagraphManager()

                # 将前端传来的数据加载到para_manager中
                if isinstance(frontend_para_manager, list):
                    for para_data in frontend_para_manager:
                        para_manager.add_paragraph_from_dict(para_data)
            except Exception as e:
                print(f"加载前端传来的para_manager失败: {str(e)}")
                para_manager = None

        # 如果没有从前端获取到有效的para_manager
        if not para_manager:
            # 尝试从后端已存储的para_manager中获取
            para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == doc_path), None)
            if not para_manager:
                return jsonify({"success": False, "message": "找不到段落管理器，请先检查格式"}), 404

        # 检查段落索引是否有效
        if para_index < 0 or para_index >= len(para_manager.paragraphs):
            return jsonify({"success": False, "message": f"段落索引超出范围，有效范围为0-{len(para_manager.paragraphs)-1}"}), 400

        # 使用AdviceAgent分析段落
        advice_agent = agents["advice"]
        result = advice_agent.analyze_paragraph_manager(para_manager, para_index, context_range)

        return jsonify({"success": True, "result": result})
    except Exception as e:
        print(f"分析段落时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500

# 增强段落内容API
@app.route('/api/enhance-paragraphs', methods=['POST'])
def enhance_paragraphs():
    try:
        data = request.get_json()
        if not data or 'doc_path' not in data:
            return jsonify({"success": False, "message": "缺少必要参数"}), 400

        doc_path = data.get('doc_path')
        para_indices = data.get('para_indices')  # 可选参数，如果不提供则增强所有段落

        # 获取段落管理器
        para_manager = None
        frontend_para_manager = data.get('para_manager', None)

        # 如果前端传来了para_manager，则使用前端的para_manager
        if frontend_para_manager:
            try:
                # 创建一个新的ParagraphManager实例
                para_manager = ParagraphManager()

                # 将前端传来的数据加载到para_manager中
                if isinstance(frontend_para_manager, list):
                    for para_data in frontend_para_manager:
                        para_manager.add_paragraph_from_dict(para_data)
            except Exception as e:
                print(f"加载前端传来的para_manager失败: {str(e)}")
                para_manager = None

        # 如果没有从前端获取到有效的para_manager
        if not para_manager:
            # 尝试从后端已存储的para_manager中获取
            para_manager = next((item['para_manager'] for item in analysised_para_manager if item['doc_path'] == doc_path), None)
            if not para_manager:
                return jsonify({"success": False, "message": "找不到段落管理器，请先检查格式"}), 404

        # 使用EditorAgent增强段落
        editor_agent = agents["editor"]
        result = editor_agent.enhance_paragraph_manager(para_manager, para_indices)

        return jsonify({"success": True, "result": result})
    except Exception as e:
        print(f"增强段落时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500

# 获取docx文档内容API
@app.route('/api/get-docx-content', methods=['GET'])
def get_docx_content():
    try:
        file_path = request.args.get('file_path')
        if not file_path or not os.path.exists(file_path):
            return jsonify({"success": False, "message": "文件不存在"}), 404

        # 直接返回文件内容
        return send_file(file_path, as_attachment=False)
    except Exception as e:
        app.logger.error(f"获取文档内容时出错: {str(e)}")
        return jsonify({"success": False, "message": f"获取文档内容时出错: {str(e)}"}), 500

if __name__ == "__main__":
    socketio.run(app, host="0.0.0.0", port=8000, debug=True)
