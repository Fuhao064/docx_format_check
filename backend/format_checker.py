import docx
from docx import Document
import extract_para_info, docx_parser
import json
from typing import Dict, List, Tuple, Union, Optional, Any
from agents.format_agent import FormatAgent
from para_type import ParagraphManager, ParsedParaType, translation_dict
import os, concurrent, re
from utils import are_values_equal, extract_number_from_string, parse_llm_json_response, are_alignments_equal, are_fonts_equal
from config_utils import load_config, save_config, update_config

def is_value_equal(expected, actual, key):
    """比较两个值是否相等，考虑到不同表示形式的数值"""
    # 对特定的key做特殊处理
    if key.lower() == 'size':
        # 处理字号的特殊情况
        expected_num = extract_number(expected)
        actual_num = extract_number(actual)
        if expected_num is not None and actual_num is not None:
            return abs(expected_num - actual_num) < 0.5  # 允许0.5pt的误差

    # 检查字符串类型的相等性
    if isinstance(expected, str) and isinstance(actual, str):
        # 清除单位并转为小写进行比较
        expected_clean = expected.lower().strip()
        actual_clean = actual.lower().strip()

        # 如果完全相等，直接返回True
        if expected_clean == actual_clean:
            return True

        # 尝试提取数值部分进行比较
        expected_num = extract_number(expected)
        actual_num = extract_number(actual)
        if expected_num is not None and actual_num is not None:
            return abs(expected_num - actual_num) < 0.1

    # 处理数值类型或混合类型
    elif isinstance(expected, (int, float)) or isinstance(actual, (int, float)):
        expected_num = extract_number(expected)
        actual_num = extract_number(actual)
        if expected_num is not None and actual_num is not None:
            return abs(expected_num - actual_num) < 0.1

    # 默认使用原有方法
    return are_values_equal(expected, actual, key)

def extract_number(value):
    """从字符串中提取数字值，支持多种单位和格式"""
    # 如果已经是数字类型，直接返回
    if isinstance(value, (int, float)):
        return float(value)

    # 转换为字符串
    value_str = str(value).lower().strip()

    # 常见单位与对应的换算系数（相对于厘米）
    unit_factors = {
        'cm': 1.0,    # 厘米
        'mm': 0.1,    # 毫米
        'dm': 10.0,   # 分米
        'm': 100.0,   # 米
        'pt': 1.0,    # 点，对于字体大小不进行换算
        'in': 2.54,   # 英寸
        'inch': 2.54, # 英寸
        '寸': 3.33     # 中国传统单位
    }

    # 尝试提取数字和单位
    match = re.search(r'([-+]?\d*\.?\d+)\s*([a-z寸]+)?', value_str)
    if not match:
        return extract_number_from_string(value)

    num_value = float(match.group(1))
    unit = match.group(2) if match.group(2) else ''  # 无单位的情况

    # 如果没有单位，直接返回数值
    if not unit:
        return num_value

    # 应用单位转换系数
    factor = unit_factors.get(unit, 1.0)
    return num_value * factor

def check_paper_format(doc_info: Dict, required_format: Dict) -> List[Dict]:
    """检查文档基本格式"""
    errors = []

    # 获取纸张设置
    paper_format = required_format.get('paper_format', {})
    if not paper_format:
        return errors

    # 检查页面大小
    expected_page_size = paper_format.get('page_size', None)
    if expected_page_size:
        actual_width = doc_info.get('page_width', None)
        actual_height = doc_info.get('page_height', None)

        expected_width = extract_number(expected_page_size.get('width', '0'))
        expected_height = extract_number(expected_page_size.get('height', '0'))

        if expected_width and expected_height:
            if actual_width and actual_height:
                width_diff = abs(actual_width - expected_width)
                height_diff = abs(actual_height - expected_height)

                if width_diff > 0.5 or height_diff > 0.5:  # 允许0.5cm的误差
                    errors.append({
                        'message': f'页面大小不符合要求，要求{expected_width}×{expected_height}厘米，实际为{actual_width}×{actual_height}厘米',
                        'location': '文档全局设置'
                    })

    # 检查页边距
    expected_margins = paper_format.get('margins', None)
    if expected_margins:
        margin_keys = ['top', 'bottom', 'left', 'right']

        for key in margin_keys:
            expected_value = extract_number(expected_margins.get(key, '0'))
            actual_value = doc_info.get(f'margin_{key}', None)

            if expected_value and actual_value:
                margin_diff = abs(actual_value - expected_value)

                if margin_diff > 0.1:  # 允许0.1cm的误差
                    margin_name = {
                        'top': '上边距',
                        'bottom': '下边距',
                        'left': '左边距',
                        'right': '右边距'
                    }.get(key, key)

                    errors.append({
                        'message': f'{margin_name}不符合要求，要求{expected_value}厘米，实际为{actual_value}厘米',
                        'location': '文档全局设置'
                    })

    return errors

def check_reference_format(doc_path: str, required_format: Dict) -> List[Dict]:
    """检查引用和参考文献格式"""
    errors = []

    try:
        doc = docx.Document(doc_path)

        # 获取参考文献格式要求
        reference_format = required_format.get('reference_format', {})
        if not reference_format:
            return errors

        # 获取引用样式
        citation_style = reference_format.get('citation_style', '').lower()
        if not citation_style:
            return errors

        # 查找参考文献部分
        references_start = False
        references = []

        for para in doc.paragraphs:
            text = para.text.strip()

            # 标识参考文献部分的开始
            if text.startswith('参考文献') or text.lower().startswith('references'):
                references_start = True
                continue

            # 收集参考文献条目
            if references_start and text:
                references.append(text)

        # 检查是否有参考文献
        if not references:
            errors.append({
                'message': '未找到参考文献部分或参考文献为空',
                'location': '文档末尾部分'
            })
            return errors

        # 根据不同引用样式检查参考文献格式
        if 'gb' in citation_style or 'gbt' in citation_style:
            # 检查GB/T格式
            ref_errors = _check_gbt_references(references, citation_style)
            for i, error in enumerate(ref_errors):
                if isinstance(error, dict):
                    errors.append(error)
                else:
                    errors.append({
                        'message': error,
                        'location': f"参考文献[{i+1}]"
                    })

        elif 'apa' in citation_style:
            # 检查APA格式
            ref_errors = _check_apa_references(references)
            for i, error in enumerate(ref_errors):
                if isinstance(error, dict):
                    errors.append(error)
                else:
                    errors.append({
                        'message': error,
                        'location': f"参考文献[{i+1}]"
                    })

        elif 'mla' in citation_style:
            # 检查MLA格式
            ref_errors = _check_mla_references(references)
            for i, error in enumerate(ref_errors):
                if isinstance(error, dict):
                    errors.append(error)
                else:
                    errors.append({
                        'message': error,
                        'location': f"参考文献[{i+1}]"
                    })

        else:
            errors.append({
                'message': f"不支持的引用样式: {citation_style}",
                'location': '参考文献格式设置'
            })

    except Exception as e:
        errors.append({
            'message': f"检查参考文献格式时出错: {str(e)}",
            'location': '参考文献部分'
        })

    return errors

def _check_gbt_references(references: List[str], standard: str) -> List[str]:
    """
    检查参考文献是否符合GB/T 7714标准
    """
    errors = []

    # 检查每条参考文献
    for i, ref in enumerate(references):
        # 检查编号格式 [1] 或 [1]
        if not re.match(r'^\[\d+\]', ref):
            errors.append(f"参考文献 #{i+1} 编号格式错误，应以[数字]开头")
            continue

        # 检查作者与题名之间的分隔符
        if '. ' not in ref and '．' not in ref:
            errors.append(f"参考文献 #{i+1} 作者与题名之间缺少正确的分隔符")

        # 检查不同类型文献的格式
        if '[J]' in ref or '[J/OL]' in ref:  # 期刊论文
            if not _check_journal_format(ref, standard):
                errors.append(f"参考文献 #{i+1} 期刊论文格式不符合{standard}标准")
        elif '[M]' in ref:  # 专著
            if not _check_book_format(ref, standard):
                errors.append(f"参考文献 #{i+1} 专著格式不符合{standard}标准")
        elif '[D]' in ref:  # 学位论文
            if not _check_thesis_format(ref, standard):
                errors.append(f"参考文献 #{i+1} 学位论文格式不符合{standard}标准")
        elif '[C]' in ref or '[C/OL]' in ref:  # 会议论文
            if not _check_conference_format(ref, standard):
                errors.append(f"参考文献 #{i+1} 会议论文格式不符合{standard}标准")
        elif '[EB/OL]' in ref:  # 电子资源
            if not _check_electronic_format(ref, standard):
                errors.append(f"参考文献 #{i+1} 电子资源格式不符合{standard}标准")
        else :errors.append(f"不符合的参考文献类型: {ref}")

    return errors

def _check_journal_format(ref: str, standard: str) -> bool:
    """检查期刊论文格式"""
    # GB/T 7714-2015格式: [序号] 作者. 题名[J]. 刊名, 出版年, 卷号(期号): 起止页码.
    # 更灵活的正则表达式，允许一些常见变体
    pattern = r'^\[\d+\]\s+[\w\s,]+\.\s+.+\[J(/OL)?\]\.\s+.+,\s+\d{4}(,\s+\d+(\(\d+\))?)?(\:\s*\d+(-\d+)?)?\.?$'
    return bool(re.match(pattern, ref, re.UNICODE))

def _check_book_format(ref: str, standard: str) -> bool:
    """检查专著格式"""
    # GB/T 7714-2015格式: [序号] 作者. 书名[M]. 版本(第1版不标注). 出版地: 出版社, 出版年: 起止页码.
    # 更灵活的模式
    pattern = r'^\[\d+\]\s+[\w\s,]+\.\s+.+\[M\]\.\s+.*(\:\s+.+,\s+\d{4}(\:\s*\d+(-\d+)?)?)?\.?$'
    return bool(re.match(pattern, ref, re.UNICODE))

def _check_thesis_format(ref: str, standard: str) -> bool:
    """检查学位论文格式"""
    # GB/T 7714-2015格式: [序号] 作者. 题名[D]. 保存地: 保存单位, 出版年.
    # 更灵活的模式
    pattern = r'^\[\d+\]\s+[\w\s,]+\.\s+.+\[D\]\.\s+.*(\:\s+.+,\s+\d{4})?\.?$'
    return bool(re.match(pattern, ref, re.UNICODE))

def _check_conference_format(ref: str, standard: str) -> bool:
    """检查会议论文格式"""
    # GB/T 7714-2015格式: [序号] 作者. 题名[C]. 会议名, 会议地点, 会议年份. 出版地: 出版者, 出版年: 起止页码.
    # 更灵活的模式
    pattern = r'^\[\d+\]\s+[\w\s,]+\.\s+.+\[C(/OL)?\]\.\s+.+(,\s+.+,\s+\d{4})?\.(\s+.+\:\s+.+,\s+\d{4}(\:\s*\d+(-\d+)?)?)?\.?$'
    return bool(re.match(pattern, ref, re.UNICODE))

def _check_electronic_format(ref: str, standard: str) -> bool:
    """检查电子资源格式"""
    # GB/T 7714-2015格式: [序号] 作者. 题名[EB/OL]. 出版地: 出版者, 出版年[引用日期]. 获取和访问路径.
    # 更灵活的模式，网址部分可选
    pattern = r'^\[\d+\]\s+[\w\s,]+\.\s+.+\[EB/OL\]\.(\s+.+\:\s+.+,\s+\d{4})?\s*(\[\d{4}-\d{2}-\d{2}\])?\.(\s+https?://.*)?$'
    return bool(re.match(pattern, ref, re.UNICODE))

def _check_apa_references(references: List[str]) -> List[str]:
    """
    检查参考文献是否符合APA标准
    """
    errors = []
    # APA格式检查逻辑
    for i, ref in enumerate(references):
        # 检查作者格式
        if not re.match(r'^[\w\s,]+\s+\(\d{4}\)', ref):
            errors.append(f"参考文献 #{i+1} 作者和年份格式不符合APA标准")

    return errors

def _check_mla_references(references: List[str]) -> List[str]:
    """
    检查参考文献是否符合MLA标准
    """
    errors = []
    # MLA格式检查逻辑
    for i, ref in enumerate(references):
        # 检查作者格式
        if not re.match(r'^[\w\s,]+\.', ref):
            errors.append(f"参考文献 #{i+1} 作者格式不符合MLA标准")

    return errors

def check_table_format(doc_path: str, required_format: Dict) -> List[Dict]:
    """检查表格格式"""
    errors = []

    try:
        doc = docx.Document(doc_path)

        # 获取表格格式要求
        table_format = required_format.get('table_format', {})
        if not table_format:
            return errors

        # 遍历文档中的表格
        table_count = 0
        for table in doc.tables:
            table_count += 1

            # 检查表格内容格式
            content_errors = _check_table_content_format(table)
            for error in content_errors:
                if isinstance(error, dict):
                    error['location'] = f"表{table_count}: {error.get('location', '')}"
                    errors.append(error)
                else:
                    errors.append({
                        'message': error,
                        'location': f"表{table_count}"
                    })

        # 遍历文档中的段落，查找表格标题
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # 检查是否为表格标题
            if text.startswith('表') or text.lower().startswith('table'):
                # 检查标题格式
                caption_errors = _check_caption_format(para, table_format, 'table')

                # 匹配表格编号
                match = re.search(r'表\s*(\d+[-.]?\d*)', text)
                table_num = match.group(1) if match else f"{i+1}"

                for error in caption_errors:
                    if isinstance(error, dict):
                        error['location'] = f"表{table_num}标题: {error.get('location', '')}"
                        errors.append(error)
                    else:
                        errors.append({
                            'message': error,
                            'location': f"表{table_num}标题"
                        })

                # 检查表格编号格式
                number_errors = _check_table_number_format(text)
                for error in number_errors:
                    if isinstance(error, dict):
                        error['location'] = f"表{table_num}编号: {error.get('location', '')}"
                        errors.append(error)
                    else:
                        errors.append({
                            'message': error,
                            'location': f"表{table_num}编号"
                        })

    except Exception as e:
        errors.append({
            'message': f"检查表格格式时出错: {str(e)}",
            'location': '全文表格'
        })

    return errors

def _check_caption_format(caption_para, required_format: Dict, caption_type: str) -> List[Dict]:
    """
    检查标题格式
    """
    errors = []

    # 检查字体
    if hasattr(caption_para, 'runs') and caption_para.runs:
        run = caption_para.runs[0]

        # 检查字体名称
        font_name = run.font.name
        required_zh_font = required_format.get('fonts', {}).get('zh_family')
        required_en_font = required_format.get('fonts', {}).get('en_family')

        if required_zh_font and not any(font in font_name for font in [required_zh_font, '黑体']):
            errors.append({
                'message': f"{caption_type}标题中文字体不符合要求，应为{required_zh_font}",
                'location': '标题字体'
            })

        # 检查字体大小
        font_size = run.font.size
        required_size = required_format.get('fonts', {}).get('size')
        if required_size and font_size:
            required_pt = extract_number(required_size)
            actual_pt = font_size.pt
            if abs(actual_pt - required_pt) > 0.5:
                errors.append({
                    'message': f"{caption_type}标题字体大小不符合要求，应为{required_size}",
                    'location': '标题字体大小'
                })

    # 检查对齐方式
    alignment = caption_para.alignment
    required_alignment = required_format.get('paragraph_format', {}).get('alignment')
    if required_alignment and required_alignment.lower() == 'center' and alignment != 1:  # 1表示居中
        errors.append({
            'message': f"{caption_type}标题未居中显示",
            'location': '标题对齐'
        })

    # 检查是否同时包含中英文标题
    if not (re.search(r'[\u4e00-\u9fa5]', caption_para.text) and
            re.search(r'[a-zA-Z]', caption_para.text)):
        errors.append({
            'message': f"{caption_type}标题应同时包含中英文",
            'location': '标题语言'
        })

    return errors

def _check_table_content_format(table) -> List[Dict]:
    """检查表格内容格式"""
    errors = []

    try:
        row_count = len(table.rows)
        col_count = len(table.columns)

        # 检查表格是否为空
        if row_count == 0 or col_count == 0:
            errors.append({
                'message': '表格为空',
                'location': '表格内容'
            })
            return errors

        # 检查表格单元格是否为空
        empty_cells = []
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if not cell.text.strip():
                    empty_cells.append(f"行{i+1}列{j+1}")

        if empty_cells:
            if len(empty_cells) <= 3:
                cell_str = '、'.join(empty_cells)
            else:
                cell_str = f"{empty_cells[0]}、{empty_cells[1]}等{len(empty_cells)}处"

            errors.append({
                'message': f'表格存在空单元格',
                'location': f"表格内容: {cell_str}"
            })

        # 检查表格行列一致性
        inconsistent_rows = []
        first_row_cell_count = len(table.rows[0].cells)

        for i, row in enumerate(table.rows):
            if len(row.cells) != first_row_cell_count:
                inconsistent_rows.append(i+1)

        if inconsistent_rows:
            if len(inconsistent_rows) <= 3:
                rows_str = '、'.join(map(str, inconsistent_rows))
            else:
                rows_str = f"{inconsistent_rows[0]}、{inconsistent_rows[1]}等{len(inconsistent_rows)}行"

            errors.append({
                'message': '表格行列不一致',
                'location': f"表格内容: 第{rows_str}"
            })

    except Exception as e:
        errors.append({
            'message': f"检查表格内容时出错: {str(e)}",
            'location': '表格内容'
        })

    return errors

def _check_table_number_format(caption_text: str) -> List[Dict]:
    """
    检查表格编号是否按章节顺序
    """
    errors = []

    # 匹配表格编号格式：表x-y 或 Table x-y
    match = re.search(r'表\s*(\d+)[-－](\d+)|Table\s*(\d+)[-－](\d+)', caption_text, re.IGNORECASE)
    if not match:
        errors.append({
            'message': f"表格编号格式不正确，应为'表x-y'或'Table x-y'格式",
            'location': '表格编号'
        })
        return errors

    # 提取章节号和表格序号
    if match.group(1) and match.group(2):  # 中文格式
        chapter_num = int(match.group(1))
        table_num = int(match.group(2))
    else:  # 英文格式
        chapter_num = int(match.group(3))
        table_num = int(match.group(4))

    # 检查章节号和表格序号是否合理
    if chapter_num <= 0 or table_num <= 0:
        errors.append({
            'message': f"表格编号中章节号和表格序号应为正整数",
            'location': '表格编号'
        })

    return errors

def check_figure_format(doc_path: str, required_format: Dict) -> List[str]:
    """
    检查图片格式是否符合要求
    """
    errors = []
    doc = docx.Document(doc_path)

    # 获取图片部分的要求格式
    figure_required_format = required_format.get('figures', {})

    # 检查图片标题和内容
    figure_count = 0

    # 遍历段落查找图片标题
    for i, para in enumerate(doc.paragraphs):
        # 查找图片标题段落
        if re.search(r'图\s*\d+[-－]\d+|Figure\s*\d+[-－]\d+', para.text, re.IGNORECASE):
            figure_count += 1

            # 检查图片标题格式
            caption_errors = _check_caption_format(para, figure_required_format.get('caption', {}), "图片")
            errors.extend(caption_errors)

            # 检查图片编号是否按章节顺序
            number_errors = _check_figure_number_format(para.text)
            errors.extend(number_errors)

            # 检查图片标题位置（应在图片下方）
            position = figure_required_format.get('caption', {}).get('position')
            if position == 'below':
                # 简单检查：如果上一段是空的，可能表示图片位置
                if i > 0 and not doc.paragraphs[i-1].text.strip():
                    pass  # 符合要求
                else:
                    errors.append(f"图片标题'{para.text}'应位于图片下方")

    # 检查是否有图片
    if figure_count == 0:
        errors.append("文档中未找到图片")

    # 打印错误信息
    if errors:
        print("\n图片格式检查错误:")
        for err in errors:
            print(f"- {err}")

    return errors

def _check_figure_number_format(caption_text: str) -> List[str]:
    """
    检查图片编号是否按章节顺序
    """
    errors = []

    # 匹配图片编号格式：图x-y 或 Figure x-y
    match = re.search(r'图\s*(\d+)[-－](\d+)|Figure\s*(\d+)[-－]\d+', caption_text, re.IGNORECASE)
    if not match:
        errors.append(f"图片编号格式不正确，应为'图x-y'或'Figure x-y'格式")
        return errors

    # 提取章节号和图片序号
    if match.group(1) and match.group(2):  # 中文格式
        chapter_num = int(match.group(1))
        figure_num = int(match.group(2))
    else:  # 英文格式
        chapter_num = int(match.group(3))
        figure_num = int(match.group(4))

    # 检查章节号和图片序号是否合理
    if chapter_num <= 0 or figure_num <= 0:
        errors.append(f"图片编号中章节号和图片序号应为正整数")

    return errors

def check_abstract(paragraph_manager:ParagraphManager):
    """检查摘要格式"""
    errors = []

    # 查找摘要段落
    abstract_paras = []
    abstract_content_paras = []

    for para in paragraph_manager.paragraphs:
        if para.type == ParsedParaType.ABSTRACT_ZH or para.content.strip().startswith('摘要'):
            abstract_paras.append(para)
        elif para.type == ParsedParaType.ABSTRACT_CONTENT_ZH:
            abstract_content_paras.append(para)

    # 检查是否存在摘要
    if not abstract_paras:
        errors.append({
            'message': '文档中缺少摘要',
            'location': '文档开头部分'
        })
        return errors

    # 检查是否有摘要内容
    if not abstract_content_paras:
        errors.append({
            'message': '文档中缺少摘要内容',
            'location': '摘要标题后'
        })
        return errors

    # 检查摘要内容长度
    for content_para in abstract_content_paras:
        abstract_text = content_para.content.strip()
        if len(abstract_text) < 10:  # 假设摘要至少应该有10个字符
            errors.append({
                'message': '摘要内容过短',
                'location': '摘要部分'
            })

    return errors

def check_keywords(paragraph_manager:ParagraphManager):
    """检查关键词格式"""
    errors = []

    # 查找关键词段落
    keyword_paras = []
    for para in paragraph_manager.paragraphs:
        para_text = para.content.strip()
        if para_text.startswith('关键词') or para_text.startswith('Keywords'):
            keyword_paras.append(para)

    # 检查是否存在关键词
    if not keyword_paras:
        errors.append({
            'message': '文档中缺少关键词',
            'location': '摘要之后'
        })
        return errors

    # 检查关键词内容
    keyword_para = keyword_paras[0]
    keyword_text = keyword_para.content.strip()

    # 检查是否有足够的关键词（假设至少3个）
    if '：' in keyword_text:
        keywords = keyword_text.split('：')[1].split('；')
    elif ':' in keyword_text:
        keywords = keyword_text.split(':')[1].split(';')
    else:
        keywords = []

    if len(keywords) < 3:
        errors.append({
            'message': '关键词数量不足，建议至少提供3个关键词',
            'location': '关键词部分'
        })

    return errors

def compare_json_formats(doc_json: Dict, required_format: Dict) -> List[Dict]:
    """
    比较文档JSON与要求格式JSON，找出不符合要求的部分

    参数:
    doc_json: 文档JSON
    required_format: 格式要求JSON

    返回:
    List[Dict]: 错误列表
    """
    errors = []

    # 检查纸张设置
    if "paper" in doc_json and "paper" in required_format:
        paper_errors = check_paper_format(doc_json["paper"], required_format)
        if paper_errors:
            errors.extend(paper_errors)

    # 检查段落格式
    if "paragraphs" in doc_json:
        paragraphs = doc_json["paragraphs"]

        # 构建段落类型映射
        para_type_map = {}
        for para in paragraphs:
            para_type = para.get("type")
            if para_type not in para_type_map:
                para_type_map[para_type] = []
            para_type_map[para_type].append(para)

        # 检查每种类型的段落格式
        for para_type, format_req in required_format.items():
            if para_type == "paper":
                continue  # 已经检查过纸张设置

            if para_type in para_type_map:
                # 检查该类型的所有段落
                for para in para_type_map[para_type]:
                    para_errors = check_paragraph_format(para, format_req, para_type)
                    if para_errors:
                        errors.extend(para_errors)

    return errors

def check_paragraph_format(para: Dict, format_req: Dict, para_type: str) -> List[Dict]:
    """
    检查单个段落的格式是否符合要求

    参数:
    para: 段落信息
    format_req: 格式要求
    para_type: 段落类型

    返回:
    List[Dict]: 错误列表
    """
    errors = []
    para_id = para.get("id", "unknown")
    para_content = para.get("content", "")
    para_meta = para.get("meta", {})

    # 检查字体设置
    if "fonts" in format_req and "fonts" in para_meta:
        font_errors = check_font_format(para_meta["fonts"], format_req["fonts"], para_type, para_id, para_content)
        if font_errors:
            errors.extend(font_errors)

    # 检查段落格式设置
    if "paragraph_format" in format_req and "paragraph_format" in para_meta:
        para_format_errors = check_para_format_settings(para_meta["paragraph_format"], format_req["paragraph_format"], para_type, para_id, para_content)
        if para_format_errors:
            errors.extend(para_format_errors)

    return errors

def extract_number(value):
    """从字符串中提取数字"""
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        match = re.search(r'([-+]?\d*\.?\d+)', value)
        if match:
            return float(match.group(1))
    return 0.0

def check_font_format(font_info: Dict, required_font: Dict, para_type: str, para_id: str, para_content: str) -> List[Dict]:
    """
    检查字体格式是否符合要求

    参数:
    font_info: 段落字体信息
    required_font: 要求的字体格式
    para_type: 段落类型
    para_id: 段落ID
    para_content: 段落内容

    返回:
    List[Dict]: 错误列表
    """
    errors = []
    type_name = translation_dict.get(para_type, para_type)
    content_preview = para_content[:20] + ("..." if len(para_content) > 20 else "")

    # 确保 font_info 中的所有属性都是集合
    for key in ['zh_family', 'en_family', 'size', 'color', 'bold', 'italic']:
        if key not in font_info or not isinstance(font_info[key], set):
            font_info[key] = set()

    # 检查中文字体
    if "zh_family" in required_font:
        required_zh = required_font["zh_family"]
        actual_zh = list(font_info.get("zh_family", set()))

        # 如果实际值为空，设置默认值
        if not actual_zh:
            # 根据段落类型设置默认字体
            if para_type in ['title_zh', 'title_en', 'heading1', 'heading2', 'heading3']:
                actual_zh = ["黑体"]  # 标题默认黑体
            else:
                actual_zh = ["宋体"]  # 正文默认宋体

        # 使用are_fonts_equal函数检查字体是否等价
        is_font_match = any(are_fonts_equal(required_zh, zh) for zh in actual_zh)

        if not is_font_match:
            errors.append({
                "id": para_id,
                "message": f"{type_name}中文字体不符合要求，应为{required_zh}，实际为{', '.join(actual_zh)}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查英文字体
    if "en_family" in required_font:
        required_en = required_font["en_family"]
        actual_en = list(font_info.get("en_family", set()))

        # 如果实际值为空，设置默认值
        if not actual_en:
            actual_en = ["Times New Roman"]  # 默认Times New Roman

        # 使用are_fonts_equal函数检查字体是否等价
        is_font_match = any(are_fonts_equal(required_en, en) for en in actual_en)

        if not is_font_match:
            errors.append({
                "id": para_id,
                "message": f"{type_name}英文字体不符合要求，应为{required_en}，实际为{', '.join(actual_en)}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查字号
    if "size" in required_font:
        required_size = extract_number(required_font["size"])
        actual_sizes = list(font_info.get("size", set()))

        # 如果实际值为空，设置默认值
        if not actual_sizes:
            # 根据段落类型设置默认字号
            if para_type in ['title_zh', 'title_en', 'heading1', 'heading2', 'heading3']:
                actual_sizes = [16.0]  # 标题默认字号16磅
            else:
                actual_sizes = [10.5]  # 正文默认字号10.5磅

        # 检查是否符合要求
        if not any(abs(extract_number(size) - required_size) < 0.5 for size in actual_sizes):
            errors.append({
                "id": para_id,
                "message": f"{type_name}字号不符合要求，应为{required_size}磅，实际为{', '.join(str(s) for s in actual_sizes)}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查加粗
    if "bold" in required_font:
        required_bold = required_font["bold"]
        actual_bold = list(font_info.get("bold", set()))

        # 如果实际值为空，则默认为不加粗
        if not actual_bold:
            actual_bold = [False]
        # 如果实际值中同时包含True和False，则使用False
        elif len(actual_bold) > 1 and False in actual_bold:
            actual_bold = [False]

        # 检查是否符合要求
        if not any(bold == required_bold for bold in actual_bold):
            errors.append({
                "id": para_id,
                "message": f"{type_name}加粗设置不符合要求，应为{'加粗' if required_bold else '不加粗'}，实际为{'加粗' if any(actual_bold) else '不加粗'}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查斜体
    if "italic" in required_font:
        required_italic = required_font["italic"]
        actual_italic = list(font_info.get("italic", set()))

        # 如果实际值为空，则默认为不斜体
        if not actual_italic:
            actual_italic = [False]
        # 如果实际值中同时包含True和False，则使用False
        elif len(actual_italic) > 1 and False in actual_italic:
            actual_italic = [False]

        # 检查是否符合要求
        if not any(italic == required_italic for italic in actual_italic):
            errors.append({
                "id": para_id,
                "message": f"{type_name}斜体设置不符合要求，应为{'斜体' if required_italic else '不斜体'}，实际为{'斜体' if any(actual_italic) else '不斜体'}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查颜色
    if "color" in required_font:
        required_color = required_font["color"].lower()
        actual_colors = [str(c).lower() for c in font_info.get("color", set())]

        # 如果实际值为空，则默认为黑色
        if not actual_colors:
            actual_colors = ["black"]

        # 处理颜色等价性
        is_color_match = False

        # 特殊处理黑色
        if required_color == "#000000" or required_color == "000000":
            is_color_match = any(c == "black" or c == "#000000" or c == "000000" for c in actual_colors)
        elif required_color == "black":
            is_color_match = any(c == "black" or c == "#000000" or c == "000000" for c in actual_colors)
        else:
            # 其他颜色的匹配
            is_color_match = any(required_color in c or c in required_color for c in actual_colors)

        # 检查是否符合要求
        if not is_color_match:
            errors.append({
                "id": para_id,
                "message": f"{type_name}颜色设置不符合要求，应为{required_color}，实际为{', '.join(actual_colors)}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    return errors

def check_para_format_settings(para_format: Dict, required_format: Dict, para_type: str, para_id: str, para_content: str) -> List[Dict]:
    """
    检查段落格式设置是否符合要求

    参数:
    para_format: 段落格式信息
    required_format: 要求的段落格式
    para_type: 段落类型
    para_id: 段落ID
    para_content: 段落内容

    返回:
    List[Dict]: 错误列表
    """
    errors = []
    type_name = translation_dict.get(para_type, para_type)
    content_preview = para_content[:20] + ("..." if len(para_content) > 20 else "")

    # 检查对齐方式
    if "alignment" in required_format:
        required_alignment = required_format["alignment"]
        actual_alignment = para_format.get("alignment")

        # 如果没有获取到对齐方式，默认为左对齐
        if not actual_alignment:
            actual_alignment = "左对齐"

        # 使用are_alignments_equal函数检查对齐方式是否等价
        is_alignment_match = are_alignments_equal(required_alignment, actual_alignment)

        if not is_alignment_match:
            errors.append({
                "id": para_id,
                "message": f"{type_name}对齐方式不符合要求，应为{required_alignment}，实际为{actual_alignment}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查行间距
    if "line_spacing" in required_format:
        required_spacing = required_format["line_spacing"]
        actual_spacing = para_format.get("line_spacing")
        if actual_spacing and not are_values_equal(required_spacing, actual_spacing, "line_spacing"):
            errors.append({
                "id": para_id,
                "message": f"{type_name}行间距不符合要求，应为{required_spacing}，实际为{actual_spacing}",
                "location": f"{type_name}: \"{content_preview}\""
            })

    # 检查缩进设置
    if "indentation" in required_format:
        indentation = required_format["indentation"]

        # 检查首行缩进
        if "first_line" in indentation:
            required_first_line = extract_number(indentation["first_line"])
            actual_first_line = para_format.get("first_line_indent")
            if actual_first_line is not None and abs(extract_number(actual_first_line) - required_first_line) > 0.1:
                errors.append({
                    "id": para_id,
                    "message": f"{type_name}首行缩进不符合要求，应为{required_first_line}厘米，实际为{actual_first_line}",
                    "location": f"{type_name}: \"{content_preview}\""
                })

        # 检查左缩进
        if "left" in indentation:
            required_left = extract_number(indentation["left"])
            actual_left = para_format.get("left_indent")
            if actual_left is not None and abs(extract_number(actual_left) - required_left) > 0.1:
                errors.append({
                    "id": para_id,
                    "message": f"{type_name}左缩进不符合要求，应为{required_left}厘米，实际为{actual_left}",
                    "location": f"{type_name}: \"{content_preview}\""
                })

        # 检查右缩进
        if "right" in indentation:
            required_right = extract_number(indentation["right"])
            actual_right = para_format.get("right_indent")
            if actual_right is not None and abs(extract_number(actual_right) - required_right) > 0.1:
                errors.append({
                    "id": para_id,
                    "message": f"{type_name}右缩进不符合要求，应为{required_right}厘米，实际为{actual_right}",
                    "location": f"{type_name}: \"{content_preview}\""
                })

        # 检查段前距
        if "space_before" in indentation:
            required_before = extract_number(indentation["space_before"])
            actual_before = para_format.get("before_spacing")
            if actual_before is not None and abs(extract_number(actual_before) - required_before) > 0.1:
                errors.append({
                    "id": para_id,
                    "message": f"{type_name}段前距不符合要求，应为{required_before}磅，实际为{actual_before}",
                    "location": f"{type_name}: \"{content_preview}\""
                })

        # 检查段后距
        if "space_after" in indentation:
            required_after = extract_number(indentation["space_after"])
            actual_after = para_format.get("after_spacing")
            if actual_after is not None and abs(extract_number(actual_after) - required_after) > 0.1:
                errors.append({
                    "id": para_id,
                    "message": f"{type_name}段后距不符合要求，应为{required_after}磅，实际为{actual_after}",
                    "location": f"{type_name}: \"{content_preview}\""
                })

    return errors

# 检查一定要求出现的段落是否出现
def check_required_paragraphs(paragraph_manager:ParagraphManager, required_format:Dict) -> List[Dict]:
    """
    检查一定要求出现的段落是否出现

    参数:
    paragraph_manager: 段落管理器
    required_format: 格式要求

    返回:
    List[Dict]: 错误列表
    """
    errors = []

    # 定义必须出现的段落类型
    required_types = [
        ("title_zh", "中文标题"),
        ("abstract_zh", "中文摘要"),
        ("abstract_content_zh", "中文摘要内容"),
        ("keywords_zh", "中文关键词"),
        ("keywords_content_zh", "中文关键词内容")
    ]

    # 检查每种必需的段落类型
    for type_key, type_name in required_types:
        if type_key in required_format:
            # 检查该类型是否存在于段落管理器中
            try:
                para_type = ParsedParaType(type_key)
                paras = paragraph_manager.get_by_type(para_type)
                if not paras:
                    errors.append({
                        "message": f"缺少{type_name}部分",
                        "location": "文档结构"
                    })
            except ValueError:
                # 如果段落类型无效，则跳过
                continue

    return errors

def check_main_format(paragraph_manager:ParagraphManager, required_format:Dict, check_abstract_keywords:bool=True) -> List[Dict]:
    """
    检查主要文档格式

    参数:
    paragraph_manager: 段落管理器
    required_format: 格式要求
    check_abstract_keywords: 是否检查摘要和关键词，默认为True
    """
    errors = []

    # 检查摘要格式
    if check_abstract_keywords:
        abstract_errors = check_abstract(paragraph_manager)
        if abstract_errors:
            errors.extend(abstract_errors)

        # 检查关键词格式
        keyword_errors = check_keywords(paragraph_manager)
        if keyword_errors:
            errors.extend(keyword_errors)

    # 获取段落格式要求
    paragraph_format = required_format.get('paragraph_format', {})
    if not paragraph_format:
        return errors

    # 获取各类型段落的格式要求
    title_format = paragraph_format.get('title', {})
    heading1_format = paragraph_format.get('heading1', {})
    heading2_format = paragraph_format.get('heading2', {})
    heading3_format = paragraph_format.get('heading3', {})
    body_format = paragraph_format.get('body', {})

    # 遍历所有段落，检查格式 - 修改为正确的列表迭代方式
    for para in paragraph_manager.paragraphs:
        para_type = para.type.value if hasattr(para, 'type') else ''
        para_text = para.content if hasattr(para, 'content') else ''
        para_attributes = para.meta if hasattr(para, 'meta') else {}

        # 根据段落类型选择相应的格式要求
        expected_format = {}
        location_prefix = ""

        if para_type == 'title_zh' or para_type == 'title_en':
            expected_format = title_format
            location_prefix = "标题"
        elif para_type == 'heading1':
            expected_format = heading1_format
            location_prefix = "一级标题"
        elif para_type == 'heading2':
            expected_format = heading2_format
            location_prefix = "二级标题"
        elif para_type == 'heading3':
            expected_format = heading3_format
            location_prefix = "三级标题"
        elif para_type == 'body':
            expected_format = body_format
            location_prefix = "正文"

        if not expected_format:
            continue

        # 递归检查格式
        try:
            format_errors = _recursive_check(para_attributes, expected_format)
            for error_info in format_errors:
                # 确保 error_info 是一个包含两个元素的元组
                if isinstance(error_info, tuple) and len(error_info) == 2:
                    # 解析错误信息
                    path_str, error_msg = error_info
                    location = f"{location_prefix}: \"{para_text[:20]}{'...' if len(para_text) > 20 else ''}\""
                    errors.append({
                        'message': error_msg,
                        'location': location
                    })
        except Exception as e:
            print(f"递归检查格式时出错: {str(e)}")
            errors.append({
                'message': f"检查格式时出错: {str(e)}",
                'location': f"{location_prefix}: \"{para_text[:20]}{'...' if len(para_text) > 20 else ''}\""
            })

    return errors

def is_value_equal(expected, actual, field_name=None):
    """检查两个值是否相等

    参数:
    expected: 期望的值
    actual: 实际的值
    field_name: 字段名称（可选）

    返回:
    bool: 是否相等
    """
    # 如果是字符串，忽略大小写
    if isinstance(expected, str) and isinstance(actual, str):
        return expected.lower() == actual.lower()
    # 如果是数字，允许小误差
    elif (isinstance(expected, (int, float)) and isinstance(actual, (int, float))) or \
         (isinstance(expected, str) and isinstance(actual, (int, float))) or \
         (isinstance(expected, (int, float)) and isinstance(actual, str)):
        try:
            expected_num = float(expected) if isinstance(expected, str) else expected
            actual_num = float(actual) if isinstance(actual, str) else actual
            return abs(expected_num - actual_num) < 0.5
        except ValueError:
            return False
    # 如果是布尔值，进行特殊处理
    elif isinstance(expected, bool) or isinstance(actual, bool):
        # 尝试将字符串转换为布尔值
        if isinstance(expected, str):
            expected = expected.lower() in ['true', 'yes', '1', 't', 'y']
        if isinstance(actual, str):
            actual = actual.lower() in ['true', 'yes', '1', 't', 'y']
        return expected == actual
    # 其他情况直接比较
    else:
        return expected == actual

    # 注意: field_name 参数在这个函数中没有被使用，但它在函数签名中被保留，以便于将来可能的扩展

def _recursive_check(actual, expected):
    """
    递归检查嵌套字典的字段。
    返回格式为[(key, error_message), ...]的列表
    """
    errors = []

    for key, expected_value in expected.items():
        actual_value = actual.get(key)

        # 如果字段不存在
        if actual_value is None:
            errors.append((key, f"缺少必需字段: '{key}'"))
            continue

        # 如果字段是列表类型，优先检查长度
        if isinstance(actual_value, list):
            if len(actual_value) > 1:
                errors.append((key, f"字段 '{key}' 存在多个值: {actual_value}"))
                continue
            if len(actual_value) == 0:
                continue
            actual_value = actual_value[0]  # 取第一个值进行比较

        # 递归处理嵌套字典
        if isinstance(expected_value, dict):
            if not isinstance(actual_value, dict):
                errors.append((key, f"字段类型不匹配: '{key}' 应为字典类型，实际为 {type(actual_value).__name__}"))
                continue

            # 递归检查嵌套字段
            deeper_errors = _recursive_check(actual_value, expected_value)
            for err_key, err_msg in deeper_errors:
                errors.append((f"{key}.{err_key}", err_msg))

        # 检查值是否匹配
        else:
            if not is_value_equal(expected_value, actual_value, field_name=key):
                errors.append((key, f"'{key}' 不匹配: 要求 {expected_value}, 实际 {actual_value}"))

    return errors

def remark_para_type(doc_path: str, format_agent: FormatAgent) -> ParagraphManager:
    # 初始化段落管理器
    paragraph_manager = ParagraphManager()

    # 读取doc的段落信息
    paragraph_manager = extract_para_info.extract_para_format_info(doc_path, paragraph_manager)

    # 文档整个内容
    doc_content = docx_parser.extract_doc_content(doc_path)

    # 定义任务函数
    def process_paragraph(para_index: int, para: Dict, doc_content_sent: bool, manager: ParagraphManager):
        try:
            if para["type"] != ParsedParaType.BODY:
                pass
            # 提取当前段落信息
            para_string = para["content"]

            # 调用大模型预测类型
            response = format_agent.predict_location(doc_content, para_string, doc_content_sent)

            # 直接使用response字典，无需再次解析
            response_dict = response

            # 将返回的location转换为ParsedParaType枚举类型
            try:
                new_type = ParsedParaType(response_dict["location"])
                # 直接访问paragraphs列表中的对应索引
                manager.paragraphs[para_index].type = new_type
            except ValueError:
                print(f"Invalid paragraph type received: {response_dict['location']}")
                manager.paragraphs[para_index].type = ParsedParaType.BODY

        except Exception as e:
            # 如果解析失败，将段落类型设置为 BODY
            print(f"Error processing paragraph: {para_string[:50]}... Error: {e}")
            manager.paragraphs[para_index].type = ParsedParaType.BODY

    # 将段落json转为中文
    paras_info_json_zh = paragraph_manager.to_chinese_dict()

    # 使用线程池并发处理，限制最大线程数为10
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        futures = []
        doc_content_sent = False

        # 创建任务列表
        tasks = []
        for i, para in enumerate(paras_info_json_zh):
            tasks.append((i, para, doc_content_sent, paragraph_manager))
            doc_content_sent = True  # 更新全文标志

        # 分批提交任务，每批最多5个任务
        batch_size = 5
        for i in range(0, len(tasks), batch_size):
            batch_tasks = tasks[i:i+batch_size]
            batch_futures = [executor.submit(process_paragraph, *task) for task in batch_tasks]
            futures.extend(batch_futures)

            # 等待当前批次完成
            for future in concurrent.futures.as_completed(batch_futures):
                try:
                    future.result()  # 确保任务完成
                except Exception as e:
                    print(f"Error in thread pool task: {e}")

    return paragraph_manager

def extract_doc_to_json(doc_path: str, format_agent: FormatAgent) -> Dict[str, Any]:
    """
    将文档内容抽取为JSON格式

    参数:
    doc_path: 文档路径
    format_agent: 格式代理

    返回:
    Dict: 包含文档结构和格式的JSON对象
    """
    try:
        # 初始化段落管理器
        paragraph_manager = ParagraphManager()

        # 提取段落格式信息
        paragraph_manager = extract_para_info.extract_para_format_info(doc_path, paragraph_manager)

        # 重新标记段落类型
        paragraph_manager = remark_para_type(doc_path, format_agent)

        # 获取文档基本信息
        doc_info = docx_parser.extract_section_info(doc_path)

        # 构建完整的文档JSON
        doc_json = {
            "paper": doc_info,
            "paragraphs": paragraph_manager.to_dict()
        }

        return doc_json
    except Exception as e:
        print(f"提取文档JSON时出错: {str(e)}")
        return {"error": str(e)}

def check_format(doc_path: str, config_path: str, format_agent: FormatAgent) -> Tuple[List[dict], ParagraphManager]:
    """
    检查文档格式是否符合要求

    参数:
    doc_path: 文档路径
    config_path: 格式配置文件路径
    format_agent: 格式代理

    返回:
    Tuple[List[dict], ParagraphManager]: 错误列表和段落管理器
    """
    try:
        # 加载配置
        required_format = load_config(config_path)

        # 提取文档为JSON格式
        doc_json = extract_doc_to_json(doc_path, format_agent)

        # 获取段落管理器
        paragraph_manager = ParagraphManager()
        paragraph_manager = extract_para_info.extract_para_format_info(doc_path, paragraph_manager)
        paragraph_manager = remark_para_type(doc_path, format_agent)

        # 收集所有错误
        errors = []

        # 方法1: 使用新的JSON比对方法
        json_errors = compare_json_formats(doc_json, required_format)
        if json_errors:
            errors.extend(json_errors)

        # 方法2: 使用传统的检查方法作为补充
        # 检查基本格式
        paper_errors = check_paper_format(doc_json["paper"], required_format)
        if paper_errors:
            # 去重处理
            for error in paper_errors:
                if error not in errors:
                    errors.append(error)

        # 检查必要出现的段落是否出现
        required_errors = check_required_paragraphs(paragraph_manager, required_format)
        if required_errors:
            # 去重处理
            for error in required_errors:
                if error not in errors:
                    errors.append(error)

        # 检查段落格式
        para_errors = check_main_format(paragraph_manager, required_format)
        if para_errors:
            # 去重处理
            for error in para_errors:
                if error not in errors:
                    errors.append(error)

        # 返回错误列表和段落管理器
        return errors, paragraph_manager

    except Exception as e:
        print(f"检查格式时出错: {str(e)}")
        return [{"message": f"检查格式时出错: {str(e)}", "location": "系统错误"}], ParagraphManager()


