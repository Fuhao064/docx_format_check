import docx
import xml.etree.ElementTree as ET
import json, re
from para_type import ParsedParaType, ParagraphManager
from docx.shared import RGBColor
from docx.oxml.ns import qn

def extract_para_format_from_style(style)-> dict:
    # 获取段落样式
    if style.type == 1:
        para_format = style.paragraph_format
        style_name = style.name
        # 创建一个字典来存储当前样式的信息
        current_style_info = {
            'style_name': style_name
        }
        # 获取对齐方式
        alignment = para_format.alignment if para_format.alignment and para_format.alignment is not None else 0
        format_alignment = analysise_alignment(alignment)
        current_style_info['alignment'] = format_alignment
        # 获取首行缩进
        first_line_indent = para_format.first_line_indent if para_format.first_line_indent and para_format.first_line_indent is not None else 0
        current_style_info['first_line_indent'] = first_line_indent
        # 获取左缩进（单位：磅）
        left_indent = para_format.left_indent if para_format.left_indent and para_format.left_indent is not None else 0
        current_style_info['left_indent'] = left_indent
        # 获取右缩进（单位：磅）
        right_indent = para_format.right_indent if para_format.right_indent and para_format.right_indent is not None else 0
        current_style_info['right_indent'] = right_indent
        # 获取段前间距（单位：磅）
        before_spacing = para_format.space_before if para_format.space_before and para_format.space_before is not None else 0
        current_style_info['before_spacing'] = before_spacing
        # 获取段后间距（单位：磅）
        after_spacing = para_format.space_after if para_format.space_after and para_format.space_after is not None else 0
        current_style_info['after_spacing'] = after_spacing
        # 获取行间距
        line_spacing = para_format.line_spacing
        if isinstance(line_spacing, float):
            current_style_info['line_spacing'] = f"{line_spacing} 倍行距"
        elif isinstance(line_spacing, int):
            current_style_info['line_spacing'] = f"{line_spacing / 20} 磅"
        else:
            current_style_info['line_spacing'] = "未设置明确行间距"
        return current_style_info

def extract_para_format_info_from_paragraph_fromat(para)-> dict:
    # 解析para数据
    current_style_info = {}
    # 初始化段落格式信息
    alignment = None
    first_line_indent = None
    left_indent = None
    right_indent = None
    before_spacing = None
    after_spacing = None
    line_spacing = None
    # 获取段落xml数据
    para_xml = para._p.xml
    # print(para_xml)
    # 命名空间
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
    }
    root = ET.fromstring(para_xml)
    pPr = root.find('.//w:pPr', namespaces)
    # 输出节点信息
    # print(ET.tostring(pPr, encoding='utf-8').decode('utf-8'))
    if pPr is not None:
        # 查找缩进节点
        ind = pPr.find('w:ind', namespaces)
        if ind is not None:
            # 使用完整的命名空间URI获取属性
            # print(ET.tostring(ind, encoding='utf-8').decode('utf-8'))
            w_ns = namespaces['w']
            first_line_indent = ind.get(f'{{{w_ns}}}firstLineChars')
            if first_line_indent is not None:
                first_line_indent = 0 if first_line_indent == 0 else int(first_line_indent) / 100
            # 如果没有首行字符缩进，则获取首行缩进值
            if first_line_indent is None:
                first_line_indent = ind.get(f'{{{w_ns}}}firstLine')
    else:
        first_line_indent = None
    # 获取段落对齐方式
    alignment = analysise_alignment(para.paragraph_format.alignment)
    # 获取左缩进
    if para.paragraph_format.left_indent is not None:
        left_indent = para.paragraph_format.left_indent
    # 获取右缩进
    if para.paragraph_format.right_indent is not None:
        right_indent = para.paragraph_format.right_indent
    # 获取段前间距
    if para.paragraph_format.space_before is not None:
        before_spacing = para.paragraph_format.space_before
    # 获取段后间距
    if para.paragraph_format.space_after is not None:
        after_spacing = para.paragraph_format.space_after
    # 获取行间距
    if para.paragraph_format.line_spacing is not None:
        line_spacing = para.paragraph_format.line_spacing

    current_style_info['alignment'] = alignment
    current_style_info['first_line_indent'] = first_line_indent
    current_style_info['left_indent'] = left_indent
    current_style_info['right_indent'] = right_indent
    current_style_info['before_spacing'] = before_spacing
    current_style_info['after_spacing'] = after_spacing
    current_style_info['line_spacing'] = line_spacing
    return current_style_info

def extract_para_format_info(doc_path, manager: ParagraphManager):

    doc = docx.Document(doc_path)
    
    # 获取所有样式
    styles_info = []
    for style in doc.styles:
        current_style_info = extract_para_format_from_style(style)
        styles_info.append(current_style_info)
    
    # 将摘要等信息拆出来
    processed_paras = pre_process_paragraphs(doc)
    # processed_paras = doc.paragraphs
    # 初始化上一个段落信息
    last_para_type = ParsedParaType.BODY
    for para in processed_paras:
        # 如果段落文本为空，则跳过
        if not para.text.strip():
            continue
        
        # 创建段落的元数据字典
        meta_data = {}
        
        # 获取段落的xml数据
        style_xml_data = para.style.element.xml
        xml_data = para._p.xml
        
        # 解析段落样式
        para_format_info = extract_para_format_info_from_paragraph_fromat(para)
        style_info = find_style_by_name(para.style.name, styles_info)
        
        # 合并段落格式信息
        current_para_format_info = para_format_info
        for key in current_para_format_info.keys():
            if current_para_format_info[key] is None:
                current_para_format_info[key] = style_info[key]
        
        # 将段落格式信息添加到元数据
        meta_data["paragraph_format"] = current_para_format_info
        
        # 解析run数据中的字体信息
        runs = para.runs
        fonts_from_runs = extract_font_info_from_runs(runs)
        
        # 解析XML中的字体信息
        fonts_from_xml = extract_font_info_from_xml(xml_data)
        if font_dict_is_empty(fonts_from_xml):
            fonts_from_xml = extract_font_info_from_xml(style_xml_data)
        
        # 合并字体信息
        fonts = merge_font_info(fonts_from_runs, fonts_from_xml)
        
        # 将字体信息添加到元数据
        meta_data["fonts"] = fonts
        
        # 动态确定段落类型
        para_type = determine_para_type(para.text, last_para_type)
        last_para_type = para_type
        # 将段落信息添加到段落管理器
        manager.add_para(
            para_type=para_type,
            content=para.text,
            meta=meta_data
        )
    
    return manager

def determine_para_type(text, last_para_type):
    """
    根据段落内容动态确定段落类型
    """
    # 如果上一个段落是关键词，则判断为关键词内容段落
    if last_para_type == ParsedParaType.KEYWORDS_ZH:
        return ParsedParaType.KEYWORDS_CONTENT_ZH
    elif last_para_type == ParsedParaType.KEYWORDS_EN:
        return ParsedParaType.KEYWORDS_CONTENT_EN
    elif last_para_type == ParsedParaType.ABSTRACT_EN:
        return ParsedParaType.ABSTRACT_CONTENT_EN
    elif last_para_type == ParsedParaType.ABSTRACT_ZH:
        return ParsedParaType.ABSTRACT_CONTENT_ZH
    # 匹配关键词
    pattern = re.compile(r'\b(摘要|Abstract|关键词|Keywords)\b\s*[:：]', re.IGNORECASE)
    match = pattern.search(text)
    if match:
        keyword = match.group(1).lower()
        if keyword == "摘要":
            return ParsedParaType.ABSTRACT_ZH
        elif keyword == "abstract":
            return ParsedParaType.ABSTRACT_EN
        elif keyword == "关键词":
            return ParsedParaType.KEYWORDS_ZH
        elif keyword == "keywords":
            return ParsedParaType.KEYWORDS_EN
    return ParsedParaType.BODY

def extract_font_info_from_runs(runs: list) -> dict:
    fonts = {
        'zh_family': set(),
        'en_family': set(),
        'size': set(),
        'color': set(),
        'bold': set(),
        'italic': set(),
        'isAllcaps': None
    }
    
    for run in runs:
        if not run.text.strip():
            continue
            
        if run.font:
            # 提取字体名称
            font_name = run.font.name
            if font_name:
                if re.search(r'[A-Za-z]+', font_name):
                    fonts['en_family'].add(font_name)
                elif re.search(r'[\u4e00-\u9fa5]+', font_name):
                    fonts['zh_family'].add(font_name)
            
            # 尝试从run._element.rPr中获取中文字体
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None and hasattr(rPr, 'rFonts'):
                    rFonts = rPr.rFonts
                    if rFonts is not None:
                        # 获取中文字体
                        east_asia_font = rFonts.get(qn('w:eastAsia'))
                        if east_asia_font:
                            fonts['zh_family'].add(east_asia_font)
                        # 获取英文字体
                        ascii_font = rFonts.get(qn('w:ascii'))
                        if ascii_font:
                            fonts['en_family'].add(ascii_font)
            
            # 提取字体大小
            if run.font.size:
                fonts['size'].add(run.font.size.pt)
            
            # 提取字体颜色
            if run.font.color and run.font.color.rgb:
                if run.font.color.rgb == RGBColor(0, 0, 0):
                    fonts['color'].add("black")
                else:
                    # 将RGB颜色对象转换为字符串
                    rgb_str = str(run.font.color.rgb)
                    fonts['color'].add(rgb_str)
            
            # 提取加粗信息
            if run.font.bold is not None:
                fonts['bold'].add(run.font.bold)
            
            # 提取斜体信息
            if run.font.italic is not None:
                fonts['italic'].add(run.font.italic)
            
            # 提取全大写信息 - 这个属性对整个段落是一致的
            if fonts['isAllcaps'] is None:
                fonts['isAllcaps'] = is_all_caps(run.text)
    
    # 如果没有找到中文字体，但有中文内容，则根据内容判断使用默认中文字体
    if not fonts['zh_family'] and any(re.search(r'[\u4e00-\u9fa5]', run.text) for run in runs if run.text.strip()):
        # 检查段落样式中是否有中文字体设置
        for run in runs:
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # 尝试从段落样式中获取中文字体
                    if hasattr(run.font, '_element') and hasattr(run.font._element, 'rPr'):
                        style_rPr = run.font._element.rPr
                        if style_rPr is not None and hasattr(style_rPr, 'rFonts'):
                            style_rFonts = style_rPr.rFonts
                            if style_rFonts is not None:
                                east_asia_font = style_rFonts.get(qn('w:eastAsia'))
                                if east_asia_font:
                                    fonts['zh_family'].add(east_asia_font)
                                    break
        
        # 如果仍然没有找到中文字体，则使用默认字体
        if not fonts['zh_family']:
            # 根据内容特征判断可能使用的字体
            for run in runs:
                if re.search(r'[\u4e00-\u9fa5]', run.text):
                    # 如果是标题样式，可能使用黑体
                    if run.bold or (run.font.bold is not None and run.font.bold):
                        fonts['zh_family'].add("黑体")
                    else:
                        # 默认使用宋体
                        fonts['zh_family'].add("宋体")
                    break
    
    return fonts

def extract_font_info_from_xml(xml_data):
    # 定义 XML 命名空间
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }
    root = ET.fromstring(xml_data)
    
    # 初始化中文字体和英文字体堆栈
    fonts = {
        'zh_family': set(),
        'en_family': set(),
        'size': set(),
        'color': set(),
        'bold': set(),
        'italic': set(),
        'isAllcaps': None
    }

    # 查找所有 w:rFonts 元素
    rFonts_elements = root.findall('.//w:rFonts', namespaces)
    # 遍历 w:rFonts 元素
    for rFonts in rFonts_elements:
        # 获取中文字体信息
        chinese_font = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
        if chinese_font:
            fonts["zh_family"].add(chinese_font)
        # 获取英文字体信息
        english_font = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
        if english_font:
            fonts["en_family"].add(english_font)

    # 查找字体大小信息 (w:sz)
    sz_elements = root.findall('.//w:sz', namespaces)
    for sz in sz_elements:
        val = sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        if val:
            # 字号单位是半点，需要除以2转换为磅
            fonts["size"].add(float(val) / 2)

    # 查找字体颜色信息 (w:color)
    color_elements = root.findall('.//w:color', namespaces)
    for color in color_elements:
        val = color.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        if val:
            if val.lower() == "000000" or val.lower() == "auto":
                fonts["color"].add("black")
            else:
                fonts["color"].add(f"#{val}")

    # 查找字体加粗信息 (w:b)
    bold_elements = root.findall('.//w:b', namespaces)
    for bold in bold_elements:
        val = bold.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        if val is None or val != "0":  # 如果没有val属性或val不等于0，则为加粗
            fonts["bold"].add(True)
        else:
            fonts["bold"].add(False)

    # 查找字体斜体信息 (w:i)
    italic_elements = root.findall('.//w:i', namespaces)
    for italic in italic_elements:
        val = italic.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        if val is None or val != "0":  # 如果没有val属性或val不等于0，则为斜体
            fonts["italic"].add(True)
        else:
            fonts["italic"].add(False)

    # 查找字体大写信息 (w:caps)
    caps_elements = root.findall('.//w:caps', namespaces)
    for caps in caps_elements:
        val = caps.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        if val is None or val != "0":  # 如果没有val属性或val不等于0，则为全大写
            fonts["isAllcaps"] = True
        else:
            fonts["isAllcaps"] = False
    
    # 如果没有找到中文字体，但XML中可能包含中文字符，则添加默认中文字体
    if not fonts["zh_family"]:
        # 检查是否有样式名称提示
        style_elements = root.findall('.//w:pStyle', namespaces)
        for style in style_elements:
            val = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if val:
                if "标题" in val or "title" in val.lower() or "heading" in val.lower():
                    fonts["zh_family"].add("黑体")
                    break
        
        # 如果仍然没有找到中文字体，则使用宋体作为默认字体
        if not fonts["zh_family"]:
            fonts["zh_family"].add("宋体")

    return fonts

def find_style_by_name(style_name: str, styles_info: list):
    for style_info in styles_info:
        if style_info is not None:
            if style_info['style_name'] == style_name:
                return style_info
    print(f"Style {style_name} not found.")
    return None

def analysise_alignment(alignment: int) -> str:
    if alignment == 0:
        return "左对齐"
    elif alignment == 1:
        return "居中"
    elif alignment == 2:
        return "右对齐"
    elif alignment == 3:
        return "两端对齐"
    elif alignment == 4:
        return "分散对齐"
    else:
        return None

def font_dict_is_empty(fonts: dict) -> bool:
    for font in fonts.values():
        if font:
            return False
    return True

def merge_font_info(fonts_from_runs: dict, fonts_from_xml: dict) -> dict:
    merged_fonts = {}
    keys = ['zh_family', 'en_family', 'size', 'color', 'bold', 'italic', 'isAllcaps']
    
    for key in keys:
        if key in ['zh_family', 'en_family', 'size', 'color']:
            # 合并集合类型，优先使用runs中的值，如果为空则使用xml的值
            merged_set = set()
            if fonts_from_runs[key]:
                merged_set.update(fonts_from_runs[key])
            if not merged_set and fonts_from_xml[key]:
                merged_set.update(fonts_from_xml[key])
            merged_fonts[key] = merged_set
        elif key in ['bold', 'italic']:
            # 处理布尔类型集合，当runs中有值时优先使用
            merged_set = set()
            if fonts_from_runs[key]:
                merged_set.update(fonts_from_runs[key])
            elif fonts_from_xml[key]:
                merged_set.update(fonts_from_xml[key])
            
            # 如果集合为空，默认为False
            if not merged_set:
                merged_set.add(False)
            
            merged_fonts[key] = merged_set
        elif key == 'isAllcaps':
            # isAllcaps是单一值，不是集合
            if fonts_from_runs[key] is not None:
                merged_fonts[key] = fonts_from_runs[key]
            elif fonts_from_xml[key] is not None:
                merged_fonts[key] = fonts_from_xml[key]
            else:
                merged_fonts[key] = False
        else:
            # 其他字段直接使用runs的值（如果存在）
            merged_fonts[key] = fonts_from_runs.get(key, fonts_from_xml.get(key))
    
    return merged_fonts

def is_all_caps(string: str) -> bool:
    # 判断是否全是大写英文字母
    english_words = re.findall(r'[A-Za-z]+', string)
    # print(english_words)
    if(english_words == []):
        return False
    # 检查提取出的每个英文单词是否都是大写
    return all(word.isupper() for word in english_words)

def split_paragraph(paragraph, split_pos):
    # 获取段落所有 run 并计算总文本
    runs = paragraph.runs
    full_text = ''.join([run.text for run in runs])
    
    # 如果文本为空或分割位置无效，则返回 None
    if not full_text or split_pos <= 0 or split_pos >= len(full_text):
        return None

    # 找到分割点所在的 run 和偏移量
    current_pos = 0
    split_run_idx = 0
    split_run_offset = 0
    for idx, run in enumerate(runs):
        run_len = len(run.text)
        if current_pos + run_len > split_pos:
            split_run_idx = idx
            split_run_offset = split_pos - current_pos
            break
        current_pos += run_len
    else:
        return None  # 分割位置在段落末尾无需处理

    # 创建新段落并插入到原段落后
    parent = paragraph._parent
    new_para = parent.add_paragraph()
    new_para.style = paragraph.style
    
    # 将新段落移动到原段落后
    paragraph._p.addnext(new_para._p)
    
    # 复制段落格式
    para_format = paragraph.paragraph_format
    new_format = new_para.paragraph_format
    for attr in ['alignment', 'left_indent', 'right_indent', 'first_line_indent', 
                 'line_spacing', 'space_before', 'space_after', 'line_spacing_rule']:
        setattr(new_format, attr, getattr(para_format, attr))

    # 处理 run 的分割并复制字体样式
    def copy_run_style(source, target):
        target.font.bold = source.font.bold
        target.font.italic = source.font.italic
        target.font.underline = source.font.underline
        if source.font.color.rgb:
            target.font.color.rgb = source.font.color.rgb
        target.font.name = source.font.name
        if source.font.size:
            target.font.size = source.font.size
        
        # 修复FutureWarning
        if (hasattr(source.font, '_element') and 
            hasattr(source.font._element, 'rPr') and 
            source.font._element.rPr is not None and 
            hasattr(source.font._element.rPr, 'rFonts') and 
            source.font._element.rPr.rFonts is not None):
            
            target.font._element.rPr.rFonts.set(qn('w:eastAsia'), 
                source.font._element.rPr.rFonts.get(qn('w:eastAsia')))

    # 分割 run
    split_run = runs[split_run_idx]
    text_before = split_run.text[:split_run_offset]
    text_after = split_run.text[split_run_offset:]
    
    # 修改原 run 的文本
    split_run.text = text_before

    # 将分割后的内容和后续 run 转移到新段落
    if text_after:
        new_run = new_para.add_run(text_after)
        copy_run_style(split_run, new_run)

    for run in runs[split_run_idx + 1:]:
        new_run = new_para.add_run(run.text)
        copy_run_style(run, new_run)
        paragraph._p.remove(run._r)  # 从原段落移除 run

    return new_para

# 预处理段落，将摘要等信息提取出来
def pre_process_paragraphs(doc):
    """
    预处理文档段落，将包含摘要/关键词的段落拆分为独立段落
    """
    while True:
        processed = False
        # 使用副本避免修改时影响迭代
        for para in list(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            # 匹配中英文关键词（摘要、Abstract、关键词、Keywords）
            pattern = re.compile(r'\b(摘要|Abstract|关键词|Keywords)\b\s*[:：]', re.IGNORECASE)
            match = pattern.search(text)
            if not match:
                continue

            # 根据匹配的关键词设置段落类型
            # keyword = match.group(1).lower()
            # if keyword == "摘要":
            #     para_type = ParsedParaType.ABSTRACT_ZH
            # elif keyword == "abstract":
            #     para_type = ParsedParaType.ABSTRACT_EN
            # elif keyword == "关键词":
            #     para_type = ParsedParaType.KEYWORDS_ZH
            # elif keyword == "keywords":
            #     para_type = ParsedParaType.KEYWORDS_EN

            # print(f"正在处理段落: {text}")
            # 计算分割位置（关键词后的位置）
            split_pos = match.end()
            
            # 分割段落
            new_para = split_paragraph(para, split_pos)
            if new_para:
                # print(f"分割后新段落: {new_para.text}")
                processed = True
        if not processed:
            break
    return doc.paragraphs

# 创建段落管理器
manager = ParagraphManager()
manager = extract_para_format_info("test.docx", manager)

chinese_data = manager.to_chinese_dict()
# print("中文格式数据：")
# 保存为json文件
with open("para_format_info.json", "w", encoding="utf-8") as f:
    json.dump(chinese_data, f, ensure_ascii=False, indent=4)
