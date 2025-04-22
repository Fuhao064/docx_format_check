import os
import sys
import unittest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# 添加项目根目录到系统路径
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from editors.format_fixer import FormatFixer, batch_fix_errors, apply_format_requirements
from preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo

class TestFormatFixer(unittest.TestCase):
    """测试格式修复器"""
    
    def setUp(self):
        """测试前准备"""
        # 创建测试文档
        self.test_doc_path = os.path.join(os.path.dirname(__file__), 'test_doc.docx')
        self.create_test_document()
        
        # 创建段落管理器
        self.para_manager = ParagraphManager()
        self.para_manager.add_para(
            para_type=ParsedParaType.HEADING1,
            content="一、结合毕业论文（设",
            meta={"段落设置": {"对齐方式": "左对齐", "行间距": "1.15"}, "字体": {"字号": 10.5}}
        )
        self.para_manager.add_para(
            para_type=ParsedParaType.BODY,
            content="大型语言模型（LLM",
            meta={"段落设置": {"对齐方式": "左对齐"}, "字体": {"字号": 10.5}}
        )
        self.para_manager.add_para(
            para_type=ParsedParaType.BODY,
            content="这些模型不仅具备广泛",
            meta={"段落设置": {"对齐方式": "左对齐"}, "字体": {"字号": 10.5}}
        )
        self.para_manager.add_para(
            para_type=ParsedParaType.HEADING1,
            content="参考文献：",
            meta={"段落设置": {"对齐方式": "左对齐", "行间距": "1.15"}, "字体": {"字号": 10.5, "加粗": True}}
        )
        
        # 创建错误列表
        self.errors = [
            {
                'message': "'字号' 不匹配: 要求 16pt, 实际 10.5",
                'location': "字体.一、结合毕业论文（设"
            },
            {
                'message': "'行间距' 不匹配: 要求 固定值 20pt, 实际 1.15",
                'location': "段落设置.一、结合毕业论文（设"
            },
            {
                'message': "'对齐方式' 不匹配: 要求 居中, 实际 左对齐",
                'location': "段落设置.一、结合毕业论文（设"
            },
            {
                'message': "'加粗' 不匹配: 要求 否, 实际 是",
                'location': "字体.参考文献："
            }
        ]
        
        # 创建格式要求
        self.requirements = {
            "heading1": {
                "fonts": {
                    "zh_family": "黑体",
                    "en_family": "Times New Roman",
                    "size": "16pt",
                    "bold": True
                },
                "paragraph_format": {
                    "alignment": "center",
                    "line_spacing": "20pt",
                    "indentation": {
                        "first_line": "0cm"
                    }
                }
            },
            "body": {
                "fonts": {
                    "zh_family": "宋体",
                    "en_family": "Times New Roman",
                    "size": "12pt",
                    "bold": False
                },
                "paragraph_format": {
                    "alignment": "left",
                    "line_spacing": "20pt",
                    "indentation": {
                        "first_line": "2cm"
                    }
                }
            }
        }
    
    def create_test_document(self):
        """创建测试文档"""
        doc = Document()
        
        # 添加标题段落
        heading = doc.add_paragraph("一、结合毕业论文（设")
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.line_spacing = 1.15
        for run in heading.runs:
            run.font.size = Pt(10.5)
        
        # 添加正文段落1
        body1 = doc.add_paragraph("大型语言模型（LLM")
        body1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in body1.runs:
            run.font.size = Pt(10.5)
        
        # 添加正文段落2
        body2 = doc.add_paragraph("这些模型不仅具备广泛")
        body2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in body2.runs:
            run.font.size = Pt(10.5)
        
        # 添加参考文献段落
        ref = doc.add_paragraph("参考文献：")
        ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ref.paragraph_format.line_spacing = 1.15
        for run in ref.runs:
            run.font.size = Pt(10.5)
            run.font.bold = True
        
        # 保存文档
        doc.save(self.test_doc_path)
    
    def tearDown(self):
        """测试后清理"""
        # 删除测试文档
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)
        
        # 删除修复后的文档
        fixed_doc_path = os.path.join(os.path.dirname(__file__), 'fixed_test_doc.docx')
        if os.path.exists(fixed_doc_path):
            os.remove(fixed_doc_path)
        
        # 删除格式化后的文档
        formatted_doc_path = os.path.join(os.path.dirname(__file__), 'formatted_test_doc.docx')
        if os.path.exists(formatted_doc_path):
            os.remove(formatted_doc_path)
    
    def test_fix_errors(self):
        """测试修复错误"""
        # 修复错误
        fixer = FormatFixer(self.test_doc_path)
        fixed_doc_path = fixer.fix_errors(self.errors, self.para_manager)
        
        # 验证修复后的文档
        self.assertTrue(os.path.exists(fixed_doc_path), "修复后的文档不存在")
        
        # 加载修复后的文档
        doc = Document(fixed_doc_path)
        
        # 验证标题段落
        heading = doc.paragraphs[0]
        self.assertEqual(heading.text, "一、结合毕业论文（设")
        self.assertEqual(heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        for run in heading.runs:
            self.assertEqual(run.font.size.pt, 16.0)
        
        # 验证参考文献段落
        ref = doc.paragraphs[3]
        self.assertEqual(ref.text, "参考文献：")
        for run in ref.runs:
            self.assertFalse(run.font.bold)
    
    def test_apply_requirements(self):
        """测试应用格式要求"""
        # 应用格式要求
        fixer = FormatFixer(self.test_doc_path)
        formatted_doc_path = fixer.fix_by_requirements(self.requirements, self.para_manager)
        
        # 验证格式化后的文档
        self.assertTrue(os.path.exists(formatted_doc_path), "格式化后的文档不存在")
        
        # 加载格式化后的文档
        doc = Document(formatted_doc_path)
        
        # 验证标题段落
        heading = doc.paragraphs[0]
        self.assertEqual(heading.text, "一、结合毕业论文（设")
        self.assertEqual(heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        for run in heading.runs:
            self.assertEqual(run.font.size.pt, 16.0)
            self.assertEqual(run.font.name, "黑体")
            self.assertTrue(run.font.bold)
        
        # 验证正文段落
        body = doc.paragraphs[1]
        self.assertEqual(body.text, "大型语言模型（LLM")
        self.assertEqual(body.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        for run in body.runs:
            self.assertEqual(run.font.size.pt, 12.0)
            self.assertEqual(run.font.name, "宋体")
            self.assertFalse(run.font.bold)

if __name__ == '__main__':
    unittest.main()
