import os
import sys
import unittest
import json
import tempfile
import requests
from docx import Document

# 添加项目根目录到系统路径
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo

class TestParaManagerExchange(unittest.TestCase):
    """测试前端和后端段落管理器数据交换"""

    def setUp(self):
        """测试前准备"""
        # API基础URL
        self.base_url = "http://localhost:8000/api"

        # 创建测试文档
        self.test_doc_path = os.path.join(os.path.dirname(__file__), 'test_doc.docx')
        self.create_test_document()

        # 创建测试配置文件
        self.test_config_path = os.path.join(os.path.dirname(__file__), 'test_config.json')
        self.create_test_config()

        # 上传文件到服务器
        self.upload_files()

    def create_test_document(self):
        """创建测试文档"""
        doc = Document()

        # 添加标题段落
        heading = doc.add_paragraph("一、结合毕业论文（设")
        heading.paragraph_format.alignment = 0  # 左对齐
        heading.paragraph_format.line_spacing = 1.15
        for run in heading.runs:
            run.font.size = int(10.5 * 20)  # Pt值是磅值的20倍，需要转换为整数

        # 添加正文段落1
        body1 = doc.add_paragraph("大型语言模型（LLM）是指具有数十亿到数万亿参数的深度学习模型，能够处理和生成自然语言。")
        body1.paragraph_format.alignment = 0  # 左对齐
        for run in body1.runs:
            run.font.size = int(10.5 * 20)

        # 添加正文段落2
        body2 = doc.add_paragraph("这些模型不仅具备广泛的知识，还能理解上下文，生成连贯的文本。")
        body2.paragraph_format.alignment = 0  # 左对齐
        for run in body2.runs:
            run.font.size = int(10.5 * 20)

        # 添加参考文献段落
        ref = doc.add_paragraph("参考文献：")
        ref.paragraph_format.alignment = 0  # 左对齐
        ref.paragraph_format.line_spacing = 1.15
        for run in ref.runs:
            run.font.size = int(10.5 * 20)
            run.font.bold = True

        # 保存文档
        doc.save(self.test_doc_path)

    def create_test_config(self):
        """创建测试配置文件"""
        config = {
            "heading1": {
                "fonts": {
                    "zh_family": "黑体",
                    "en_family": "Times New Roman",
                    "size": 16,
                    "bold": True
                },
                "paragraph_format": {
                    "alignment": "center",
                    "line_spacing": "20pt",
                    "first_line_indent": "0cm"
                }
            },
            "body": {
                "fonts": {
                    "zh_family": "宋体",
                    "en_family": "Times New Roman",
                    "size": 12,
                    "bold": False
                },
                "paragraph_format": {
                    "alignment": "left",
                    "line_spacing": "20pt",
                    "first_line_indent": "2cm"
                }
            }
        }

        with open(self.test_config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=4)

    def upload_files(self):
        """上传文件到服务器"""
        # 上传文档
        with open(self.test_doc_path, 'rb') as f:
            files = {'file': (os.path.basename(self.test_doc_path), f)}
            response = requests.post(f"{self.base_url}/upload-files", files=files)
            if response.status_code == 200:
                self.doc_server_path = response.json()['file']['path']
                print(f"文档上传成功: {self.doc_server_path}")
            else:
                print(f"文档上传失败: {response.text}")
                self.doc_server_path = self.test_doc_path

        # 上传配置文件
        with open(self.test_config_path, 'rb') as f:
            files = {'file': (os.path.basename(self.test_config_path), f)}
            response = requests.post(f"{self.base_url}/upload-format", files=files)
            if response.status_code == 200:
                self.config_server_path = response.json()['file_path']
                print(f"配置文件上传成功: {self.config_server_path}")
            else:
                print(f"配置文件上传失败: {response.text}")
                self.config_server_path = self.test_config_path

    def tearDown(self):
        """测试后清理"""
        # 删除测试文档
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)

        # 删除测试配置文件
        if os.path.exists(self.test_config_path):
            os.remove(self.test_config_path)

    def test_para_manager_exchange(self):
        """测试段落管理器数据交换"""
        # 步骤1: 从后端获取段落管理器
        data = {
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path
        }
        response = requests.post(f"{self.base_url}/check-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("para_manager", result, "应该包含段落管理器")

        # 获取段落管理器
        para_manager = result.get("para_manager", [])
        self.assertGreater(len(para_manager), 0, "段落管理器应该包含段落")

        # 步骤2: 修改段落管理器
        # 修改第一个段落的内容
        if len(para_manager) > 0:
            para_manager[0]["content"] = "修改后的标题"

        # 步骤3: 将修改后的段落管理器发送回后端
        data = {
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path,
            "para_manager": para_manager
        }
        response = requests.post(f"{self.base_url}/apply-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")

        # 验证响应内容类型
        self.assertEqual(response.headers['Content-Type'], 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', "应该返回docx文件")

        # 保存返回的文件
        output_path = os.path.join(os.path.dirname(__file__), 'output_test_doc.docx')
        with open(output_path, 'wb') as f:
            f.write(response.content)

        # 验证文件是否存在
        self.assertTrue(os.path.exists(output_path), "应该生成输出文件")

        # 验证文件大小
        self.assertGreater(os.path.getsize(output_path), 0, "文件大小应该大于0")

        # 步骤4: 验证修改是否生效
        # 打开修改后的文档
        doc = Document(output_path)

        # 验证第一个段落的内容是否已修改
        if len(doc.paragraphs) > 0:
            self.assertEqual(doc.paragraphs[0].text, "修改后的标题", "段落内容应该已修改")

        # 删除输出文件
        if os.path.exists(output_path):
            os.remove(output_path)

    def test_para_manager_consistency(self):
        """测试段落管理器一致性"""
        # 步骤1: 从后端获取段落管理器
        data = {
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path
        }
        response = requests.post(f"{self.base_url}/check-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("para_manager", result, "应该包含段落管理器")

        # 获取段落管理器
        para_manager1 = result.get("para_manager", [])

        # 步骤2: 再次从后端获取段落管理器
        response = requests.post(f"{self.base_url}/check-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("para_manager", result, "应该包含段落管理器")

        # 获取段落管理器
        para_manager2 = result.get("para_manager", [])

        # 步骤3: 验证两次获取的段落管理器是否一致
        self.assertEqual(len(para_manager1), len(para_manager2), "段落数量应该一致")

        # 验证每个段落的内容是否一致
        for i in range(len(para_manager1)):
            self.assertEqual(para_manager1[i]["content"], para_manager2[i]["content"], f"段落{i}的内容应该一致")
            self.assertEqual(para_manager1[i]["type"], para_manager2[i]["type"], f"段落{i}的类型应该一致")

    def test_para_manager_with_different_endpoints(self):
        """测试不同端点使用段落管理器"""
        # 步骤1: 从后端获取段落管理器
        data = {
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path
        }
        response = requests.post(f"{self.base_url}/check-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("para_manager", result, "应该包含段落管理器")

        # 获取段落管理器
        para_manager = result.get("para_manager", [])

        # 步骤2: 使用段落管理器调用段落分析API
        data = {
            "doc_path": self.doc_server_path,
            "para_index": 1,
            "para_manager": para_manager
        }
        response = requests.post(f"{self.base_url}/analyze-paragraph", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("result", result, "应该包含分析结果")

        # 步骤3: 使用段落管理器调用段落增强API
        data = {
            "doc_path": self.doc_server_path,
            "para_indices": [1],
            "para_manager": para_manager
        }
        response = requests.post(f"{self.base_url}/enhance-paragraphs", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("result", result, "应该包含增强结果")

        # 步骤4: 使用段落管理器调用发送消息API
        data = {
            "message": "分析这个文档的格式问题",
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path,
            "para_manager": para_manager
        }
        response = requests.post(f"{self.base_url}/send-message", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("message", result, "应该包含消息")

if __name__ == '__main__':
    unittest.main()
