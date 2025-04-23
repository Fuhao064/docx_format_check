import os
import sys
import unittest
import json
import tempfile
import requests
from docx import Document

# 添加项目根目录到系统路径
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo

class TestAPIEndpoints(unittest.TestCase):
    """测试API端点"""

    def setUp(self):
        """测试前准备"""
        # API基础URL
        self.base_url = "http://localhost:8000/api"

        # 创建测试文档
        self.test_doc_path = os.path.join(os.path.dirname(__file__), 'test_doc.docx')
        self.create_test_document()

        # 创建测试配置文件
        self.test_config_path = os.path.join(os.path.dirname(__file__), 'test_config.json')
        self.create_test_config()

        # 创建段落管理器
        self.para_manager = self.create_para_manager()

        # 上传文件到服务器
        self.upload_files()

    def create_test_document(self):
        """创建测试文档"""
        doc = Document()

        # 添加标题段落
        heading = doc.add_paragraph("一、结合毕业论文（设")
        heading.paragraph_format.alignment = 0  # 左对齐
        heading.paragraph_format.line_spacing = 1.15
        for run in heading.runs:
            run.font.size = int(10.5 * 20)  # Pt值是磅值的20倍，需要转换为整数

        # 添加正文段落1
        body1 = doc.add_paragraph("大型语言模型（LLM）是指具有数十亿到数万亿参数的深度学习模型，能够处理和生成自然语言。")
        body1.paragraph_format.alignment = 0  # 左对齐
        for run in body1.runs:
            run.font.size = int(10.5 * 20)

        # 添加正文段落2
        body2 = doc.add_paragraph("这些模型不仅具备广泛的知识，还能理解上下文，生成连贯的文本。")
        body2.paragraph_format.alignment = 0  # 左对齐
        for run in body2.runs:
            run.font.size = int(10.5 * 20)

        # 添加参考文献段落
        ref = doc.add_paragraph("参考文献：")
        ref.paragraph_format.alignment = 0  # 左对齐
        ref.paragraph_format.line_spacing = 1.15
        for run in ref.runs:
            run.font.size = int(10.5 * 20)
            run.font.bold = True

        # 保存文档
        doc.save(self.test_doc_path)

    def create_test_config(self):
        """创建测试配置文件"""
        config = {
            "heading1": {
                "fonts": {
                    "zh_family": "黑体",
                    "en_family": "Times New Roman",
                    "size": 16,
                    "bold": True
                },
                "paragraph_format": {
                    "alignment": "center",
                    "line_spacing": "20pt",
                    "first_line_indent": "0cm"
                }
            },
            "body": {
                "fonts": {
                    "zh_family": "宋体",
                    "en_family": "Times New Roman",
                    "size": 12,
                    "bold": False
                },
                "paragraph_format": {
                    "alignment": "left",
                    "line_spacing": "20pt",
                    "first_line_indent": "2cm"
                }
            }
        }

        with open(self.test_config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=4)

    def create_para_manager(self):
        """创建段落管理器"""
        para_manager = []

        # 添加标题段落
        para_manager.append({
            "id": "para0",
            "type": "heading1",
            "content": "一、结合毕业论文（设",
            "meta": {
                "paragraph_format": {
                    "alignment": "left",
                    "line_spacing": "1.15"
                },
                "fonts": {
                    "size": 10.5
                }
            }
        })

        # 添加正文段落1
        para_manager.append({
            "id": "para1",
            "type": "body",
            "content": "大型语言模型（LLM）是指具有数十亿到数万亿参数的深度学习模型，能够处理和生成自然语言。",
            "meta": {
                "paragraph_format": {
                    "alignment": "left"
                },
                "fonts": {
                    "size": 10.5
                }
            }
        })

        # 添加正文段落2
        para_manager.append({
            "id": "para2",
            "type": "body",
            "content": "这些模型不仅具备广泛的知识，还能理解上下文，生成连贯的文本。",
            "meta": {
                "paragraph_format": {
                    "alignment": "left"
                },
                "fonts": {
                    "size": 10.5
                }
            }
        })

        # 添加参考文献段落
        para_manager.append({
            "id": "para3",
            "type": "heading1",
            "content": "参考文献：",
            "meta": {
                "paragraph_format": {
                    "alignment": "left",
                    "line_spacing": "1.15"
                },
                "fonts": {
                    "size": 10.5,
                    "bold": True
                }
            }
        })

        return para_manager

    def upload_files(self):
        """上传文件到服务器"""
        # 上传文档
        with open(self.test_doc_path, 'rb') as f:
            files = {'file': (os.path.basename(self.test_doc_path), f)}
            response = requests.post(f"{self.base_url}/upload-files", files=files)
            if response.status_code == 200:
                self.doc_server_path = response.json()['file']['path']
                print(f"文档上传成功: {self.doc_server_path}")
            else:
                print(f"文档上传失败: {response.text}")
                self.doc_server_path = self.test_doc_path

        # 上传配置文件
        with open(self.test_config_path, 'rb') as f:
            files = {'file': (os.path.basename(self.test_config_path), f)}
            response = requests.post(f"{self.base_url}/upload-format", files=files)
            if response.status_code == 200:
                self.config_server_path = response.json()['file_path']
                print(f"配置文件上传成功: {self.config_server_path}")
            else:
                print(f"配置文件上传失败: {response.text}")
                self.config_server_path = self.test_config_path

    def tearDown(self):
        """测试后清理"""
        # 删除测试文档
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)

        # 删除测试配置文件
        if os.path.exists(self.test_config_path):
            os.remove(self.test_config_path)

    def test_check_format_endpoint(self):
        """测试检查格式API端点"""
        # 发送请求
        data = {
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path
        }
        response = requests.post(f"{self.base_url}/check-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("errors", result, "应该包含错误信息")
        self.assertIn("para_manager", result, "应该包含段落管理器")

        # 打印错误信息，便于调试
        print("检测到的错误:")
        for error in result.get("errors", []):
            print(f"- {error.get('message')} (位置: {error.get('location')})")

    def test_send_message_endpoint(self):
        """测试发送消息API端点"""
        # 发送请求
        data = {
            "message": "分析这个文档的格式问题",
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path,
            "para_manager": self.para_manager
        }
        response = requests.post(f"{self.base_url}/send-message", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("message", result, "应该包含消息")

        # 打印响应消息，便于调试
        print("响应消息:")
        print(result.get("message"))

    def test_analyze_paragraph_endpoint(self):
        """测试段落分析API端点"""
        # 发送请求
        data = {
            "doc_path": self.doc_server_path,
            "para_index": 1,
            "context_range": 1,
            "para_manager": self.para_manager
        }
        response = requests.post(f"{self.base_url}/analyze-paragraph", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("result", result, "应该包含分析结果")

        # 验证分析结果
        analysis = result.get("result", {})
        self.assertIn("original", analysis, "应该包含原始内容")
        self.assertIn("suggestions", analysis, "应该包含建议")
        self.assertIn("improved_version", analysis, "应该包含改进版本")

        # 打印分析结果，便于调试
        print("段落分析结果:")
        print(f"原始内容: {analysis.get('original')}")
        print("建议:")
        for suggestion in analysis.get('suggestions', []):
            print(f"- 问题: {suggestion.get('issue')}")
            print(f"  解决方案: {suggestion.get('solution')}")
        print(f"改进版本: {analysis.get('improved_version')}")

    def test_enhance_paragraphs_endpoint(self):
        """测试段落增强API端点"""
        # 发送请求
        data = {
            "doc_path": self.doc_server_path,
            "para_indices": [1],
            "para_manager": self.para_manager
        }
        response = requests.post(f"{self.base_url}/enhance-paragraphs", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")
        result = response.json()
        self.assertTrue(result.get("success"), "应该返回成功状态")
        self.assertIn("result", result, "应该包含增强结果")

        # 验证增强结果
        enhancement = result.get("result", {})
        self.assertIn("enhanced_paragraphs", enhancement, "应该包含增强后的段落")
        self.assertIn("total_paragraphs", enhancement, "应该包含总段落数")
        self.assertIn("processed_count", enhancement, "应该包含处理的段落数")

        # 验证处理的段落数
        self.assertEqual(enhancement.get("processed_count"), 1, "应该处理1个段落")

        # 验证增强后的段落
        enhanced_paras = enhancement.get("enhanced_paragraphs", [])
        self.assertEqual(len(enhanced_paras), 1, "应该有1个增强后的段落")

        # 验证增强后的段落内容
        enhanced_para = enhanced_paras[0]
        self.assertIn("original", enhanced_para, "应该包含原始内容")
        self.assertIn("enhanced", enhanced_para, "应该包含增强后的内容")

        # 打印增强结果，便于调试
        print("段落增强结果:")
        print(f"原始内容: {enhanced_para.get('original')}")
        print(f"增强内容: {enhanced_para.get('enhanced')}")

    def test_apply_format_endpoint(self):
        """测试应用格式API端点"""
        # 发送请求
        data = {
            "doc_path": self.doc_server_path,
            "config_path": self.config_server_path,
            "para_manager": self.para_manager,
            "original_filename": os.path.basename(self.test_doc_path)
        }
        response = requests.post(f"{self.base_url}/apply-format", json=data)

        # 验证响应
        self.assertEqual(response.status_code, 200, "应该返回200状态码")

        # 验证响应内容类型
        self.assertEqual(response.headers['Content-Type'], 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', "应该返回docx文件")

        # 保存返回的文件
        output_path = os.path.join(os.path.dirname(__file__), 'output_test_doc.docx')
        with open(output_path, 'wb') as f:
            f.write(response.content)

        # 验证文件是否存在
        self.assertTrue(os.path.exists(output_path), "应该生成输出文件")

        # 验证文件大小
        self.assertGreater(os.path.getsize(output_path), 0, "文件大小应该大于0")

        # 删除输出文件
        if os.path.exists(output_path):
            os.remove(output_path)

if __name__ == '__main__':
    unittest.main()
