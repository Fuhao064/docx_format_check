from editors.document_marker import mark_document_errors
from docx import Document
import os

# 创建一个测试文档
doc = Document()
doc.add_paragraph("这是一个测试段落。行间距不正确，对齐方式不正确，字体不正确。")
doc.add_paragraph("这是第二个测试段落。")
test_doc_path = "test_document.docx"
doc.save(test_doc_path)

# 创建测试错误列表
test_errors = [
    {
        "message": "'行间距' 不匹配: 要求 固定值 20pt, 实际 未设置明确行间距",
        "location": "段落设置.这是一个测试段落"
    },
    {
        "message": "'对齐方式' 不匹配: 要求 左对齐, 实际 两端对齐",
        "location": "段落设置.这是一个测试段落"
    },
    {
        "message": "'中文字体' 不匹配: 要求 宋体, 实际 未知",
        "location": "字体.这是一个测试段落"
    }
]

# 测试标记文档
output_path = mark_document_errors(test_doc_path, test_errors)
print(f"标记后的文档已保存到: {output_path}")

# 清理测试文件
try:
    os.remove(test_doc_path)
    print(f"已删除测试文档: {test_doc_path}")
except:
    pass
