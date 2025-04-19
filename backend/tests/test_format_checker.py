import unittest
from unittest.mock import patch, MagicMock, mock_open
import sys
import os
import json
import re
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加父级目录到系统路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 导入被测试的模块
import format_checker
from backend.preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo

class TestFormatChecker(unittest.TestCase):
    """测试format_checker模块的功能"""
    
    def setUp(self):
        """测试前的准备工作"""
        # 模拟文档信息
        self.mock_doc_info = {
            'page_width': 21.0,
            'page_height': 29.7,
            'margin_top': 2.5,
            'margin_bottom': 2.5,
            'margin_left': 3.0,
            'margin_right': 2.0
        }
        
        # 模拟格式要求
        self.mock_format = {
            'paper_format': {
                'page_size': {
                    'width': '21cm',
                    'height': '29.7cm'
                },
                'margins': {
                    'top': '2.5cm',
                    'bottom': '2.5cm',
                    'left': '3.0cm',
                    'right': '2.0cm'
                }
            },
            'paragraph_format': {
                'title': {
                    'fonts': {
                        'zh_family': '黑体',
                        'en_family': 'Times New Roman',
                        'size': '18pt'
                    },
                    'paragraph_format': {
                        'alignment': 'center',
                        'line_spacing': '1.5'
                    }
                },
                'heading1': {
                    'fonts': {
                        'zh_family': '黑体',
                        'size': '16pt'
                    }
                },
                'body': {
                    'fonts': {
                        'zh_family': '宋体',
                        'en_family': 'Times New Roman',
                        'size': '12pt'
                    },
                    'paragraph_format': {
                        'first_line_indent': '2em'
                    }
                }
            }
        }
        
        # 创建一个模拟的段落管理器
        self.paragraph_manager = ParagraphManager()
        # 添加模拟段落
        self.paragraph_manager.add_para(
            ParsedParaType.TITLE_ZH, 
            "论文标题", 
            {"fonts": {"zh_family": ["黑体"], "size": 18}, "paragraph_format": {"alignment": "center"}}
        )
        self.paragraph_manager.add_para(
            ParsedParaType.ABSTRACT_ZH, 
            "摘要", 
            {"fonts": {"zh_family": ["宋体"], "size": 16}}
        )
        self.paragraph_manager.add_para(
            ParsedParaType.ABSTRACT_CONTENT_ZH, 
            "这是一个测试摘要内容。", 
            {"fonts": {"zh_family": ["宋体"], "size": 12}}
        )
        self.paragraph_manager.add_para(
            ParsedParaType.KEYWORDS_ZH, 
            "关键词：测试；单元测试；格式", 
            {"fonts": {"zh_family": ["宋体"], "size": 12}}
        )
        self.paragraph_manager.add_para(
            ParsedParaType.HEADING1, 
            "1 引言", 
            {"fonts": {"zh_family": ["黑体"], "size": 16}}
        )
        self.paragraph_manager.add_para(
            ParsedParaType.BODY, 
            "这是正文内容。This is body text.", 
            {"fonts": {"zh_family": ["宋体"], "en_family": ["Times New Roman"], "size": 12}}
        )
    
    def test_is_value_equal(self):
        """测试is_value_equal函数"""
        # 测试字符串比较
        self.assertTrue(format_checker.is_value_equal("center", "CENTER", "alignment"), "应该能够比较不同大小写")
        self.assertTrue(format_checker.is_value_equal("2.5cm", "2.5", "margin"), "应该能够比较带单位的数值")
        
        # 测试数值比较
        self.assertTrue(format_checker.is_value_equal(2.5, 2.51, "margin"), "应该允许小误差")  # 允许小误差
        self.assertTrue(format_checker.is_value_equal("2.5", 2.5, "margin"), "应该能够比较字符串和数字")
        
        # 测试负面案例
        self.assertFalse(format_checker.is_value_equal("center", "left", "alignment"), "不同值应该返回False")
        self.assertFalse(format_checker.is_value_equal("3.0cm", 2.5, "margin"), "差异较大的值应该返回False")
    
    def test_extract_number(self):
        """测试extract_number函数"""
        self.assertEqual(format_checker.extract_number("2.5cm"), 2.5)
        self.assertEqual(format_checker.extract_number("2.5"), 2.5)
        self.assertEqual(format_checker.extract_number(2.5), 2.5)
        self.assertEqual(format_checker.extract_number("-2.5"), -2.5)
        self.assertIsNone(format_checker.extract_number("no-number"))
    
    def test_check_paper_format(self):
        """测试check_paper_format函数"""
        # 测试符合要求的情况
        errors = format_checker.check_paper_format(self.mock_doc_info, self.mock_format)
        self.assertEqual(len(errors), 0)
        
        # 测试不符合要求的情况
        mock_doc_info_wrong = self.mock_doc_info.copy()
        mock_doc_info_wrong['page_width'] = 25.0  # 与要求不符
        errors = format_checker.check_paper_format(mock_doc_info_wrong, self.mock_format)
        self.assertGreater(len(errors), 0)
        self.assertIn('页面大小不符合要求', errors[0]['message'])
    
    def test_check_abstract(self):
        """测试check_abstract函数"""
        # 创建一个新的带有正确摘要类型的段落管理器
        abstract_manager = ParagraphManager()
        abstract_manager.add_para(
            ParsedParaType.ABSTRACT_ZH, 
            "摘要", 
            {"fonts": {"zh_family": ["宋体"], "size": 16}}
        )
        abstract_manager.add_para(
            ParsedParaType.ABSTRACT_CONTENT_ZH, 
            "这是一个完整的测试摘要内容，长度应该足够通过检测。", 
            {"fonts": {"zh_family": ["宋体"], "size": 12}}
        )
        
        # 打印段落类型，用于调试
        print("\n调试信息 - 摘要测试:")
        for i, para in enumerate(abstract_manager.paragraphs):
            print(f"段落 {i}: 类型={para.type.value}, 内容={para.content[:20]}...")
        
        # 测试正常情况
        errors = format_checker.check_abstract(abstract_manager)
        if errors:
            print("摘要检查错误:", errors)
        self.assertEqual(len(errors), 0)
        
        # 测试没有摘要的情况
        no_abstract_manager = ParagraphManager()
        no_abstract_manager.add_para(ParsedParaType.TITLE_ZH, "标题", {})
        no_abstract_manager.add_para(ParsedParaType.BODY, "正文", {})
        errors = format_checker.check_abstract(no_abstract_manager)
        self.assertGreater(len(errors), 0)
        self.assertIn('缺少摘要', errors[0]['message'])
    
    def test_check_keywords(self):
        """测试check_keywords函数"""
        # 测试正常情况
        errors = format_checker.check_keywords(self.paragraph_manager)
        self.assertEqual(len(errors), 0)
        
        # 测试没有关键词的情况
        no_keywords_manager = ParagraphManager()
        no_keywords_manager.add_para(ParsedParaType.TITLE_ZH, "标题", {})
        no_keywords_manager.add_para(ParsedParaType.ABSTRACT_ZH, "摘要", {})
        no_keywords_manager.add_para(ParsedParaType.BODY, "正文", {})
        errors = format_checker.check_keywords(no_keywords_manager)
        self.assertGreater(len(errors), 0)
        self.assertIn('缺少关键词', errors[0]['message'])
        
        # 测试关键词不足的情况
        few_keywords_manager = ParagraphManager()
        few_keywords_manager.add_para(ParsedParaType.KEYWORDS_ZH, "关键词：测试", {})
        errors = format_checker.check_keywords(few_keywords_manager)
        self.assertGreater(len(errors), 0)
        self.assertIn('关键词数量不足', errors[0]['message'])
    
    @patch('docx.Document')
    def test_check_reference_format(self, mock_document):
        """测试check_reference_format函数"""
        # 模拟文档对象
        mock_doc = MagicMock()
        mock_document.return_value = mock_doc
        
        # 模拟段落对象
        mock_paragraphs = [
            MagicMock(text="参考文献"),
            MagicMock(text="[1] 作者. 论文题目[J]. 期刊名称, 2020, 40(1): 100-110."),
            MagicMock(text="[2] 作者. 书名[M]. 北京: 出版社, 2019.")
        ]
        mock_doc.paragraphs = mock_paragraphs
        
        # 测试GB/T格式
        mock_format = {'reference_format': {'citation_style': 'GB/T 7714'}}
        errors = format_checker.check_reference_format("dummy_path.docx", mock_format)
        self.assertEqual(len(errors), 0)
        
        # 测试不支持的引用格式
        mock_format = {'reference_format': {'citation_style': 'unknown_style'}}
        errors = format_checker.check_reference_format("dummy_path.docx", mock_format)
        self.assertGreater(len(errors), 0)
        self.assertIn('不支持的引用样式', errors[0]['message'])
    
    def test_check_main_format(self):
        """测试check_main_format函数"""
        # 创建一个与格式要求完全匹配的段落管理器
        matched_manager = ParagraphManager()
        matched_manager.add_para(
            ParsedParaType.TITLE_ZH, 
            "论文标题", 
            {"fonts": {"zh_family": ["黑体"], "en_family": ["Times New Roman"], "size": 18}, 
             "paragraph_format": {"alignment": "center", "line_spacing": "1.5"}}
        )
        matched_manager.add_para(
            ParsedParaType.HEADING1, 
            "1 引言", 
            {"fonts": {"zh_family": ["黑体"], "size": 16}}
        )
        matched_manager.add_para(
            ParsedParaType.BODY, 
            "这是正文内容。This is body text.", 
            {"fonts": {"zh_family": ["宋体"], "en_family": ["Times New Roman"], "size": 12},
             "paragraph_format": {"first_line_indent": "2em"}}
        )
        
        # 打印格式要求和段落元数据，用于调试
        print("\n调试信息 - 主格式测试:")
        print("标题格式要求:", self.mock_format['paragraph_format']['title'])
        print("正文格式要求:", self.mock_format['paragraph_format']['body'])
        print("标题段落元数据:", matched_manager.paragraphs[0].meta)
        
        # 测试所有格式正确的情况 - 关闭摘要和关键词检查
        errors = format_checker.check_main_format(matched_manager, self.mock_format, check_abstract_keywords=False)
        if errors:
            print("主格式检查错误:", errors)
        self.assertEqual(len(errors), 0)
        
        # 测试段落格式错误的情况
        wrong_format_manager = ParagraphManager()
        wrong_format_manager.add_para(
            ParsedParaType.TITLE_ZH, 
            "错误格式标题", 
            {"fonts": {"zh_family": ["宋体"], "size": 14}}  # 与要求不符
        )
        errors = format_checker.check_main_format(wrong_format_manager, self.mock_format, check_abstract_keywords=False)
        self.assertGreater(len(errors), 0)

    @patch('format_checker.FormatAgent')
    def test_remark_para_type(self, mock_format_agent):
        """测试remark_para_type函数"""
        # 创建模拟的FormatAgent
        mock_agent = MagicMock()
        mock_format_agent.return_value = mock_agent
        
        # 模拟predict_location方法
        mock_agent.predict_location.return_value = {"location": "heading1"}
        
        # 创建测试文档路径
        test_doc_path = "test_doc.docx"
        
        # 测试前打补丁
        with patch('extract_para_info.extract_para_format_info', return_value=self.paragraph_manager), \
             patch('docx_parser.extract_doc_content', return_value="测试文档内容"):
            
            # 调用被测试函数
            result = format_checker.remark_para_type(test_doc_path, mock_agent)
            
            # 验证结果
            self.assertIsInstance(result, ParagraphManager)
            mock_agent.predict_location.assert_called()

    @patch('format_checker.check_paper_format')
    @patch('format_checker.check_main_format')
    @patch('format_checker.check_table_format')
    @patch('format_checker.check_figure_format')
    @patch('format_checker.check_reference_format')
    @patch('format_checker.remark_para_type')
    @patch('format_checker.load_config')
    @patch('docx_parser.extract_section_info')
    @patch('extract_para_info.extract_para_format_info')
    def test_check_format(self, mock_extract_para, mock_extract_section, mock_load_config,
                         mock_remark_para_type, mock_check_reference, mock_check_figure,
                         mock_check_table, mock_check_main, mock_check_paper):
        """测试check_format函数"""
        # 设置各个模拟函数的返回值
        mock_load_config.return_value = self.mock_format
        mock_extract_section.return_value = self.mock_doc_info
        mock_extract_para.return_value = self.paragraph_manager
        mock_remark_para_type.return_value = self.paragraph_manager
        
        # 设置各个检查函数的返回值
        mock_check_paper.return_value = []
        mock_check_main.return_value = []
        mock_check_table.return_value = []
        mock_check_figure.return_value = []
        mock_check_reference.return_value = []
        
        # 创建测试参数
        test_doc_path = "test_doc.docx"
        test_config_path = "test_config.json"
        mock_format_agent = MagicMock()
        
        # 调用被测试函数
        errors, manager = format_checker.check_format(test_doc_path, test_config_path, mock_format_agent)
        
        # 验证结果
        self.assertEqual(len(errors), 0)
        self.assertEqual(manager, self.paragraph_manager)
        
        # 验证所有模拟函数都被调用
        mock_load_config.assert_called_once_with(test_config_path)
        mock_extract_section.assert_called_once_with(test_doc_path)
        mock_extract_para.assert_called_once()
        mock_remark_para_type.assert_called_once()
        mock_check_paper.assert_called_once()
        mock_check_main.assert_called_once()
        mock_check_table.assert_called_once()
        mock_check_figure.assert_called_once()
        mock_check_reference.assert_called_once()

if __name__ == '__main__':
    unittest.main() 