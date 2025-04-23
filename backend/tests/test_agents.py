import os
import sys
import unittest
import json
import tempfile
from docx import Document

# 添加项目根目录到系统路径
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from agents.format_agent import FormatAgent
from agents.advice_agent import AdviceAgent
from agents.editor_agent import EditorAgent
from agents.communicate_agent import CommunicateAgent
from preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo
from utils.config_utils import load_config

class TestAgents(unittest.TestCase):
    """测试修改后的代理功能"""

    def setUp(self):
        """测试前准备"""
        # 创建测试文档
        self.test_doc_path = os.path.join(os.path.dirname(__file__), 'test_doc.docx')
        self.create_test_document()

        # 创建测试配置文件
        self.test_config_path = os.path.join(os.path.dirname(__file__), 'test_config.json')
        self.create_test_config()

        # 创建段落管理器
        self.para_manager = ParagraphManager()
        self.para_manager.add_para(
            para_type=ParsedParaType.HEADING1,
            content="一、结合毕业论文（设",
            meta={"paragraph_format": {"alignment": "left", "line_spacing": "1.15"}, "fonts": {"size": 10.5}}
        )
        self.para_manager.add_para(
            para_type=ParsedParaType.BODY,
            content="大型语言模型（LLM）是指具有数十亿到数万亿参数的深度学习模型，能够处理和生成自然语言。",
            meta={"paragraph_format": {"alignment": "left"}, "fonts": {"size": 10.5}}
        )
        self.para_manager.add_para(
            para_type=ParsedParaType.BODY,
            content="这些模型不仅具备广泛的知识，还能理解上下文，生成连贯的文本。",
            meta={"paragraph_format": {"alignment": "left"}, "fonts": {"size": 10.5}}
        )
        self.para_manager.add_para(
            para_type=ParsedParaType.HEADING1,
            content="参考文献：",
            meta={"paragraph_format": {"alignment": "left", "line_spacing": "1.15"}, "fonts": {"size": 10.5, "bold": True}}
        )

        # 初始化代理
        self.format_agent = FormatAgent("qwen-plus")
        self.advice_agent = AdviceAgent("qwen-plus")
        self.editor_agent = EditorAgent("qwen-plus")
        self.communicate_agent = CommunicateAgent("qwen-plus")

    def create_test_document(self):
        """创建测试文档"""
        doc = Document()

        # 添加标题段落
        heading = doc.add_paragraph("一、结合毕业论文（设")
        heading.paragraph_format.alignment = 0  # 左对齐
        heading.paragraph_format.line_spacing = 1.15
        for run in heading.runs:
            run.font.size = int(10.5 * 20)  # Pt值是磅值的20倍，需要转换为整数

        # 添加正文段落1
        body1 = doc.add_paragraph("大型语言模型（LLM）是指具有数十亿到数万亿参数的深度学习模型，能够处理和生成自然语言。")
        body1.paragraph_format.alignment = 0  # 左对齐
        for run in body1.runs:
            run.font.size = int(10.5 * 20)

        # 添加正文段落2
        body2 = doc.add_paragraph("这些模型不仅具备广泛的知识，还能理解上下文，生成连贯的文本。")
        body2.paragraph_format.alignment = 0  # 左对齐
        for run in body2.runs:
            run.font.size = int(10.5 * 20)

        # 添加参考文献段落
        ref = doc.add_paragraph("参考文献：")
        ref.paragraph_format.alignment = 0  # 左对齐
        ref.paragraph_format.line_spacing = 1.15
        for run in ref.runs:
            run.font.size = int(10.5 * 20)
            run.font.bold = True

        # 保存文档
        doc.save(self.test_doc_path)

    def create_test_config(self):
        """创建测试配置文件"""
        config = {
            "heading1": {
                "fonts": {
                    "zh_family": "黑体",
                    "en_family": "Times New Roman",
                    "size": 16,
                    "bold": True
                },
                "paragraph_format": {
                    "alignment": "center",
                    "line_spacing": "20pt",
                    "first_line_indent": "0cm"
                }
            },
            "body": {
                "fonts": {
                    "zh_family": "宋体",
                    "en_family": "Times New Roman",
                    "size": 12,
                    "bold": False
                },
                "paragraph_format": {
                    "alignment": "left",
                    "line_spacing": "20pt",
                    "first_line_indent": "2cm"
                }
            }
        }

        with open(self.test_config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=4)

    def tearDown(self):
        """测试后清理"""
        # 删除测试文档
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)

        # 删除测试配置文件
        if os.path.exists(self.test_config_path):
            os.remove(self.test_config_path)

    def test_format_agent_analyze_paragraph(self):
        """测试FormatAgent的段落分析功能"""
        # 加载配置
        config = load_config(self.test_config_path)

        # 分析标题段落
        heading_para = self.para_manager.paragraphs[0]
        errors = self.format_agent.analyze_paragraph(heading_para, config)

        # 验证错误检测
        self.assertTrue(len(errors) > 0, "应该检测到格式错误")

        # 检查是否检测到字体大小错误
        font_size_error = False
        alignment_error = False
        for error in errors:
            if "字号" in error["message"]:
                font_size_error = True
            if "对齐方式" in error["message"]:
                alignment_error = True

        self.assertTrue(font_size_error, "应该检测到字体大小错误")
        self.assertTrue(alignment_error, "应该检测到对齐方式错误")

    def test_format_agent_check_paragraph_manager(self):
        """测试FormatAgent的段落管理器检查功能"""
        # 检查段落管理器
        errors = self.format_agent.check_paragraph_manager(self.para_manager, self.test_config_path)

        # 验证错误检测
        self.assertTrue(len(errors) > 0, "应该检测到格式错误")

        # 打印错误信息，便于调试
        print("检测到的错误:")
        for error in errors:
            print(f"- {error['message']} (位置: {error['location']})")

    def test_format_agent_fix_paragraph_manager(self):
        """测试FormatAgent的段落管理器修复功能"""
        # 修复段落管理器
        fixed_manager = self.format_agent.fix_paragraph_manager(self.para_manager, self.test_config_path)

        # 验证修复结果
        self.assertEqual(len(fixed_manager.paragraphs), len(self.para_manager.paragraphs), "段落数量应该保持不变")

        # 检查标题段落是否已修复
        heading_para = fixed_manager.paragraphs[0]
        self.assertEqual(heading_para.type, ParsedParaType.HEADING1, "段落类型应该保持不变")
        self.assertEqual(heading_para.content, "一、结合毕业论文（设", "段落内容应该保持不变")

        # 检查标题段落的格式是否已修复
        heading_meta = heading_para.meta
        self.assertIn("paragraph_format", heading_meta, "应该包含段落格式信息")
        self.assertIn("fonts", heading_meta, "应该包含字体信息")

        # 验证段落格式
        para_format = heading_meta["paragraph_format"]
        self.assertEqual(para_format.get("alignment"), "center", "对齐方式应该已修复")

        # 验证字体格式
        fonts = heading_meta["fonts"]
        self.assertEqual(fonts.get("size"), 16, "字体大小应该已修复")
        self.assertEqual(fonts.get("bold"), True, "加粗设置应该已修复")

    def test_advice_agent_analyze_paragraph_manager(self):
        """测试AdviceAgent的段落管理器分析功能"""
        # 分析段落
        result = self.advice_agent.analyze_paragraph_manager(self.para_manager, 1, 1)

        # 验证分析结果
        self.assertIn("original", result, "应该包含原始内容")
        self.assertIn("suggestions", result, "应该包含建议")
        self.assertIn("improved_version", result, "应该包含改进版本")
        self.assertIn("para_type", result, "应该包含段落类型")

        # 打印分析结果，便于调试
        print("段落分析结果:")
        print(f"原始内容: {result.get('original')}")
        print(f"段落类型: {result.get('para_type')}")
        print("建议:")
        for suggestion in result.get('suggestions', []):
            print(f"- 问题: {suggestion.get('issue')}")
            print(f"  解决方案: {suggestion.get('solution')}")
        print(f"改进版本: {result.get('improved_version')}")

    def test_editor_agent_enhance_paragraph_manager(self):
        """测试EditorAgent的段落管理器增强功能"""
        # 增强段落
        result = self.editor_agent.enhance_paragraph_manager(self.para_manager, [1])

        # 验证增强结果
        self.assertIn("enhanced_paragraphs", result, "应该包含增强后的段落")
        self.assertIn("total_paragraphs", result, "应该包含总段落数")
        self.assertIn("processed_count", result, "应该包含处理的段落数")

        # 验证处理的段落数
        self.assertEqual(result.get("processed_count"), 1, "应该处理1个段落")

        # 验证增强后的段落
        enhanced_paras = result.get("enhanced_paragraphs", [])
        self.assertEqual(len(enhanced_paras), 1, "应该有1个增强后的段落")

        # 验证增强后的段落内容
        enhanced_para = enhanced_paras[0]
        self.assertIn("original", enhanced_para, "应该包含原始内容")
        self.assertIn("enhanced", enhanced_para, "应该包含增强后的内容")
        self.assertIn("para_type", enhanced_para, "应该包含段落类型")
        self.assertIn("index", enhanced_para, "应该包含段落索引")

        # 打印增强结果，便于调试
        print("段落增强结果:")
        print(f"原始内容: {enhanced_para.get('original')}")
        print(f"增强内容: {enhanced_para.get('enhanced')}")

    def test_editor_agent_create_modified_paragraph_manager(self):
        """测试EditorAgent的创建修改后的段落管理器功能"""
        # 创建修改列表
        modifications = [
            {
                "index": 1,
                "content": "这是修改后的内容，用于测试EditorAgent的创建修改后的段落管理器功能。"
            }
        ]

        # 创建修改后的段落管理器
        modified_manager = self.editor_agent.create_modified_paragraph_manager(self.para_manager, modifications)

        # 验证修改后的段落管理器
        self.assertEqual(len(modified_manager.paragraphs), len(self.para_manager.paragraphs), "段落数量应该保持不变")

        # 验证修改后的段落内容
        modified_para = modified_manager.paragraphs[1]
        self.assertEqual(modified_para.content, modifications[0]["content"], "段落内容应该已修改")
        self.assertEqual(modified_para.type, self.para_manager.paragraphs[1].type, "段落类型应该保持不变")

    def test_communicate_agent_get_response(self):
        """测试CommunicateAgent的获取响应功能"""
        # 获取响应
        doc_content = "这是测试文档内容，用于测试CommunicateAgent的获取响应功能。"
        response = self.communicate_agent.get_response("分析这个段落的格式问题", doc_content, self.para_manager, self.test_config_path)

        # 验证响应
        self.assertIsNotNone(response, "应该返回响应")
        self.assertIsInstance(response, str, "响应应该是字符串")

        # 打印响应，便于调试
        print("CommunicateAgent响应:")
        print(response)

if __name__ == '__main__':
    unittest.main()
