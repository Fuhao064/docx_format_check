import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION, WD_SECTION
import json
import os
from typing import Dict, List, Union, Optional, Tuple
from para_type import ParsedParaType, ParagraphManager, ParaInfo
import re
from utils import parse_color

class FormatEditorWithErrors:
    """
    文档格式编辑器，用于修改docx文档的格式，包括段落样式、字体样式等
    同时支持标记格式错误
    """
    def __init__(self, doc_path: str):
        """
        初始化格式编辑器
        
        Args:
            doc_path: 文档路径
        """
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        
    def save(self, output_path: Optional[str] = None):
        """
        保存文档
        
        Args:
            output_path: 输出路径，如果为None则覆盖原文件
        """
        if output_path is None:
            output_path = self.doc_path
        self.doc.save(output_path)
    
    def _parse_color(self, color_str: str) -> Tuple[int, int, int]:
        return parse_color(color_str)
    
    def mark_errors_in_document(self, errors: List[Dict], output_path: Optional[str] = None):
        """
        在文档中标记格式错误
        
        Args:
            errors: 格式错误列表，每个错误是一个字典，包含type, location, message等字段
            output_path: 输出路径，如果为None则覆盖原文件
        """
        # 按段落索引分组错误
        para_errors = {}
        for error in errors:
            error_type = error.get('type', '')
            location = error.get('location', '')
            message = error.get('message', '')
            
            # 提取段落索引
            para_index = -1
            if error_type == 'paragraph':
                # 从location中提取段落索引，格式如"第10段：正文"
                match = re.search(r'第(\d+)段', location)
                if match:
                    para_index = int(match.group(1)) - 1  # 转为0-based索引
            
            # 将错误添加到对应段落
            if para_index >= 0 and para_index < len(self.doc.paragraphs):
                if para_index not in para_errors:
                    para_errors[para_index] = []
                para_errors[para_index].append(message)
        
        # 处理段落错误
        for para_index, error_messages in para_errors.items():
            paragraph = self.doc.paragraphs[para_index]
            
            # 清除段落中的所有runs
            text = paragraph.text
            for i in range(len(paragraph.runs)-1, -1, -1):
                paragraph.runs[i].text = ""
            
            # 添加带有红色标记的文本
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # 设置为红色
            
            # 添加错误注释
            error_text = "\n【错误】: " + "; ".join(error_messages)
            comment_run = paragraph.add_run(error_text)
            comment_run.font.color.rgb = RGBColor(255, 0, 0)  # 设置为红色
            comment_run.font.bold = True
        
        # 保存文档
        self.save(output_path)
    
    def mark_errors_with_comments(self, errors: List[Dict], output_path: Optional[str] = None):
        """
        在文档中使用注释标记格式错误
        
        Args:
            errors: 格式错误列表，每个错误是一个字典，包含type, location, message等字段
            output_path: 输出路径，如果为None则覆盖原文件
        """
        # 按段落索引分组错误
        para_errors = {}
        for error in errors:
            error_type = error.get('type', '')
            location = error.get('location', '')
            message = error.get('message', '')
            
            # 提取段落索引
            para_index = -1
            if error_type == 'paragraph':
                # 从location中提取段落索引，格式如"第10段：正文"
                match = re.search(r'第(\d+)段', location)
                if match:
                    para_index = int(match.group(1)) - 1  # 转为0-based索引
            
            # 将错误添加到对应段落
            if para_index >= 0 and para_index < len(self.doc.paragraphs):
                if para_index not in para_errors:
                    para_errors[para_index] = []
                para_errors[para_index].append(message)
        
        # 处理段落错误
        for para_index, error_messages in para_errors.items():
            paragraph = self.doc.paragraphs[para_index]
            
            # 将段落文本设置为红色
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)  # 设置为红色
            
            # 如果段落没有runs，添加一个红色的run
            if len(paragraph.runs) == 0:
                run = paragraph.add_run(paragraph.text)
                run.font.color.rgb = RGBColor(255, 0, 0)  # 设置为红色
        
        # 保存文档
        self.save(output_path)

def mark_errors_in_document(doc_path: str, errors: List[Dict], output_path: Optional[str] = None):
    """
    在文档中标记格式错误
    
    Args:
        doc_path: 文档路径
        errors: 格式错误列表
        output_path: 输出路径，如果为None则覆盖原文件
    """
    editor = FormatEditorWithErrors(doc_path)
    editor.mark_errors_in_document(errors, output_path)
    return True