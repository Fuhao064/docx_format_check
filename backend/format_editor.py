from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
import json
import os
import re
import requests
from typing import Dict, List, Optional, Union, Tuple
from para_type import ParagraphManager, ParsedParaType, ParaInfo

# 全局映射字典
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
}

LINE_SPACING_MAP = {
    "1.0": 1.0,
    "1.5": 1.5,
    "2.0": 2.0,
    "single": 1.0,
    "double": 2.0
}

def load_config(config_path: str) -> Dict:
    """加载配置文件"""
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def apply_font_settings(run, font_settings: Dict, is_chinese: bool = True):
    """
    应用字体设置到文本运行对象
    
    Args:
        run: 文本运行对象
        font_settings: 字体设置字典
        is_chinese: 是否为中文
    """
    # 设置字体族
    if is_chinese and 'zh_family' in font_settings:
        run.font.name = font_settings['zh_family']
    elif not is_chinese and 'en_family' in font_settings:
        run.font.name = font_settings['en_family']
    
    # 设置字体大小
    if 'size' in font_settings:
        size_text = font_settings['size']
        if isinstance(size_text, str) and 'pt' in size_text:
            size = float(size_text.replace('pt', ''))
            run.font.size = Pt(size)
        else:
            run.font.size = Pt(float(size_text))
    
    # 设置加粗
    if 'bold' in font_settings:
        run.font.bold = font_settings['bold']
    
    # 设置斜体
    if 'italic' in font_settings:
        run.font.italic = font_settings['italic']
    
    # 设置全部大写
    if 'isAllCaps' in font_settings:
        run.font.all_caps = font_settings['isAllCaps']
    
    # 设置字体颜色
    if 'color' in font_settings:
        color = font_settings['color'].lstrip('#')
        if len(color) == 6:
            r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)

def apply_paragraph_format(paragraph, format_settings: Dict):
    """
    应用段落格式设置
    
    Args:
        paragraph: 段落对象
        format_settings: 格式设置字典
    """
    # 设置行间距
    if 'line_spacing' in format_settings:
        spacing = format_settings['line_spacing']
        if spacing in LINE_SPACING_MAP:
            if spacing == "1.5":
                paragraph.paragraph_format.line_spacing = 1.5
            elif spacing == "2.0" or spacing == "double":
                paragraph.paragraph_format.line_spacing = 2.0
            else:
                paragraph.paragraph_format.line_spacing = 1.0
        else:
            # 尝试直接设置具体数值
            try:
                paragraph.paragraph_format.line_spacing = float(spacing)
            except ValueError:
                paragraph.paragraph_format.line_spacing = 1.5  # 默认值
    
    # 设置对齐方式
    if 'alignment' in format_settings:
        alignment = format_settings['alignment'].lower()
        if alignment in ALIGNMENT_MAP:
            paragraph.paragraph_format.alignment = ALIGNMENT_MAP[alignment]
    
    # 设置缩进
    if 'indentation' in format_settings:
        indentation = format_settings['indentation']
        
        # 首行缩进
        if 'first_line' in indentation:
            first_line = indentation['first_line']
            if isinstance(first_line, str) and 'cm' in first_line:
                paragraph.paragraph_format.first_line_indent = Cm(float(first_line.replace('cm', '')))
            else:
                paragraph.paragraph_format.first_line_indent = Cm(float(first_line))
        
        # 左缩进
        if 'left' in indentation:
            left = indentation['left']
            if isinstance(left, str) and 'cm' in left:
                paragraph.paragraph_format.left_indent = Cm(float(left.replace('cm', '')))
            else:
                paragraph.paragraph_format.left_indent = Cm(float(left))
        
        # 右缩进
        if 'right' in indentation:
            right = indentation['right']
            if isinstance(right, str) and 'cm' in right:
                paragraph.paragraph_format.right_indent = Cm(float(right.replace('cm', '')))
            else:
                paragraph.paragraph_format.right_indent = Cm(float(right))
        
        # 段前距
        if 'space_before' in indentation:
            space_before = indentation['space_before']
            if isinstance(space_before, str) and 'cm' in space_before:
                paragraph.paragraph_format.space_before = Cm(float(space_before.replace('cm', '')))
            else:
                paragraph.paragraph_format.space_before = Cm(float(space_before))
        
        # 段后距
        if 'space_after' in indentation:
            space_after = indentation['space_after']
            if isinstance(space_after, str) and 'cm' in space_after:
                paragraph.paragraph_format.space_after = Cm(float(space_after.replace('cm', '')))
            else:
                paragraph.paragraph_format.space_after = Cm(float(space_after))

def is_chinese_char(char: str) -> bool:
    """判断字符是否为中文"""
    return '\u4e00' <= char <= '\u9fff'

def split_text_by_language(text: str) -> List[Tuple[str, bool]]:
    """
    将文本按中英文分段
    
    Args:
        text: 要分段的文本
    
    Returns:
        List[Tuple[str, bool]]: 文本段列表，每项是(文本段, 是否中文)
    """
    result = []
    current_text = ""
    current_is_chinese = None
    
    for char in text:
        is_chinese = is_chinese_char(char)
        
        # 首个字符或当前字符类型与前面相同
        if current_is_chinese is None or is_chinese == current_is_chinese:
            current_text += char
            current_is_chinese = is_chinese
        else:
            # 字符类型发生变化，保存前一段并开始新段
            if current_text:
                result.append((current_text, current_is_chinese))
            current_text = char
            current_is_chinese = is_chinese
    
    # 添加最后一段
    if current_text:
        result.append((current_text, current_is_chinese))
    
    return result

def get_openai_image_caption(image_path: str) -> str:
    """
    使用OpenAI API获取图片题注
    
    Args:
        image_path: 图片路径
        
    Returns:
        str: 生成的题注
    """
    try:
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            return "图片题注（需要设置OpenAI API密钥）"
        
        # 读取图片并转为base64
        import base64
        with open(image_path, "rb") as image_file:
            encoded_image = base64.b64encode(image_file.read()).decode('utf-8')
        
        # 调用OpenAI API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        payload = {
            "model": "gpt-4-vision-preview",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "请为这张图片生成一个简短的中文学术题注，不超过20字："
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded_image}"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 100
        }
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload
        )
        
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            return f"图 {os.path.basename(image_path)}"
        
    except Exception as e:
        return f"图 {os.path.basename(image_path)}"

def set_paper_format(doc: Document, config: Dict):
    """
    设置文档页面格式
    
    Args:
        doc: 文档对象
        config: 配置字典
    """
    paper_config = config.get('paper', {})
    
    # 获取文档的节
    section = doc.sections[0]
    
    # 设置纸张大小
    if 'size' in paper_config:
        size = paper_config['size']
        # 标准纸张大小设置
        if size == 'A4':
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        elif size == 'Letter':
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        # 可以添加更多纸张尺寸...
    
    # 设置页面方向
    if 'orientation' in paper_config:
        orientation = paper_config['orientation']
        if orientation.lower() == 'portrait':
            section.orientation = WD_ORIENT.PORTRAIT
        elif orientation.lower() == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
    
    # 设置页边距
    if 'margins' in paper_config:
        margins = paper_config['margins']
        
        for margin_name, margin_value in margins.items():
            if isinstance(margin_value, str) and 'cm' in margin_value:
                margin_cm = float(margin_value.replace('cm', ''))
            else:
                margin_cm = float(margin_value)
            
            if margin_name == 'top':
                section.top_margin = Cm(margin_cm)
            elif margin_name == 'bottom':
                section.bottom_margin = Cm(margin_cm)
            elif margin_name == 'left':
                section.left_margin = Cm(margin_cm)
            elif margin_name == 'right':
                section.right_margin = Cm(margin_cm)

def format_document(config: Dict, para_manager: ParagraphManager, output_path: str = None, doc_path: str = None) -> str:
    """
    根据配置和段落管理器格式化文档
    
    Args:
        config: 格式配置字典或配置文件路径
        para_manager: 段落管理器实例
        output_path: 输出文档路径，如果为None则生成临时文件
        doc_path: 可选的输入文档路径，如果提供则在该文档基础上修改
        
    Returns:
        str: 输出文档的路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)
    
    # 创建或打开文档
    if doc_path:
        doc = Document(doc_path)
        # 清空文档内容
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
    else:
        doc = Document()
    
    # 设置页面格式
    set_paper_format(doc, config)
    
    # 按照段落类型和配置格式化文档
    for para_info in para_manager.paragraphs:
        # 创建新段落
        paragraph = doc.add_paragraph()
        
        # 获取段落类型
        para_type = para_info.type.value
        
        # 获取段落格式设置
        if para_type in config:
            format_config = config[para_type]
            
            # 应用段落格式
            if 'paragraph_format' in format_config:
                apply_paragraph_format(paragraph, format_config['paragraph_format'])
            
            # 处理文本内容并应用字体设置
            if 'fonts' in format_config:
                font_settings = format_config['fonts']
                
                # 拆分中英文
                segments = split_text_by_language(para_info.content)
                
                # 为每段添加相应的格式
                for text, is_chinese in segments:
                    run = paragraph.add_run(text)
                    apply_font_settings(run, font_settings, is_chinese)
            else:
                # 无特殊字体设置，使用默认体
                paragraph.add_run(para_info.content)
        else:
            # 使用默认体
            paragraph.add_run(para_info.content)
            
        # 特殊处理图片段落
        if para_type == 'figures' and para_info.meta and para_info.meta.get('image_path'):
            image_path = para_info.meta['image_path']
            
            # 添加图片
            if os.path.exists(image_path):
                # 创建图片段落（不是题注段落）
                img_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['figures']:
                    apply_paragraph_format(img_paragraph, config['figures']['paragraph_format'])
                
                # 添加图片，设置合适的宽度（可根据需要调整）
                img_paragraph.add_run().add_picture(image_path, width=Inches(6))
                
                # 生成题注
                caption = para_info.content
                if not caption:
                    caption = get_openai_image_caption(image_path)
                
                # 添加题注段落
                caption_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['figures']:
                    apply_paragraph_format(caption_paragraph, config['figures']['paragraph_format'])
                
                # 设置题注文本
                caption_run = caption_paragraph.add_run(f"图 {para_info.meta.get('figure_number', '')} {caption}")
                
                # 应用题注字体格式
                if 'caption' in config['figures'] and 'fonts' in config['figures']['caption']:
                    apply_font_settings(caption_run, config['figures']['caption']['fonts'])
    
    # 确定输出路径
    if output_path is None:
        output_path = f"output_{os.path.basename(doc_path)}" if doc_path else "formatted_document.docx"
    
    # 保存文档
    doc.save(output_path)
    
    return output_path

def add_figure_caption(config: Dict, doc_path: str, image_path: str, figure_number: int, caption: Optional[str] = None, 
                      output_path: Optional[str] = None) -> str:
    """
    为文档中的图片添加题注
    
    Args:
        config: 格式配置字典或配置文件路径
        doc_path: 文档路径
        image_path: 图片路径
        figure_number: 图片编号
        caption: 题注文本，如果为None则自动生成
        output_path: 输出文档路径，如果为None则覆盖原文档
        
    Returns:
        str: 输出文档的路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)
    
    # 打开文档
    doc = Document(doc_path)
    
    # 自动生成题注
    if caption is None:
        caption = get_openai_image_caption(image_path)
    
    # 添加图片
    paragraph = doc.add_paragraph()
    if 'paragraph_format' in config['figures']:
        apply_paragraph_format(paragraph, config['figures']['paragraph_format'])
    
    # 添加图片，设置合适的宽度
    paragraph.add_run().add_picture(image_path, width=Inches(6))
    
    # 获取题注位置设置
    position = "below"
    if 'caption' in config['figures'] and 'position' in config['figures']['caption']:
        position = config['figures']['caption']['position']
    
    # 如果题注在上方，则在图片前添加题注
    if position.lower() == "above":
        # 在当前段落之前插入新段落
        p_index = len(doc.paragraphs) - 1
        for i in range(p_index, 0, -1):
            doc.paragraphs[i]._p.addprevious(doc.paragraphs[i-1]._p)
    else:
        # 添加下方题注
        caption_paragraph = doc.add_paragraph()
        if 'paragraph_format' in config['figures']:
            apply_paragraph_format(caption_paragraph, config['figures']['paragraph_format'])
        
        caption_text = f"图 {figure_number} {caption}"
        caption_run = caption_paragraph.add_run(caption_text)
        
        # 应用题注字体格式
        if 'caption' in config['figures'] and 'fonts' in config['figures']['caption']:
            apply_font_settings(caption_run, config['figures']['caption']['fonts'])
    
    # 保存文档
    output_file = output_path or doc_path
    doc.save(output_file)
    
    return output_file

def format_table_caption(config: Dict, doc_path: str, table_number: int, caption: str, output_path: Optional[str] = None) -> str:
    """
    为文档中的表格添加题注
    
    Args:
        config: 格式配置字典或配置文件路径
        doc_path: 文档路径
        table_number: 表格编号
        caption: 题注文本
        output_path: 输出文档路径，如果为None则覆盖原文档
        
    Returns:
        str: 输出文档的路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)
    
    # 打开文档
    doc = Document(doc_path)
    
    # 获取题注位置设置
    position = "above"
    if 'caption' in config['tables'] and 'position' in config['tables']['caption']:
        position = config['tables']['caption']['position']
    
    # 添加题注段落
    caption_paragraph = doc.add_paragraph()
    if 'paragraph_format' in config['tables']:
        apply_paragraph_format(caption_paragraph, config['tables']['paragraph_format'])
    
    # 设置题注文本
    caption_text = f"表 {table_number} {caption}"
    caption_run = caption_paragraph.add_run(caption_text)
    
    # 应用题注字体格式
    if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
        apply_font_settings(caption_run, config['tables']['caption']['fonts'])
    
    # 保存文档
    output_file = output_path or doc_path
    doc.save(output_file)
    
    return output_file

def generate_formatted_doc(config: Dict, para_manager: ParagraphManager, output_path: str) -> str:
    """
    根据段落管理器和配置文件生成格式化的文档
    
    Args:
        config: 格式配置字典或配置文件路径
        para_manager: 段落管理器实例，包含所有需要格式化的段落
        output_path: 输出文档路径
        
    Returns:
        str: 生成的文档路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)
    
    # 创建新文档
    doc = Document()
    
    # 设置页面格式
    set_paper_format(doc, config)
    
    # 图表编号计数器
    figure_number = 1
    table_number = 1
    
    # 遍历所有段落并应用相应格式
    for para_info in para_manager.paragraphs:
        # 获取段落类型
        para_type = para_info.type.value
        
        # 处理图片类型的段落
        if para_type == 'figures' and para_info.meta and 'image_path' in para_info.meta:
            image_path = para_info.meta['image_path']
            if os.path.exists(image_path):
                # 添加图片段落
                img_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['figures']:
                    apply_paragraph_format(img_paragraph, config['figures']['paragraph_format'])
                
                # 添加图片
                width = Inches(6)  # 默认宽度
                if 'width' in para_info.meta:
                    width_str = para_info.meta['width']
                    if isinstance(width_str, str) and 'inches' in width_str.lower():
                        width = Inches(float(width_str.lower().replace('inches', '')))
                    elif isinstance(width_str, (int, float)):
                        width = Inches(float(width_str))
                
                img_paragraph.add_run().add_picture(image_path, width=width)
                
                # 生成题注
                caption = para_info.content
                if not caption:
                    caption = get_openai_image_caption(image_path)
                
                # 设置图片编号
                fig_num = para_info.meta.get('figure_number', figure_number)
                
                # 获取题注位置设置
                position = "below"
                if 'caption' in config['figures'] and 'position' in config['figures']['caption']:
                    position = config['figures']['caption']['position']
                
                # 添加题注段落
                caption_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['figures']:
                    apply_paragraph_format(caption_paragraph, config['figures']['paragraph_format'])
                
                # 设置题注文本
                caption_text = f"图 {fig_num} {caption}"
                caption_run = caption_paragraph.add_run(caption_text)
                
                # 应用题注字体格式
                if 'caption' in config['figures'] and 'fonts' in config['figures']['caption']:
                    apply_font_settings(caption_run, config['figures']['caption']['fonts'])
                
                # 更新图片编号
                figure_number += 1
            continue
        
        # 处理表格类型的段落
        elif para_type == 'tables' and para_info.meta and 'table_data' in para_info.meta:
            table_data = para_info.meta['table_data']
            
            # 获取题注位置设置
            position = "above"
            if 'caption' in config['tables'] and 'position' in config['tables']['caption']:
                position = config['tables']['caption']['position']
            
            # 如果题注在表格上方，先添加题注
            if position.lower() == "above":
                # 添加题注段落
                caption_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['tables']:
                    apply_paragraph_format(caption_paragraph, config['tables']['paragraph_format'])
                
                # 设置题注文本
                caption = para_info.content
                tab_num = para_info.meta.get('table_number', table_number)
                caption_text = f"表 {tab_num} {caption}"
                caption_run = caption_paragraph.add_run(caption_text)
                
                # 应用题注字体格式
                if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                    apply_font_settings(caption_run, config['tables']['caption']['fonts'])
            
            # 添加表格
            if isinstance(table_data, list) and len(table_data) > 0:
                num_rows = len(table_data)
                num_cols = len(table_data[0]) if isinstance(table_data[0], list) else 1
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                
                # 应用表格样式
                if 'caption' in config['tables'] and 'border' in config['tables']['caption']:
                    border_config = config['tables']['caption']['border']
                    if 'style' in border_config and border_config['style'] == 'solid':
                        table.style = 'Table Grid'
                
                # 填充表格数据
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data if isinstance(row_data, list) else [row_data]):
                        cell = table.cell(i, j)
                        cell.text = str(cell_data)
                        
                        # 应用单元格格式
                        for paragraph in cell.paragraphs:
                            if 'paragraph_format' in config['tables']:
                                apply_paragraph_format(paragraph, config['tables']['paragraph_format'])
                            
                            for run in paragraph.runs:
                                if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                                    apply_font_settings(run, config['tables']['caption']['fonts'])
            
            # 如果题注在表格下方，在表格后添加题注
            if position.lower() == "below":
                # 添加题注段落
                caption_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['tables']:
                    apply_paragraph_format(caption_paragraph, config['tables']['paragraph_format'])
                
                # 设置题注文本
                caption = para_info.content
                tab_num = para_info.meta.get('table_number', table_number)
                caption_text = f"表 {tab_num} {caption}"
                caption_run = caption_paragraph.add_run(caption_text)
                
                # 应用题注字体格式
                if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                    apply_font_settings(caption_run, config['tables']['caption']['fonts'])
            
            # 更新表格编号
            table_number += 1
            continue
        
        # 创建普通段落
        paragraph = doc.add_paragraph()
        
        # 获取段落格式设置
        if para_type in config:
            format_config = config[para_type]
            
            # 应用段落格式
            if 'paragraph_format' in format_config:
                apply_paragraph_format(paragraph, format_config['paragraph_format'])
            
            # 处理文本内容并应用字体设置
            if 'fonts' in format_config:
                font_settings = format_config['fonts']
                
                # 拆分中英文
                segments = split_text_by_language(para_info.content)
                
                # 为每段添加相应的格式
                for text, is_chinese in segments:
                    run = paragraph.add_run(text)
                    apply_font_settings(run, font_settings, is_chinese)
            else:
                # 无特殊字体设置，使用默认体
                paragraph.add_run(para_info.content)
        else:
            # 使用默认体
            paragraph.add_run(para_info.content)
    
    # 保存文档
    doc.save(output_path)
    
    return output_path

# 使用示例
# config = load_config("config.json")
# para_manager = ParagraphManager()
# para_manager.add_para(ParsedParaType.TITLE_ZH, "测试文档标题")
# output_path = generate_formatted_doc(config, para_manager, "output.docx")
