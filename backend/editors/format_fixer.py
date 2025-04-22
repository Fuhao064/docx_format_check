from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
import json
import os
import re
from typing import Dict, List, Optional, Union, Tuple, Any
from preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo

# 全局映射字典
ALIGNMENT_MAP = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
}

LINE_SPACING_MAP = {
    "1.0": 1.0,
    "1.5": 1.5,
    "2.0": 2.0,
    "single": 1.0,
    "double": 2.0
}

class FormatFixer:
    """格式修复器，用于根据检查出的错误自动修复文档格式问题"""

    def __init__(self, doc_path: str = None):
        """
        初始化格式修复器

        Args:
            doc_path: 文档路径，如果提供则加载该文档
        """
        self.doc = None
        self.doc_path = doc_path
        if doc_path:
            self.load_document(doc_path)

        # 错误类型映射
        self.error_type_map = {
            "字号": self._fix_font_size,
            "行间距": self._fix_line_spacing,
            "对齐方式": self._fix_alignment,
            "加粗": self._fix_bold,
            "斜体": self._fix_italic,
            "首行缩进": self._fix_first_line_indent,
            "字体": self._fix_font_family
        }

    def load_document(self, doc_path: str) -> None:
        """
        加载文档

        Args:
            doc_path: 文档路径
        """
        self.doc = Document(doc_path)
        self.doc_path = doc_path

    def fix_errors(self, errors: List[Dict], para_manager: ParagraphManager, output_path: Optional[str] = None) -> str:
        """
        根据错误列表修复文档格式

        Args:
            errors: 错误列表，每个错误是一个字典
            para_manager: 段落管理器，包含段落信息
            output_path: 输出文档路径，如果为None则生成临时文件

        Returns:
            str: 修复后的文档路径
        """
        if not self.doc:
            raise ValueError("未加载文档，请先调用load_document方法")

        # 按段落分组错误
        errors_by_para = self._group_errors_by_paragraph(errors)

        # 遍历段落管理器中的段落
        for i, para_info in enumerate(para_manager.paragraphs):
            para_content = para_info.content
            para_type = para_info.type

            # 查找对应的段落错误
            para_errors = self._find_para_errors(para_content, errors_by_para)

            if not para_errors:
                continue

            # 查找文档中对应的段落
            doc_para = self._find_matching_paragraph(para_content)
            if not doc_para:
                continue

            # 修复段落错误
            self._fix_paragraph_errors(doc_para, para_errors, para_info)

        # 保存修复后的文档
        if output_path is None:
            output_path = f"fixed_{os.path.basename(self.doc_path)}"

        self.doc.save(output_path)
        return output_path

    def fix_by_requirements(self, requirements: Dict, para_manager: ParagraphManager, output_path: Optional[str] = None) -> str:
        """
        根据格式要求直接应用格式

        Args:
            requirements: 格式要求字典
            para_manager: 段落管理器，包含段落信息
            output_path: 输出文档路径，如果为None则生成临时文件

        Returns:
            str: 修复后的文档路径
        """
        if not self.doc:
            raise ValueError("未加载文档，请先调用load_document方法")

        # 遍历段落管理器中的段落
        for i, para_info in enumerate(para_manager.paragraphs):
            para_content = para_info.content
            para_type = para_info.type.value

            # 查找文档中对应的段落
            doc_para = self._find_matching_paragraph(para_content)
            if not doc_para:
                continue

            # 获取该类型段落的格式要求
            if para_type in requirements:
                para_requirements = requirements[para_type]
                self._apply_paragraph_requirements(doc_para, para_requirements, para_info)

        # 保存修复后的文档
        if output_path is None:
            output_path = f"formatted_{os.path.basename(self.doc_path)}"

        self.doc.save(output_path)
        return output_path

    def _group_errors_by_paragraph(self, errors: List[Dict]) -> Dict[str, List[Dict]]:
        """
        按段落分组错误

        Args:
            errors: 错误列表

        Returns:
            Dict[str, List[Dict]]: 按段落分组的错误字典
        """
        errors_by_para = {}

        for error in errors:
            location = error.get('location', '')
            if '.' in location:
                # 提取段落内容的前几个字作为标识
                para_identifier = location.split('.', 1)[1].strip()

                if para_identifier not in errors_by_para:
                    errors_by_para[para_identifier] = []

                errors_by_para[para_identifier].append(error)

        return errors_by_para

    def _find_para_errors(self, para_content: str, errors_by_para: Dict[str, List[Dict]]) -> List[Dict]:
        """
        查找段落对应的错误

        Args:
            para_content: 段落内容
            errors_by_para: 按段落分组的错误字典

        Returns:
            List[Dict]: 段落对应的错误列表
        """
        # 取段落内容的前30个字符作为匹配依据
        para_start = para_content[:30] if len(para_content) > 30 else para_content

        for para_id, errors in errors_by_para.items():
            if para_start.startswith(para_id) or para_id.startswith(para_start):
                return errors

        return []

    def _find_matching_paragraph(self, para_content: str) -> Optional[Any]:
        """
        在文档中查找匹配的段落

        Args:
            para_content: 段落内容

        Returns:
            Optional[Any]: 匹配的段落对象，如果未找到则返回None
        """
        # 取段落内容的前30个字符作为匹配依据
        para_start = para_content[:30] if len(para_content) > 30 else para_content

        for para in self.doc.paragraphs:
            if para.text.startswith(para_start) or para_start.startswith(para.text[:30]):
                return para

        return None

    def _fix_paragraph_errors(self, paragraph: Any, errors: List[Dict], para_info: ParaInfo) -> None:
        """
        修复段落错误

        Args:
            paragraph: 段落对象
            errors: 错误列表
            para_info: 段落信息
        """
        for error in errors:
            error_message = error.get('message', '')

            # 提取错误类型和要求值
            error_type, required_value, actual_value = self._parse_error_message(error_message)

            if error_type and required_value:
                # 调用对应的修复方法
                fix_method = self.error_type_map.get(error_type)
                if fix_method:
                    fix_method(paragraph, required_value, para_info)

    def _apply_paragraph_requirements(self, paragraph: Any, requirements: Dict, para_info: ParaInfo) -> None:
        """
        应用段落格式要求

        Args:
            paragraph: 段落对象
            requirements: 格式要求字典
            para_info: 段落信息
        """
        # 应用字体设置
        if 'fonts' in requirements:
            font_settings = requirements['fonts']
            self._apply_font_settings(paragraph, font_settings, para_info)

        # 应用段落格式
        if 'paragraph_format' in requirements:
            para_format = requirements['paragraph_format']
            self._apply_paragraph_format(paragraph, para_format, para_info)

    def _apply_font_settings(self, paragraph: Any, font_settings: Dict, para_info: ParaInfo) -> None:
        """
        应用字体设置

        Args:
            paragraph: 段落对象
            font_settings: 字体设置字典
            para_info: 段落信息
        """
        # 如果段落没有runs，则添加一个run
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)
            paragraph.text = ""

        # 应用字体设置到所有runs
        for run in paragraph.runs:
            # 设置字体族
            if 'zh_family' in font_settings:
                run.font.name = font_settings['zh_family']

            # 设置字体大小
            if 'size' in font_settings:
                size_text = font_settings['size']
                if isinstance(size_text, str) and 'pt' in size_text:
                    size = float(size_text.replace('pt', ''))
                    run.font.size = Pt(size)
                else:
                    run.font.size = Pt(float(size_text))

            # 设置加粗
            if 'bold' in font_settings:
                run.font.bold = font_settings['bold']

            # 设置斜体
            if 'italic' in font_settings:
                run.font.italic = font_settings['italic']

    def _apply_paragraph_format(self, paragraph: Any, para_format: Dict, para_info: ParaInfo) -> None:
        """
        应用段落格式

        Args:
            paragraph: 段落对象
            para_format: 段落格式字典
            para_info: 段落信息
        """
        # 设置对齐方式
        if 'alignment' in para_format:
            alignment = para_format['alignment'].lower()
            if alignment in ALIGNMENT_MAP:
                paragraph.paragraph_format.alignment = ALIGNMENT_MAP[alignment]

        # 设置行间距
        if 'line_spacing' in para_format:
            spacing = para_format['line_spacing']
            if isinstance(spacing, str) and 'pt' in spacing:
                # 固定值行间距
                pt_value = float(spacing.replace('pt', ''))
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(pt_value)
            elif spacing in LINE_SPACING_MAP:
                # 倍数行间距
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph.paragraph_format.line_spacing = LINE_SPACING_MAP[spacing]
            else:
                # 尝试直接设置具体数值
                try:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph.paragraph_format.line_spacing = float(spacing)
                except ValueError:
                    pass

        # 设置首行缩进
        if 'indentation' in para_format and 'first_line' in para_format['indentation']:
            first_line = para_format['indentation']['first_line']
            if isinstance(first_line, str) and 'cm' in first_line:
                paragraph.paragraph_format.first_line_indent = Cm(float(first_line.replace('cm', '')))
            else:
                paragraph.paragraph_format.first_line_indent = Cm(float(first_line))

    def _parse_error_message(self, error_message: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        """
        解析错误消息，提取错误类型和要求值

        Args:
            error_message: 错误消息

        Returns:
            Tuple[Optional[str], Optional[str], Optional[str]]: 错误类型、要求值和实际值
        """
        # 匹配格式：'错误类型' 不匹配: 要求 要求值, 实际 实际值
        pattern = r"'([^']+)'\s+不匹配:\s+要求\s+([^,]+),\s+实际\s+(.+)"
        match = re.match(pattern, error_message)

        if match:
            error_type = match.group(1)
            required_value = match.group(2)
            actual_value = match.group(3)
            return error_type, required_value, actual_value

        return None, None, None

    def _fix_font_size(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复字体大小

        Args:
            paragraph: 段落对象
            required_value: 要求的字体大小
            para_info: 段落信息
        """
        # 提取字体大小数值
        size_match = re.search(r'(\d+(\.\d+)?)\s*pt', required_value)
        if not size_match:
            return

        size = float(size_match.group(1))

        # 应用到所有runs
        for run in paragraph.runs:
            run.font.size = Pt(size)

    def _fix_line_spacing(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复行间距

        Args:
            paragraph: 段落对象
            required_value: 要求的行间距
            para_info: 段落信息
        """
        # 处理固定值行间距
        if '固定值' in required_value or 'Fixed value' in required_value:
            # 提取pt值
            pt_match = re.search(r'(\d+(\.\d+)?)\s*pt', required_value)
            if pt_match:
                pt_value = float(pt_match.group(1))
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(pt_value)
        else:
            # 处理倍数行间距
            for spacing_name, spacing_value in LINE_SPACING_MAP.items():
                if spacing_name in required_value:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph.paragraph_format.line_spacing = spacing_value
                    break

    def _fix_alignment(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复对齐方式

        Args:
            paragraph: 段落对象
            required_value: 要求的对齐方式
            para_info: 段落信息
        """
        if required_value in ALIGNMENT_MAP:
            paragraph.paragraph_format.alignment = ALIGNMENT_MAP[required_value]

    def _fix_bold(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复加粗样式

        Args:
            paragraph: 段落对象
            required_value: 要求的加粗设置
            para_info: 段落信息
        """
        is_bold = required_value.lower() == '是' or required_value.lower() == 'true'

        # 应用到所有runs
        for run in paragraph.runs:
            run.font.bold = is_bold

    def _fix_italic(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复斜体样式

        Args:
            paragraph: 段落对象
            required_value: 要求的斜体设置
            para_info: 段落信息
        """
        is_italic = required_value.lower() == '是' or required_value.lower() == 'true'

        # 应用到所有runs
        for run in paragraph.runs:
            run.font.italic = is_italic

    def _fix_first_line_indent(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复首行缩进

        Args:
            paragraph: 段落对象
            required_value: 要求的首行缩进
            para_info: 段落信息
        """
        # 提取缩进值
        indent_match = re.search(r'(\d+(\.\d+)?)\s*(cm|字符)', required_value)
        if not indent_match:
            return

        indent_value = float(indent_match.group(1))
        indent_unit = indent_match.group(3)

        if indent_unit == 'cm':
            paragraph.paragraph_format.first_line_indent = Cm(indent_value)
        elif indent_unit == '字符':
            # 假设一个字符宽度为0.5cm
            paragraph.paragraph_format.first_line_indent = Cm(indent_value * 0.5)

    def _fix_font_family(self, paragraph: Any, required_value: str, para_info: ParaInfo) -> None:
        """
        修复字体族

        Args:
            paragraph: 段落对象
            required_value: 要求的字体族
            para_info: 段落信息
        """
        # 应用到所有runs
        for run in paragraph.runs:
            run.font.name = required_value

# 批量修复文档中的格式错误
def batch_fix_errors(doc_path: str, errors: List[Dict], para_manager: ParagraphManager, output_path: Optional[str] = None) -> str:
    """
    批量修复文档中的格式错误

    Args:
        doc_path: 文档路径
        errors: 错误列表
        para_manager: 段落管理器
        output_path: 输出文档路径

    Returns:
        str: 修复后的文档路径
    """
    fixer = FormatFixer(doc_path)
    return fixer.fix_errors(errors, para_manager, output_path)

# 根据格式要求应用格式
def apply_format_requirements(doc_path: str, requirements: Dict, para_manager: ParagraphManager, output_path: Optional[str] = None) -> str:
    """
    根据格式要求应用格式

    Args:
        doc_path: 文档路径
        requirements: 格式要求字典
        para_manager: 段落管理器
        output_path: 输出文档路径

    Returns:
        str: 修复后的文档路径
    """
    fixer = FormatFixer(doc_path)
    return fixer.fix_by_requirements(requirements, para_manager, output_path)
