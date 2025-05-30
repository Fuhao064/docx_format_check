from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
import json
import os
import re
import requests
from typing import Dict, List, Optional, Union, Tuple
from backend.preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo
from docx.oxml.ns import qn  # 导入qn函数，用于XML命名空间

# 全局映射字典
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY
}

LINE_SPACING_MAP = {
    "1.0": 1.0,
    "1.5": 1.5,
    "2.0": 2.0,
    "single": 1.0,
    "double": 2.0
}

def load_config(config_path: str) -> Dict:
    """加载配置文件"""
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def apply_font_settings(run, font_settings: Dict, is_chinese: bool = True):
    """
    应用字体设置到文本运行对象

    Args:
        run: 文本运行对象
        font_settings: 字体设置字典
        is_chinese: 是否为中文
    """
    # 设置字体族
    if 'en_family' in font_settings:
        # 先设置英文字体（对所有文本设置默认字体）
        run.font.name = font_settings['en_family']

    # 然后使用XML操作设置中文字体
    if 'zh_family' in font_settings:
        try:
            # 确保元素存在
            if not hasattr(run, 'element') or not hasattr(run.element, 'rPr'):
                print(f"运行元素或rPr不存在，无法设置中文字体")
                return

            # 如果运行元素的rPr不存在，添加它
            if not run.element.rPr:
                run.element._add_rPr()

            # 设置东亚字体（中文字体）
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_settings['zh_family'])
            print(f"成功设置中文字体: {font_settings['zh_family']}")
        except Exception as e:
            print(f"设置中文字体时出错: {str(e)}")

    # 设置字体大小
    if 'size' in font_settings:
        size_text = font_settings['size']
        # 如果值是"Unknown"，使用默认值
        if size_text == "Unknown" or size_text == "unknown":
            # 使用默认字体大小（10.5pt）
            run.font.size = Pt(10.5)
            print(f"字体大小值为'Unknown'，使用默认值10.5pt")
        elif isinstance(size_text, str) and 'pt' in size_text:
            size = float(size_text.replace('pt', ''))
            run.font.size = Pt(size)
        else:
            try:
                run.font.size = Pt(float(size_text))
            except (ValueError, TypeError):
                # 如果无法转换为浮点数，使用默认值
                run.font.size = Pt(10.5)
                print(f"无法将字体大小值'{size_text}'转换为浮点数，使用默认值10.5pt")

    # 设置加粗
    if 'bold' in font_settings:
        run.font.bold = font_settings['bold']

    # 设置斜体
    if 'italic' in font_settings:
        run.font.italic = font_settings['italic']

    # 设置全部大写
    if 'isAllCaps' in font_settings:
        run.font.all_caps = font_settings['isAllCaps']

    # 设置字体颜色
    if 'color' in font_settings:
        color = font_settings['color'].lstrip('#')
        if len(color) == 6:
            r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)

def apply_paragraph_format(paragraph, format_settings: Dict):
    """
    应用段落格式设置

    Args:
        paragraph: 段落对象
        format_settings: 格式设置字典
    """
    # 设置行间距
    if 'line_spacing' in format_settings:
        spacing = format_settings['line_spacing']
        if isinstance(spacing, str) and ('固定值' in spacing or 'Fixed value' in spacing):
            # 处理固定值行间距
            pt_match = re.search(r'(\d+(\.\d+)?)\s*pt', spacing)
            if pt_match:
                pt_value = float(pt_match.group(1))
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(pt_value)
        elif spacing in LINE_SPACING_MAP:
            # 处理倍数行间距
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            if spacing == "1.5":
                paragraph.paragraph_format.line_spacing = 1.5
            elif spacing == "2.0" or spacing == "double":
                paragraph.paragraph_format.line_spacing = 2.0
            else:
                paragraph.paragraph_format.line_spacing = 1.0
        else:
            # 尝试直接设置具体数值
            try:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph.paragraph_format.line_spacing = float(spacing)
            except ValueError:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph.paragraph_format.line_spacing = 1.5  # 默认值

    # 设置对齐方式
    if 'alignment' in format_settings:
        alignment = format_settings['alignment'].lower()
        if alignment in ALIGNMENT_MAP:
            paragraph.paragraph_format.alignment = ALIGNMENT_MAP[alignment]

    # 设置缩进
    if 'indentation' in format_settings:
        indentation = format_settings['indentation']

        # 首行缩进
        if 'first_line' in indentation:
            first_line = indentation['first_line']
            # 如果值是"Unknown"，使用默认值
            if first_line == "Unknown" or first_line == "unknown":
                # 使用默认缩进值（0.5cm）
                paragraph.paragraph_format.first_line_indent = Cm(0.5)
                print(f"首行缩进值为'Unknown'，使用默认值0.5cm")
            # 处理字符单位的缩进
            elif isinstance(first_line, str) and 'character' in first_line:
                # 使用XML操作设置字符单位的缩进
                try:
                    # 提取字符数量
                    char_match = re.search(r'(\d+(\.\d+)?)', first_line)
                    if char_match:
                        char_value = int(float(char_match.group(1)) * 100)  # 转换为整数，并乘以100（Word中的字符单位是1/100字符）
                        # 设置首行缩进为0（清除原有的缩进设置）
                        paragraph.paragraph_format.first_line_indent = 0
                        # 使用XML操作设置字符单位的缩进
                        paragraph_format = paragraph.paragraph_format
                        if not paragraph_format._element.pPr:
                            paragraph_format._element._add_pPr()
                        if not paragraph_format._element.pPr.ind:
                            paragraph_format._element.pPr._add_ind()
                        paragraph_format._element.pPr.ind.set(qn("w:firstLineChars"), str(char_value))
                        print(f"设置首行缩进为{char_match.group(1)}字符")
                    else:
                        # 如果无法提取字符数量，使用默认值
                        paragraph.paragraph_format.first_line_indent = Cm(0.5)
                        print(f"无法提取字符数量，使用默认值0.5cm")
                except Exception as e:
                    # 如果设置失败，使用默认值
                    paragraph.paragraph_format.first_line_indent = Cm(0.5)
                    print(f"设置字符单位缩进失败: {str(e)}，使用默认值0.5cm")
            elif isinstance(first_line, str) and 'cm' in first_line:
                paragraph.paragraph_format.first_line_indent = Cm(float(first_line.replace('cm', '')))
            else:
                try:
                    paragraph.paragraph_format.first_line_indent = Cm(float(first_line))
                except (ValueError, TypeError):
                    # 如果无法转换为浮点数，尝试处理字符单位
                    try:
                        # 检查是否包含"字符"或"character"
                        if isinstance(first_line, str) and ("字符" in first_line or "character" in first_line.lower()):
                            # 提取字符数量
                            char_match = re.search(r'(\d+(\.\d+)?)', first_line)
                            if char_match:
                                char_value = int(float(char_match.group(1)) * 100)  # 转换为整数，并乘以100
                                # 设置首行缩进为0（清除原有的缩进设置）
                                paragraph.paragraph_format.first_line_indent = 0
                                # 使用XML操作设置字符单位的缩进
                                paragraph_format = paragraph.paragraph_format
                                if not paragraph_format._element.pPr:
                                    paragraph_format._element._add_pPr()
                                if not paragraph_format._element.pPr.ind:
                                    paragraph_format._element.pPr._add_ind()
                                paragraph_format._element.pPr.ind.set(qn("w:firstLineChars"), str(char_value))
                                print(f"设置首行缩进为{char_match.group(1)}字符")
                                return
                        # 如果不是字符单位或无法提取字符数量，使用默认值
                        paragraph.paragraph_format.first_line_indent = Cm(0.5)
                        print(f"无法将首行缩进值'{first_line}'转换为浮点数或字符单位，使用默认值0.5cm")
                    except Exception as e:
                        # 如果设置失败，使用默认值
                        paragraph.paragraph_format.first_line_indent = Cm(0.5)
                        print(f"设置首行缩进失败: {str(e)}，使用默认值0.5cm")

        # 左缩进
        if 'left' in indentation:
            left = indentation['left']
            # 如果值是"Unknown"，使用默认值
            if left == "Unknown" or left == "unknown":
                # 使用默认缩进值（0cm）
                paragraph.paragraph_format.left_indent = Cm(0)
                print(f"左缩进值为'Unknown'，使用默认值0cm")
            elif isinstance(left, str) and 'cm' in left:
                paragraph.paragraph_format.left_indent = Cm(float(left.replace('cm', '')))
            else:
                try:
                    paragraph.paragraph_format.left_indent = Cm(float(left))
                except (ValueError, TypeError):
                    # 如果无法转换为浮点数，使用默认值
                    paragraph.paragraph_format.left_indent = Cm(0)
                    print(f"无法将左缩进值'{left}'转换为浮点数，使用默认值0cm")

        # 右缩进
        if 'right' in indentation:
            right = indentation['right']
            # 如果值是"Unknown"，使用默认值
            if right == "Unknown" or right == "unknown":
                # 使用默认缩进值（0cm）
                paragraph.paragraph_format.right_indent = Cm(0)
                print(f"右缩进值为'Unknown'，使用默认值0cm")
            elif isinstance(right, str) and 'cm' in right:
                paragraph.paragraph_format.right_indent = Cm(float(right.replace('cm', '')))
            else:
                try:
                    paragraph.paragraph_format.right_indent = Cm(float(right))
                except (ValueError, TypeError):
                    # 如果无法转换为浮点数，使用默认值
                    paragraph.paragraph_format.right_indent = Cm(0)
                    print(f"无法将右缩进值'{right}'转换为浮点数，使用默认值0cm")

        # 段前距
        if 'space_before' in indentation:
            space_before = indentation['space_before']
            # 如果值是"Unknown"，使用默认值
            if space_before == "Unknown" or space_before == "unknown":
                # 使用默认段前距值（0cm）
                paragraph.paragraph_format.space_before = Cm(0)
                print(f"段前距值为'Unknown'，使用默认值0cm")
            elif isinstance(space_before, str) and 'cm' in space_before:
                paragraph.paragraph_format.space_before = Cm(float(space_before.replace('cm', '')))
            else:
                try:
                    paragraph.paragraph_format.space_before = Cm(float(space_before))
                except (ValueError, TypeError):
                    # 如果无法转换为浮点数，使用默认值
                    paragraph.paragraph_format.space_before = Cm(0)
                    print(f"无法将段前距值'{space_before}'转换为浮点数，使用默认值0cm")

        # 段后距
        if 'space_after' in indentation:
            space_after = indentation['space_after']
            # 如果值是"Unknown"，使用默认值
            if space_after == "Unknown" or space_after == "unknown":
                # 使用默认段后距值（0cm）
                paragraph.paragraph_format.space_after = Cm(0)
                print(f"段后距值为'Unknown'，使用默认值0cm")
            elif isinstance(space_after, str) and 'cm' in space_after:
                paragraph.paragraph_format.space_after = Cm(float(space_after.replace('cm', '')))
            else:
                try:
                    paragraph.paragraph_format.space_after = Cm(float(space_after))
                except (ValueError, TypeError):
                    # 如果无法转换为浮点数，使用默认值
                    paragraph.paragraph_format.space_after = Cm(0)
                    print(f"无法将段后距值'{space_after}'转换为浮点数，使用默认值0cm")

def is_chinese_char(char: str) -> bool:
    """判断字符是否为中文"""
    return '\u4e00' <= char <= '\u9fff'

def split_text_by_language(text: str) -> List[Tuple[str, bool]]:
    """
    将文本按中英文分段

    Args:
        text: 要分段的文本

    Returns:
        List[Tuple[str, bool]]: 文本段列表，每项是(文本段, 是否中文)
    """
    result = []
    current_text = ""
    current_is_chinese = None

    for char in text:
        is_chinese = is_chinese_char(char)

        # 首个字符或当前字符类型与前面相同
        if current_is_chinese is None or is_chinese == current_is_chinese:
            current_text += char
            current_is_chinese = is_chinese
        else:
            # 字符类型发生变化，保存前一段并开始新段
            if current_text:
                result.append((current_text, current_is_chinese))
            current_text = char
            current_is_chinese = is_chinese

    # 添加最后一段
    if current_text:
        result.append((current_text, current_is_chinese))

    return result

def get_openai_image_caption(image_path: str) -> str:
    """
    使用OpenAI API获取图片题注

    Args:
        image_path: 图片路径

    Returns:
        str: 生成的题注
    """
    try:
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            return "图片题注（需要设置OpenAI API密钥）"

        # 读取图片并转为base64
        import base64
        with open(image_path, "rb") as image_file:
            encoded_image = base64.b64encode(image_file.read()).decode('utf-8')

        # 调用OpenAI API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }

        payload = {
            "model": "gpt-4-vision-preview",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "请为这张图片生成一个简短的中文学术题注，不超过20字："
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded_image}"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 100
        }

        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload
        )

        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            return f"图 {os.path.basename(image_path)}"

    except Exception:
        return f"图 {os.path.basename(image_path)}"

def set_paper_format(doc: Document, config: Dict):
    """
    设置文档页面格式

    Args:
        doc: 文档对象
        config: 配置字典
    """
    paper_config = config.get('paper', {})

    # 获取文档的节
    section = doc.sections[0]

    # 设置纸张大小
    if 'size' in paper_config:
        size = paper_config['size']
        # 标准纸张大小设置
        if size == 'A4':
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        elif size == 'Letter':
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        # 可以添加更多纸张尺寸...

    # 设置页面方向
    if 'orientation' in paper_config:
        orientation = paper_config['orientation']
        if orientation.lower() == 'portrait':
            section.orientation = WD_ORIENT.PORTRAIT
        elif orientation.lower() == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE

    # 设置页边距
    if 'margins' in paper_config:
        margins = paper_config['margins']

        for margin_name, margin_value in margins.items():
            # 如果值是"Unknown"，使用默认值
            if margin_value == "Unknown" or margin_value == "unknown":
                # 使用默认边距值（2.54cm = 1英寸）
                margin_cm = 2.54
                print(f"边距值为'Unknown'，使用默认值{margin_cm}cm")
            elif isinstance(margin_value, str) and 'cm' in margin_value:
                margin_cm = float(margin_value.replace('cm', ''))
            else:
                try:
                    margin_cm = float(margin_value)
                except (ValueError, TypeError):
                    # 如果无法转换为浮点数，使用默认值
                    margin_cm = 2.54
                    print(f"无法将边距值'{margin_value}'转换为浮点数，使用默认值{margin_cm}cm")

            if margin_name == 'top':
                section.top_margin = Cm(margin_cm)
            elif margin_name == 'bottom':
                section.bottom_margin = Cm(margin_cm)
            elif margin_name == 'left':
                section.left_margin = Cm(margin_cm)
            elif margin_name == 'right':
                section.right_margin = Cm(margin_cm)

def format_document(config: Dict, para_manager: ParagraphManager, output_path: str = None, doc_path: str = None) -> str:
    """
    根据配置和段落管理器格式化文档

    Args:
        config: 格式配置字典或配置文件路径
        para_manager: 段落管理器实例
        output_path: 输出文档路径，如果为None则生成临时文件
        doc_path: 可选的输入文档路径，如果提供则在该文档基础上修改

    Returns:
        str: 输出文档的路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)

    # 创建或打开文档
    if doc_path:
        doc = Document(doc_path)
        # 清空文档内容
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
    else:
        doc = Document()

    # 设置页面格式
    set_paper_format(doc, config)

    # 按照段落类型和配置格式化文档
    for para_info in para_manager.paragraphs:
        # 创建新段落
        paragraph = doc.add_paragraph()

        # 获取段落类型
        para_type = para_info.type.value

        # 获取段落格式设置
        if para_type in config:
            format_config = config[para_type]

            # 应用段落格式
            if 'paragraph_format' in format_config:
                apply_paragraph_format(paragraph, format_config['paragraph_format'])

            # 处理文本内容并应用字体设置
            if 'fonts' in format_config:
                font_settings = format_config['fonts']

                # 拆分中英文
                segments = split_text_by_language(para_info.content)

                # 为每段添加相应的格式
                for text, is_chinese in segments:
                    run = paragraph.add_run(text)
                    apply_font_settings(run, font_settings, is_chinese)
            else:
                # 无特殊字体设置，使用默认体
                paragraph.add_run(para_info.content)
        else:
            # 使用默认体
            paragraph.add_run(para_info.content)

        # 特殊处理图片段落
        if para_type == 'figures' and para_info.meta and para_info.meta.get('image_path'):
            image_path = para_info.meta['image_path']

            # 添加图片
            if os.path.exists(image_path):
                # 创建图片段落（不是题注段落）
                img_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['figures']:
                    apply_paragraph_format(img_paragraph, config['figures']['paragraph_format'])

                # 添加图片，设置合适的宽度（可根据需要调整）
                img_paragraph.add_run().add_picture(image_path, width=Inches(6))

                # 生成题注
                caption = para_info.content
                if not caption:
                    caption = get_openai_image_caption(image_path)

                # 添加题注段落
                caption_paragraph = doc.add_paragraph()
                if 'paragraph_format' in config['figures']:
                    apply_paragraph_format(caption_paragraph, config['figures']['paragraph_format'])

                # 设置题注文本
                caption_run = caption_paragraph.add_run(f"图 {para_info.meta.get('figure_number', '')} {caption}")

                # 应用题注字体格式
                if 'caption' in config['figures'] and 'fonts' in config['figures']['caption']:
                    apply_font_settings(caption_run, config['figures']['caption']['fonts'])

    # 确定输出路径
    if output_path is None:
        output_path = f"output_{os.path.basename(doc_path)}" if doc_path else "formatted_document.docx"

    # 保存文档
    doc.save(output_path)

    return output_path

def add_figure_caption(config: Dict, doc_path: str, image_path: str, figure_number: int, caption: Optional[str] = None,
                      output_path: Optional[str] = None) -> str:
    """
    为文档中的图片添加题注

    Args:
        config: 格式配置字典或配置文件路径
        doc_path: 文档路径
        image_path: 图片路径
        figure_number: 图片编号
        caption: 题注文本，如果为None则自动生成
        output_path: 输出文档路径，如果为None则覆盖原文档

    Returns:
        str: 输出文档的路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)

    # 打开文档
    doc = Document(doc_path)

    # 自动生成题注
    if caption is None:
        caption = get_openai_image_caption(image_path)

    # 添加图片
    paragraph = doc.add_paragraph()
    if 'paragraph_format' in config['figures']:
        apply_paragraph_format(paragraph, config['figures']['paragraph_format'])

    # 添加图片，设置合适的宽度
    paragraph.add_run().add_picture(image_path, width=Inches(6))

    # 获取题注位置设置
    position = "below"
    if 'caption' in config['figures'] and 'position' in config['figures']['caption']:
        position = config['figures']['caption']['position']

    # 如果题注在上方，则在图片前添加题注
    if position.lower() == "above":
        # 在当前段落之前插入新段落
        p_index = len(doc.paragraphs) - 1
        for i in range(p_index, 0, -1):
            doc.paragraphs[i]._p.addprevious(doc.paragraphs[i-1]._p)
    else:
        # 添加下方题注
        caption_paragraph = doc.add_paragraph()
        if 'paragraph_format' in config['figures']:
            apply_paragraph_format(caption_paragraph, config['figures']['paragraph_format'])

        caption_text = f"图 {figure_number} {caption}"
        caption_run = caption_paragraph.add_run(caption_text)

        # 应用题注字体格式
        if 'caption' in config['figures'] and 'fonts' in config['figures']['caption']:
            apply_font_settings(caption_run, config['figures']['caption']['fonts'])

    # 保存文档
    output_file = output_path or doc_path
    doc.save(output_file)

    return output_file

def format_table_caption(config: Dict, doc_path: str, table_number: int, caption: str, output_path: Optional[str] = None) -> str:
    """
    为文档中的表格添加题注

    Args:
        config: 格式配置字典或配置文件路径
        doc_path: 文档路径
        table_number: 表格编号
        caption: 题注文本
        output_path: 输出文档路径，如果为None则覆盖原文档

    Returns:
        str: 输出文档的路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)

    # 打开文档
    doc = Document(doc_path)

    # 获取题注位置设置
    position = "above"
    if 'caption' in config['tables'] and 'position' in config['tables']['caption']:
        position = config['tables']['caption']['position']

    # 添加题注段落
    caption_paragraph = doc.add_paragraph()
    if 'paragraph_format' in config['tables']:
        apply_paragraph_format(caption_paragraph, config['tables']['paragraph_format'])

    # 设置题注文本
    caption_text = f"表 {table_number} {caption}"
    caption_run = caption_paragraph.add_run(caption_text)

    # 应用题注字体格式
    if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
        apply_font_settings(caption_run, config['tables']['caption']['fonts'])

    # 保存文档
    output_file = output_path or doc_path
    doc.save(output_file)

    return output_file

def generate_formatted_doc(config: Dict, para_manager: ParagraphManager, output_path: str, errors: Optional[List[Dict]] = None, doc_path: Optional[str] = None) -> str:
    """
    根据段落管理器和配置文件生成格式化的文档，并根据错误信息进行修改
    保留原文档中所有元素的相对位置

    Args:
        config: 格式配置字典或配置文件路径
        para_manager: 段落管理器实例，包含所有需要格式化的段落
        output_path: 输出文档路径
        errors: 错误信息列表（可选）
        doc_path: 原文档路径（可选），如果提供则直接在原文档上修改文本内容

    Returns:
        str: 生成的文档路径
    """
    # 如果config是字符串路径，则加载配置
    if isinstance(config, str):
        config = load_config(config)

    # 判断是否有原文档
    if doc_path and os.path.exists(doc_path):
        # 直接打开原文档进行修改，而不是创建新文档
        doc = Document(doc_path)
        print(f"正在基于原文档 {doc_path} 进行格式化")

        # 提取原文档中的所有段落和元素
        original_paragraphs = list(doc.paragraphs)
        original_tables = list(doc.tables)

        # 为段落管理器中的段落创建一个索引
        para_index = 0
        table_index = 0

        # 逐个段落处理，修改格式但保持内容和相对位置
        for i, para in enumerate(original_paragraphs):
            # 跳过空段落
            if not para.text.strip():
                continue

            # 获取段落信息
            if para_index < len(para_manager.paragraphs):
                para_info = para_manager.paragraphs[para_index]
                para_type = para_info.type.value

                # 判断是否为表格标题段落
                is_table_caption = False
                if i < len(original_paragraphs) - 1:
                    # 检查下一个元素是否为表格
                    next_para = original_paragraphs[i+1]
                    if not next_para.text.strip() and table_index < len(original_tables):
                        # 可能是表格标题
                        if para.text.strip().startswith("表"):
                            is_table_caption = True

                # 判断是否为图片
                is_figure = False
                for run in para.runs:
                    if hasattr(run, "_r") and run._r.xpath(".//a:blip"):
                        is_figure = True
                        break

                # 判断是否为图片标题
                is_figure_caption = False
                if para.text.strip().startswith("图") and not is_figure:
                    is_figure_caption = True

                # 判断是否为公式
                is_equation = False
                if para.text.strip().startswith("式") or para.text.strip().startswith("公式"):
                    is_equation = True

                # 应用格式设置
                # 根据段落类型或特殊类型应用格式
                if para_type in config:
                    format_config = config[para_type]

                    # 应用段落格式
                    if 'paragraph_format' in format_config:
                        apply_paragraph_format(para, format_config['paragraph_format'])

                    # 更新段落文本内容
                    if not is_figure and not is_table_caption and not is_figure_caption and not is_equation:
                        # 清空段落内容准备重新设置
                        for run in list(para.runs):
                            p = run._element
                            p.getparent().remove(p)

                        # 处理文本内容并应用字体设置
                        if 'fonts' in format_config:
                            font_settings = format_config['fonts']

                            # 拆分中英文
                            segments = split_text_by_language(para_info.content)

                            # 为每段添加相应的格式
                            for text, is_chinese in segments:
                                run = para.add_run(text)
                                apply_font_settings(run, font_settings, is_chinese)
                        else:
                            # 无特殊字体设置，使用默认体
                            para.add_run(para_info.content)

                            # 如果是参考文献类型，根据配置设置字体
                            if para_type == 'references' or para_type == 'references_content':
                                # 尝试从配置中获取参考文献的字体设置
                                ref_config = config.get('references', {}) or config.get('references_content', {})
                                ref_fonts = ref_config.get('fonts', {})

                                if ref_fonts:
                                    # 如果配置中有字体设置，使用配置中的设置
                                    for run in para.runs:
                                        apply_font_settings(run, ref_fonts)
                                else:
                                    # 如果配置中没有字体设置，使用默认值
                                    for run in para.runs:
                                        # 设置英文字体
                                        run.font.name = 'Times New Roman'
                                        # 设置中文字体
                                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                        # 设置字体大小（如果没有特别指定）
                                        if not hasattr(run.font, 'size') or run.font.size is None:
                                            run.font.size = Pt(10.5)

                if errors:
                    # 标记错误段落
                    for error in errors:
                        if error['location'] and error['location'][:20] in para_info.content[:20]:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0) # 红色标记错误
                            break # 找到第一个匹配的错误就跳出循环

                # 增加索引，处理下一个段落
                if not is_figure and not is_table_caption and not is_figure_caption and not is_equation:
                    para_index += 1

        # 处理表格
        for table in original_tables:
            # 应用表格样式
            if 'tables' in config and 'caption' in config['tables'] and 'border' in config['tables']['caption']:
                border_config = config['tables']['caption']['border']
                if 'style' in border_config and border_config['style'] == 'solid':
                    table.style = 'Table Grid'

            # 应用单元格格式
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'tables' in config and 'paragraph_format' in config['tables']:
                            apply_paragraph_format(paragraph, config['tables']['paragraph_format'])

                        for run in paragraph.runs:
                            if 'tables' in config and 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                                apply_font_settings(run, config['tables']['caption']['fonts'])
    else:
        # 如果没有原文档，则创建新文档（原有逻辑）
        doc = Document()

        # 设置页面格式
        set_paper_format(doc, config)

        # 判断是否使用原文档中的图片和表格
        use_original_media = doc_path and os.path.exists(doc_path)

        # 如果使用原文档中的图片和表格，则提取原文档中的图片和表格
        original_images = []
        original_tables = []

        if use_original_media:
            try:
                # 打开原文档
                original_doc = Document(doc_path)

                # 提取原文档中的图片
                try:
                    import zipfile
                    import tempfile
                    import base64
                    import io
                    from PIL import Image

                    # 创建临时目录来存储提取的图片
                    temp_dir = tempfile.mkdtemp()

                    # 打开docx文件（实际上是一个zip文件）
                    with zipfile.ZipFile(doc_path, 'r') as docx_zip:
                        # 提取所有图片文件
                        for item in docx_zip.namelist():
                            if item.startswith('word/media/'):
                                # 提取图片数据
                                image_data = docx_zip.read(item)
                                image_name = os.path.basename(item)

                                # 保存图片到临时目录
                                image_path = os.path.join(temp_dir, image_name)
                                with open(image_path, 'wb') as f:
                                    f.write(image_data)

                                # 获取图片尺寸
                                try:
                                    with Image.open(io.BytesIO(image_data)) as img:
                                        width, height = img.size
                                        # 转换为英寸
                                        width_inches = width / 96  # 假设96dpi
                                except Exception as e:
                                    print(f"无法获取图片尺寸: {str(e)}")
                                    width_inches = 6  # 默认宽度
                                    height = 0

                                # 将图片信息添加到列表中
                                original_images.append({
                                    'path': image_path,
                                    'width': width_inches,
                                    'binary_data': image_data
                                })

                    print(f"从原文档中提取了 {len(original_images)} 个图片")
                except Exception as e:
                    print(f"提取原文档中的图片失败: {str(e)}")

                # 提取原文档中的表格
                try:
                    for i, table in enumerate(original_doc.tables):
                        # 提取表格数据
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text)
                            table_data.append(row_data)

                        # 将表格信息添加到列表中
                        original_tables.append({
                            'data': table_data,
                            'style': table.style.name if table.style else 'Table Grid',
                            'position': i
                        })

                    print(f"从原文档中提取了 {len(original_tables)} 个表格")
                except Exception as e:
                    print(f"提取原文档中的表格失败: {str(e)}")
            except Exception as e:
                print(f"打开原文档失败: {str(e)}")
                use_original_media = False

        # 图表编号计数器
        figure_number = 1
        table_number = 1

        # 如果有原文档中的图片和表格，则将它们添加到原始图片和表格列表中
        if use_original_media and original_images:
            print(f"从原文档中提取了 {len(original_images)} 个图片")
            # 将原文档中的图片信息保存到变量中，供后续使用

        if use_original_media and original_tables:
            print(f"从原文档中提取了 {len(original_tables)} 个表格")
            # 将原文档中的表格信息保存到变量中，供后续使用

        # 清空段落管理器中的段落，以便重新添加
        if use_original_media:
            # 保存非图片和表格的段落
            non_media_paragraphs = []
            for para in para_manager.paragraphs:
                if para.type.value != 'figures' and para.type.value != 'tables':
                    non_media_paragraphs.append(para)

            # 清空段落管理器
            para_manager.paragraphs = []

            # 添加非图片和表格的段落
            for para in non_media_paragraphs:
                para_manager.paragraphs.append(para)

            # 添加原文档中的图片
            print(f"将原文档中的 {len(original_images)} 个图片添加到段落管理器中")
            for i, image in enumerate(original_images):
                # 创建图片段落的元数据
                meta_data = {
                    'image_path': image['path'],
                    'width': image['width'],
                    'binary_data': image['binary_data'],
                    'figure_number': i + 1
                }

                # 将图片添加到段落管理器中
                para_manager.add_para(
                    para_type=ParsedParaType.FIGURES,
                    content=f"图片 {i+1}",  # 默认题注
                    meta=meta_data
                )

            # 添加原文档中的表格
            print(f"将原文档中的 {len(original_tables)} 个表格添加到段落管理器中")
            for i, table in enumerate(original_tables):
                # 创建表格段落的元数据
                meta_data = {
                    'table_data': table['data'],
                    'table_number': i + 1,
                    'style': table['style'],
                    'position': table['position']
                }

                # 将表格添加到段落管理器中
                para_manager.add_para(
                    para_type=ParsedParaType.TABLES,
                    content=f"表格 {i+1}",  # 默认题注
                    meta=meta_data
                )

        # 遍历所有段落并应用相应格式
        for para_info in para_manager.paragraphs:
            # 获取段落类型
            para_type = para_info.type.value

            # 如果是图片或表格段落，则正常处理
            # 我们已经将原文档中的图片和表格添加到段落管理器中，所以不需要跳过

            # 处理图片类型的段落
            if para_type == 'figures' and para_info.meta and 'image_path' in para_info.meta:
                image_path = para_info.meta['image_path']

                # 正常处理图片
                # 我们已经将原文档中的图片添加到段落管理器中

                # 如果图片路径不存在，尝试使用二进制数据
                if not os.path.exists(image_path):
                    print(f"图片路径 {image_path} 不存在，尝试使用二进制数据")

                    # 尝试使用二进制数据
                    if 'binary_data' in para_info.meta and para_info.meta['binary_data']:
                        try:
                            # 创建临时文件来保存图片
                            import tempfile
                            import base64
                            temp_dir = tempfile.mkdtemp()
                            temp_path = os.path.join(temp_dir, f"image_{figure_number}.png")

                            # 解码二进制数据并保存为图片
                            binary_data = para_info.meta['binary_data']
                            if isinstance(binary_data, str):
                                binary_data = base64.b64decode(binary_data)

                            with open(temp_path, 'wb') as f:
                                f.write(binary_data)

                            image_path = temp_path
                            print(f"使用二进制数据创建图片: {image_path}")
                        except Exception as e:
                            print(f"使用二进制数据创建图片失败: {str(e)}")

                if os.path.exists(image_path):
                    # 添加图片段落
                    img_paragraph = doc.add_paragraph()
                    if 'paragraph_format' in config['figures']:
                        apply_paragraph_format(img_paragraph, config['figures']['paragraph_format'])

                    # 添加图片
                    width = Inches(6)  # 默认宽度
                    if 'width' in para_info.meta:
                        width_str = para_info.meta['width']
                        if isinstance(width_str, str) and 'inches' in width_str.lower():
                            width = Inches(float(width_str.lower().replace('inches', '')))
                        elif isinstance(width_str, (int, float)):
                            width = Inches(float(width_str))
                        print(f"使用段落管理器中的图片宽度: {width_str}")

                    img_paragraph.add_run().add_picture(image_path, width=width)

                    # 生成题注
                    caption = para_info.content
                    if not caption:
                        caption = get_openai_image_caption(image_path)

                    # 设置图片编号
                    fig_num = para_info.meta.get('figure_number', figure_number)

                    # 获取题注位置设置
                    position = "below"
                    if 'caption' in config['figures'] and 'position' in config['figures']['caption']:
                        position = config['figures']['caption']['position']

                    # 添加题注段落
                    caption_paragraph = doc.add_paragraph()
                    if 'paragraph_format' in config['figures']:
                        apply_paragraph_format(caption_paragraph, config['figures']['paragraph_format'])

                    # 设置题注文本
                    caption_text = f"图 {fig_num} {caption}"
                    caption_run = caption_paragraph.add_run(caption_text)

                    # 应用题注字体格式
                    if 'caption' in config['figures'] and 'fonts' in config['figures']['caption']:
                        apply_font_settings(caption_run, config['figures']['caption']['fonts'])

                    # 更新图片编号
                    figure_number += 1
                continue

            # 处理表格类型的段落
            elif para_type == 'tables' and para_info.meta and 'table_data' in para_info.meta:
                table_data = para_info.meta['table_data']

                # 获取题注位置设置
                position = "above"
                if 'caption' in config['tables'] and 'position' in config['tables']['caption']:
                    position = config['tables']['caption']['position']

                # 如果题注在表格上方，先添加题注
                if position.lower() == "above":
                    # 添加题注段落
                    caption_paragraph = doc.add_paragraph()
                    if 'paragraph_format' in config['tables']:
                        apply_paragraph_format(caption_paragraph, config['tables']['paragraph_format'])

                    # 设置题注文本
                    caption = para_info.content
                    tab_num = para_info.meta.get('table_number', table_number)
                    caption_text = f"表 {tab_num} {caption}"
                    caption_run = caption_paragraph.add_run(caption_text)

                    # 应用题注字体格式
                    if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                        apply_font_settings(caption_run, config['tables']['caption']['fonts'])

                # 正常处理表格
                # 我们已经将原文档中的表格添加到段落管理器中
                # 创建新表格
                if isinstance(table_data, list) and len(table_data) > 0:
                    num_rows = len(table_data)
                    num_cols = len(table_data[0]) if isinstance(table_data[0], list) else 1

                    table = doc.add_table(rows=num_rows, cols=num_cols)

                    # 应用表格样式
                    if 'caption' in config['tables'] and 'border' in config['tables']['caption']:
                        border_config = config['tables']['caption']['border']
                        if 'style' in border_config and border_config['style'] == 'solid':
                            table.style = 'Table Grid'

                    # 填充表格数据
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data if isinstance(row_data, list) else [row_data]):
                            cell = table.cell(i, j)
                            cell.text = str(cell_data)

                            # 应用单元格格式
                            for paragraph in cell.paragraphs:
                                if 'paragraph_format' in config['tables']:
                                    apply_paragraph_format(paragraph, config['tables']['paragraph_format'])

                                for run in paragraph.runs:
                                    if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                                        apply_font_settings(run, config['tables']['caption']['fonts'])

                # 如果题注在表格下方，在表格后添加题注
                if position.lower() == "below":
                    # 添加题注段落
                    caption_paragraph = doc.add_paragraph()
                    if 'paragraph_format' in config['tables']:
                        apply_paragraph_format(caption_paragraph, config['tables']['paragraph_format'])

                    # 设置题注文本
                    caption = para_info.content
                    tab_num = para_info.meta.get('table_number', table_number)
                    caption_text = f"表 {tab_num} {caption}"
                    caption_run = caption_paragraph.add_run(caption_text)

                    # 应用题注字体格式
                    if 'caption' in config['tables'] and 'fonts' in config['tables']['caption']:
                        apply_font_settings(caption_run, config['tables']['caption']['fonts'])

                # 更新表格编号
                table_number += 1
                continue

            # 创建普通段落
            paragraph = doc.add_paragraph()

            # 获取段落格式设置
            if para_type in config:
                format_config = config[para_type]

                # 应用段落格式
                if 'paragraph_format' in format_config:
                    apply_paragraph_format(paragraph, format_config['paragraph_format'])

                # 处理文本内容并应用字体设置
                if 'fonts' in format_config:
                    font_settings = format_config['fonts']

                    # 拆分中英文
                    segments = split_text_by_language(para_info.content)

                    # 为每段添加相应的格式
                    for text, is_chinese in segments:
                        run = paragraph.add_run(text)
                        apply_font_settings(run, font_settings, is_chinese)
                else:
                    # 无特殊字体设置，使用默认体
                    run = paragraph.add_run(para_info.content)

                    # 如果是参考文献类型，根据配置设置字体
                    if para_type == 'references' or para_type == 'references_content':
                        # 尝试从配置中获取参考文献的字体设置
                        ref_config = config.get('references', {}) or config.get('references_content', {})
                        ref_fonts = ref_config.get('fonts', {})

                        if ref_fonts:
                            # 如果配置中有字体设置，使用配置中的设置
                            apply_font_settings(run, ref_fonts)
                        else:
                            # 如果配置中没有字体设置，使用默认值
                            # 设置英文字体
                            run.font.name = 'Times New Roman'
                            # 设置中文字体
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            # 设置字体大小（如果没有特别指定）
                            if not hasattr(run.font, 'size') or run.font.size is None:
                                run.font.size = Pt(10.5)

                if errors:
                    for error in errors:
                        if error['location'] and error['location'][:20] in para_info.content[:20]:
                            run.font.color.rgb = RGBColor(255, 0, 0) # 红色标记错误
                            break # 找到第一个匹配的错误就跳出循环
            else:
                # 使用默认体
                run = paragraph.add_run(para_info.content)

                # 如果是参考文献类型，根据配置设置字体
                if para_type == 'references' or para_type == 'references_content':
                    # 尝试从配置中获取参考文献的字体设置
                    ref_config = config.get('references', {}) or config.get('references_content', {})
                    ref_fonts = ref_config.get('fonts', {})

                    if ref_fonts:
                        # 如果配置中有字体设置，使用配置中的设置
                        apply_font_settings(run, ref_fonts)
                    else:
                        # 如果配置中没有字体设置，使用默认值
                        # 设置英文字体
                        run.font.name = 'Times New Roman'
                        # 设置中文字体
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        # 设置字体大小（如果没有特别指定）
                        if not hasattr(run.font, 'size') or run.font.size is None:
                            run.font.size = Pt(10.5)

    # 保存文档
    doc.save(output_path)

    # 返回文档路径
    return output_path

# 使用示例
# config = load_config("config.json")
# para_manager = ParagraphManager()
# para_manager.add_para(ParsedParaType.TITLE_ZH, "测试文档标题")
# output_path = generate_formatted_doc(config, para_manager, "output.docx")
