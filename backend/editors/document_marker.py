from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Union, Tuple, Any
from backend.preparation.para_type import ParagraphManager, ParsedParaType, ParaInfo
from docx.oxml.ns import qn  # 导入qn函数，用于XML命名空间

def mark_document_errors(doc_path: str, errors: List[Dict], para_manager: Optional[ParagraphManager] = None, output_path: Optional[str] = None) -> str:
    """
    标记文档中的格式错误，将错误内容标记为红色，并在内容后面添加错误说明

    Args:
        doc_path: 文档路径
        errors: 错误列表
        para_manager: 段落管理器（可选）
        output_path: 输出文档路径

    Returns:
        str: 标记后的文档路径
    """
    # 加载文档
    doc = Document(doc_path)

    # 检查错误列表是否为空
    if not errors or len(errors) == 0:
        # 添加标题 - 使用更通用的方式
        title_para = doc.add_paragraph('文档格式分析报告')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)
        # 设置字体
        title_run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加基本信息
        doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n文档: {os.path.basename(doc_path)}')

        # 添加无错误提示
        doc.add_paragraph('没有发现格式问题。').bold = True

        # 保存文档
        if output_path is None:
            output_path = os.path.join(os.path.dirname(doc_path), f"marked_{os.path.basename(doc_path)}")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

    # 添加标题 - 使用更通用的方式
    title_para = doc.add_paragraph('文档格式分析报告')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    # 设置字体
    title_run.font.name = 'Times New Roman'  # 先设置英文字体
    # 设置中文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 添加基本信息
    doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n文档: {os.path.basename(doc_path)}')

    # 添加错误摘要 - 使用更通用的方式
    subtitle_para = doc.add_paragraph(f'发现 {len(errors)} 个格式问题:')
    subtitle_run = subtitle_para.runs[0]
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    # 设置字体
    subtitle_run.font.name = 'Times New Roman'  # 先设置英文字体
    # 设置中文字体
    subtitle_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    for i, error in enumerate(errors):
        p = doc.add_paragraph()
        error_run = p.add_run(f"{i+1}. {error.get('message', '')}")
        error_run.bold = True
        # 设置字体
        error_run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        error_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        if 'location' in error and error['location']:
            location_run = p.add_run(f" (位置: {error['location']})")
            location_run.italic = True
            # 设置字体
            location_run.font.name = 'Times New Roman'  # 先设置英文字体
            # 设置中文字体
            location_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 按段落分组错误
    errors_by_para = {}
    for error in errors:
        location = error.get('location', '')
        if location and '.' in location:
            # 提取段落内容的前几个字作为标识
            parts = location.split('.', 1)
            error_type = parts[0].strip()
            para_identifier = parts[1].strip()

            if para_identifier not in errors_by_para:
                errors_by_para[para_identifier] = []

            # 添加错误类型到错误信息中
            error_with_type = error.copy()
            if 'error_type' not in error_with_type:
                error_with_type['error_type'] = error_type

            errors_by_para[para_identifier].append(error_with_type)

    # 处理段落错误
    for para in doc.paragraphs:
        para_text = para.text
        if not para_text.strip():
            continue

        # 查找匹配的错误
        para_errors = []
        for para_id, err_list in errors_by_para.items():
            # 使用前30个字符作为匹配依据
            para_start = para_text[:30] if len(para_text) > 30 else para_text
            if para_start.startswith(para_id) or para_id.startswith(para_start):
                para_errors = err_list
                break

        if not para_errors:
            continue

        # 清除段落中的所有运行
        try:
            p = para._p
            for r in para.runs:
                p.remove(r._r)
        except Exception as e:
            print(f"清除段落运行时出错: {str(e)}")
            # 使用替代方法：创建新段落并替换
            new_para = doc.add_paragraph()
            doc._body._body.insert_before(new_para._p, para._p)
            doc._body._body.remove(para._p)
            para = new_para

        # 添加红色标记的文本
        run = para.add_run(para_text)
        run.font.color.rgb = RGBColor(255, 0, 0)
        # 设置字体
        run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加错误说明
        error_text = " （"
        for i, error in enumerate(para_errors):
            error_message = error.get('message', '')
            error_type = error.get('error_type', '')

            # 解析错误消息，提取错误类型、要求值和实际值
            parsed_error = parse_error_message(error_message)
            if parsed_error[0]:
                error_type, required_value, actual_value = parsed_error
                error_text += f"{error_type}应为{required_value}，实际为{actual_value}"
            else:
                error_text += error_message

            if i < len(para_errors) - 1:
                error_text += "；"

        error_text += "）"
        error_run = para.add_run(error_text)
        error_run.font.color.rgb = RGBColor(255, 0, 0)
        error_run.italic = True
        # 设置字体
        error_run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        error_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 处理表格错误
    table_errors = [error for error in errors if '表' in error.get('location', '')]
    if table_errors:
        # 添加表格错误摘要
        summary_para = doc.add_paragraph()
        title_run = summary_para.add_run("表格错误摘要：")
        title_run.bold = True
        # 设置字体
        title_run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for error in table_errors:
            error_run = summary_para.add_run(f"\n- {error.get('location', '')}: {error.get('message', '')}")
            error_run.font.color.rgb = RGBColor(255, 0, 0)
            # 设置字体
            error_run.font.name = 'Times New Roman'  # 先设置英文字体
            # 设置中文字体
            error_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 处理图片错误
    image_errors = [error for error in errors if '图' in error.get('location', '')]
    if image_errors:
        # 添加图片错误摘要
        summary_para = doc.add_paragraph()
        title_run = summary_para.add_run("图片错误摘要：")
        title_run.bold = True
        # 设置字体
        title_run.font.name = 'Times New Roman'  # 先设置英文字体
        # 设置中文字体
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for error in image_errors:
            error_run = summary_para.add_run(f"\n- {error.get('location', '')}: {error.get('message', '')}")
            error_run.font.color.rgb = RGBColor(255, 0, 0)
            # 设置字体
            error_run.font.name = 'Times New Roman'  # 先设置英文字体
            # 设置中文字体
            error_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 确定输出路径
    if output_path is None:
        # 使用与原文档相同的目录
        output_dir = os.path.dirname(doc_path)
        output_path = os.path.join(output_dir, f"marked_{os.path.basename(doc_path)}")

    # 保存文档
    try:
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"文档已成功标记并保存到: {output_path}")
    except Exception as e:
        print(f"保存标记文档时出错: {str(e)}")
        # 尝试使用临时目录
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_output_path = os.path.join(temp_dir, f"marked_{os.path.basename(doc_path)}")
        doc.save(temp_output_path)
        print(f"文档已保存到临时目录: {temp_output_path}")
        output_path = temp_output_path

    return output_path

def parse_error_message(error_message: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """
    解析错误消息，提取错误类型和要求值

    Args:
        error_message: 错误消息

    Returns:
        Tuple[Optional[str], Optional[str], Optional[str]]: 错误类型、要求值和实际值
    """
    # 匹配格式：'错误类型' 不匹配: 要求 要求值, 实际 实际值
    pattern = r"'([^']+)'\s+不匹配:\s+要求\s+([^,]+),\s+实际\s+(.+)"
    match = re.match(pattern, error_message)

    if match:
        error_type = match.group(1)
        required_value = match.group(2)
        actual_value = match.group(3)
        return error_type, required_value, actual_value

    return None, None, None
