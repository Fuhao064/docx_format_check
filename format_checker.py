import docx 
import extract_para_info, docx_parser
import json
from typing import Dict
from format_agent import LLMs
from para_type import ParagraphManager,ParsedParaType
import os,concurrent,re
def load_config(config_path: str) -> Dict:
    """加载JSON配置文件"""
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"配置文件 {config_path} 未找到，将创建新文件")
        return {}

def save_config(config_path: str, config: Dict) -> None:
    """保存配置文件"""
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=4)

def update_config(config_path: str, new_config: Dict) -> None:
    """深度更新配置文件"""
    try:
        current_config = load_config(config_path)
        # 实现深度更新（示例为简单实现，复杂情况需要递归函数）
        for key, value in new_config.items():
            if isinstance(value, dict) and key in current_config:
                current_config[key].update(value)
            else:
                current_config[key] = value
        save_config(config_path, current_config)
        print("配置文件更新成功！")
    except json.JSONDecodeError:
        print(f"配置文件 {config_path} 格式错误")
    except Exception as e:
        print(f"发生错误：{e}")
def extract_number(value_str):
    """从字符串中提取数值"""
    try:
        # 使用正则表达式提取数字（包括小数点）
        number = re.findall(r'[\d.]+', str(value_str))[0]
        return float(number)
    except (IndexError, ValueError):
        return None

def is_value_equal(expected_str: str, actual_str: str, key: str) -> bool:
    """比较两个值是否相等，考虑特殊情况"""
    # 转换为小写进行比较
    expected_lower = str(expected_str).lower()
    actual_lower = str(actual_str).lower()
    
    # 如果完全相同（忽略大小写），直接返回True
    if expected_lower == actual_lower:
        return True
    
    # 对于方向值的特殊处理
    if key == 'orientation':
        return expected_lower == actual_lower
        
    # 提取数值进行比较
    expected_value = extract_number(expected_str)
    actual_value = extract_number(actual_str)
    
    if expected_value is not None and actual_value is not None:
        # 对于边距值，允许0.01厘米的误差
        margin_keys = ('top', 'bottom', 'left', 'right')
        if any(margin_key in key.lower() for margin_key in margin_keys):
            return abs(expected_value - actual_value) <= 0.01
            
    return False

def check_paper_format(doc_info: Dict, required_format: Dict) -> None:
    paper_required_format = required_format.get('paper', {})
    errors = []
    
    # 遍历要求格式中的每个字段
    for key, expected in paper_required_format.items():
        # 检查字段是否存在
        if key not in doc_info:
            errors.append(f"缺少必需字段: '{key}'")
            continue
            
        actual = doc_info[key]
        
        # 递归处理嵌套字典
        if isinstance(expected, dict):
            if not isinstance(actual, dict):
                errors.append(f"字段类型不匹配: '{key}' 应为字典类型，实际为 {type(actual).__name__}")
                continue
                
            # 递归调用检查嵌套字段
            nested_errors = check_paper_format(actual, {"paper": expected})
            errors.extend([f"{key}.{err}" for err in nested_errors])
                
        # 检查值是否匹配
        else:
            if not is_value_equal(expected, actual, key):
                errors.append(f"'{key}' 不匹配: 要求 {expected}, 实际 {actual}")
    
    # 打印错误信息
    if errors:
        print("\n页面格式检查错误:")
        for err in errors:
            print(f"- {err}")
    
    return errors

def remark_para_type(doc_path: str, llm: LLMs) -> ParagraphManager:
    # 初始化段落管理器
    paragraph_manager = ParagraphManager()
    
    # 读取doc的段落信息
    paragraph_manager = extract_para_info.extract_para_format_info(doc_path, paragraph_manager)
    
    # 文档整个内容
    doc_content = docx_parser.extract_doc_content(doc_path)
    
    # 定义任务函数
    def process_paragraph(para_index: int, para: Dict, doc_content_sent: bool, manager: ParagraphManager):
        try:
            if para["type"] != ParsedParaType.BODY:
                pass
            # 提取当前段落信息
            para_string = para["content"]
            
            # 调用大模型预测类型
            response = llm.predict_location(doc_content, para_string, doc_content_sent)
            response = response.strip()  # 去除首尾空白字符
            
            # 解析 JSON 字符串
            response_dict = json.loads(response)
            
            # 将返回的location转换为ParsedParaType枚举类型
            try:
                new_type = ParsedParaType(response_dict["location"])
                # 直接访问paragraphs列表中的对应索引
                manager.paragraphs[para_index].type = new_type
            except ValueError:
                print(f"Invalid paragraph type received: {response_dict['location']}")
                manager.paragraphs[para_index].type = ParsedParaType.BODY
            
        except Exception as e:
            # 如果解析失败，将段落类型设置为 BODY
            print(f"Error processing paragraph: {para_string[:50]}... Error: {e}")
            manager.paragraphs[para_index].type = ParsedParaType.BODY

    # 将段落json转为中文
    paras_info_json_zh = paragraph_manager.to_chinese_dict()
    
    # 使用线程池并发处理
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = []
        doc_content_sent = False
        
        # 提交任务到线程池，同时传入段落索引
        for i, para in enumerate(paras_info_json_zh):
            future = executor.submit(process_paragraph, i, para, doc_content_sent, paragraph_manager)
            futures.append(future)
            doc_content_sent = True  # 更新全文标志
        
        # 等待所有任务完成
        for future in concurrent.futures.as_completed(futures):
            try:
                future.result()  # 确保任务完成
            except Exception as e:
                print(f"Error in thread pool task: {e}")
    
    return paragraph_manager
    
def check_main_format(paragraph_manager:ParagraphManager, required_format:dict):
    errors = []
    paras_format_info = paragraph_manager.to_dict()
    for para in paras_format_info:
        # 和required_format进行对比
        para_type = para["type"]
        para_meta = para["meta"]
        required_meta = required_format.get(para_type)
        # 遍历要求格式中的每个字段
    for key, expected in required_meta.items():
        # 检查字段是否存在
        if key not in para_meta:
            errors.append(f"缺少必需字段: '{key}'")
            continue
            
        actual = para_meta[key]
        
        # 递归处理嵌套字典
        if isinstance(expected, dict):
            if not isinstance(actual, dict):
                errors.append(f"字段类型不匹配: '{key}' 应为字典类型，实际为 {type(actual).__name__}")
                continue
                
            # 递归调用检查嵌套字段
            nested_errors = check_paper_format(actual, {"paper": expected})
            errors.extend([f"{key}.{err}" for err in nested_errors])
                
        # 检查值是否匹配
        else:
            if not is_value_equal(expected, actual, key):
                errors.append(f"'{key}' 不匹配: 要求 {expected}, 实际 {actual}")
    
    # 打印错误信息
    if errors:
        print("\n页面格式检查错误:")
        for err in errors:
            print(f"- {err}")
    return 




def check_format(doc_path:str, required_format:dict, llm:LLMs):
    doc_info = docx_parser.extract_section_info(doc_path)
    # 初始化段落管理器
    paragraph_manager = ParagraphManager()
    # 读取doc的段落信息
    paragraph_manager = extract_para_info.extract_para_format_info(doc_path, paragraph_manager)
    # 重分配段落类型
    # paragraph_manager = remark_para_type(doc_path, llm)
    paragraph_manager = ParagraphManager.build_from_json_file("result.json")
    # paras_info_json_zh = paragraph_manager.to_chinese_dict()
    # 对齐段落信息和格式信息同Config
    # print(paras_info_dict)
    # 检查格式是否符合要求
    check_paper_format(doc_info, required_format)
    #  检查段落格式
    check_main_format(paragraph_manager, required_format)
    #  检查字体格式
    # check_run_format(runs_info, required_format)
    # # 检查表格样式
    # check_table_format(doc_path, required_format)
    # # 检查图片样式
    # check_image_format(doc_path, required_format)
    # # 检查引用样式和参考文献
    # check_reference_format(doc_path, required_format)

llm = LLMs()
llm.set_model('deepseek-r1')
required_format = load_config('config.json')
check_format('test.docx', required_format, llm)
